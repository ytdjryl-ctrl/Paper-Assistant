from __future__ import annotations

import csv
import json
import re
import zipfile
from pathlib import Path
from typing import Iterable, List, Tuple
from xml.etree import ElementTree

from .models import SourceFile


TEXT_SUFFIXES = {".txt", ".md", ".py", ".yaml", ".yml", ".json", ".csv", ".tsv", ".log", ".xml", ".html", ".htm"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"}


def _read_text(path: Path) -> Tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding), f"text:{encoding}"
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore"), "text:utf-8-ignore"


def _extract_pdf(path: Path) -> Tuple[str, str]:
    errors = []
    try:
        import fitz  # type: ignore

        with fitz.open(str(path)) as document:
            text = "\n\n".join(page.get_text() for page in document)
        if text.strip():
            return text, "pdf:pymupdf"
    except Exception as exc:  # pragma: no cover - backend availability is environment-specific
        errors.append(f"PyMuPDF: {exc}")

    try:
        from pdfminer.high_level import extract_text  # type: ignore

        text = extract_text(str(path)) or ""
        if text.strip():
            return text, "pdf:pdfminer"
    except Exception as exc:  # pragma: no cover
        errors.append(f"pdfminer: {exc}")

    try:
        import PyPDF2  # type: ignore

        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text, "pdf:pypdf2"
    except Exception as exc:  # pragma: no cover
        errors.append(f"PyPDF2: {exc}")

    raise RuntimeError("; ".join(errors) or "no PDF text extractor returned content")


def _xml_text_from_zip(path: Path, member_pattern: str) -> str:
    chunks: List[str] = []
    with zipfile.ZipFile(path) as archive:
        members = sorted(name for name in archive.namelist() if re.search(member_pattern, name))
        for name in members:
            root = ElementTree.fromstring(archive.read(name))
            current: List[str] = []
            for element in root.iter():
                tag = element.tag.rsplit("}", 1)[-1]
                if tag == "t" and element.text:
                    current.append(element.text)
                elif tag in {"p", "tr"} and current:
                    chunks.append(" ".join(current).strip())
                    current = []
            if current:
                chunks.append(" ".join(current).strip())
    return "\n".join(chunk for chunk in chunks if chunk)


def _extract_docx(path: Path) -> Tuple[str, str]:
    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    lines.append(" | ".join(values))
        text = "\n".join(lines)
        if text.strip():
            return text, "docx:python-docx"
    except Exception:
        pass

    text = _xml_text_from_zip(path, r"^word/(document|footnotes|endnotes)\.xml$")
    if not text.strip():
        raise RuntimeError("DOCX did not contain extractable text")
    return text, "docx:openxml"


def _extract_pptx(path: Path) -> Tuple[str, str]:
    text = _xml_text_from_zip(path, r"^ppt/slides/slide\d+\.xml$")
    if not text.strip():
        raise RuntimeError("PPTX did not contain extractable slide text")
    return text, "pptx:openxml"


def _extract_xlsx(path: Path, max_rows_per_sheet: int = 500, max_columns: int = 40) -> Tuple[str, str]:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("openpyxl is required for XLSX extraction") from exc

    workbook = load_workbook(filename=str(path), read_only=True, data_only=False)
    lines: List[str] = []
    try:
        for sheet in workbook.worksheets:
            lines.append(f"## Sheet: {sheet.title}")
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), 1):
                values = ["" if value is None else str(value) for value in row[:max_columns]]
                if any(values):
                    lines.append(" | ".join(values))
                if row_index >= max_rows_per_sheet:
                    lines.append(f"[sheet truncated after {max_rows_per_sheet} rows]")
                    break
    finally:
        workbook.close()
    return "\n".join(lines), "xlsx:openpyxl"


def _extract_image(path: Path) -> Tuple[str, str]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore

        text = pytesseract.image_to_string(Image.open(str(path)), lang="chi_sim+eng")
        if text.strip():
            return text, "image:tesseract"
    except Exception:
        pass
    return f"[Image evidence: {path.name}. OCR text was not available; inspect the original image before making visual claims.]", "image:metadata"


def extract_file(path: Path) -> Tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        if suffix == ".json":
            raw, method = _read_text(path)
            try:
                return json.dumps(json.loads(raw), ensure_ascii=False, indent=2), method
            except Exception:
                return raw, method
        if suffix in {".csv", ".tsv"}:
            delimiter = "\t" if suffix == ".tsv" else ","
            rows = []
            with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
                for row_index, row in enumerate(csv.reader(handle, delimiter=delimiter), 1):
                    rows.append(" | ".join(row[:40]))
                    if row_index >= 1000:
                        rows.append("[table truncated after 1000 rows]")
                        break
            return "\n".join(rows), f"table:{suffix.lstrip('.')}"
        return _read_text(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pptx":
        return _extract_pptx(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    if suffix in IMAGE_SUFFIXES:
        return _extract_image(path)
    if suffix in {".doc", ".xls", ".ppt"}:
        raise RuntimeError(f"legacy {suffix} extraction is not supported; convert it to an Office Open XML format")
    raise RuntimeError(f"unsupported file type: {suffix or '(no extension)'}")


def extract_source_files(
    source_files: Iterable[SourceFile],
    workspace_path: Path,
    max_chars_per_file: int = 50000,
) -> List[str]:
    evidence_dir = workspace_path / "evidence" / "extracted"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    warnings: List[str] = []
    manifest = []

    for index, source in enumerate(source_files, 1):
        try:
            text, method = extract_file(source.path)
            text = text.strip()
            source.truncated = len(text) > max_chars_per_file
            source.extracted_text = text[:max_chars_per_file]
            source.text_preview = source.extracted_text[:2500]
            source.extraction_method = method
            safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", source.path.stem)[:80] or "source"
            extracted_path = evidence_dir / f"{index:03d}_{safe_stem}.txt"
            extracted_path.write_text(source.extracted_text + "\n", encoding="utf-8")
            source.extracted_path = extracted_path.relative_to(workspace_path).as_posix()
        except Exception as exc:
            source.extraction_error = str(exc)
            warnings.append(f"Could not extract {source.rel_path}: {exc}")

        manifest.append(
            {
                "source_id": f"S{index}",
                "path": source.rel_path,
                "kind": source.kind,
                "size_bytes": source.size_bytes,
                "extraction_method": source.extraction_method,
                "extracted_path": source.extracted_path,
                "extracted_chars": len(source.extracted_text),
                "truncated": source.truncated,
                "error": source.extraction_error,
            }
        )

    (workspace_path / "evidence" / "evidence_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return warnings

