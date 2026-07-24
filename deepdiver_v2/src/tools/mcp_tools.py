# Copyright (c) 2025 Huawei Technologies Co., Ltd. All rights reserved.
# Copyright (c) 2026 South China Sea Institute of Oceanology, Chinese Academy of Sciences (SCSIO, CAS). All rights reserved.
import os
import json
import random
import subprocess
import requests
import re
import shutil
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import logging
from dataclasses import dataclass
from urllib.parse import urlparse
import tempfile
import time
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from datetime import datetime, timedelta
import dateutil.parser
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from urllib.parse import quote
from collections import Counter
import inspect
import sys
from functools import wraps
from typing import Optional

import feedparser
import urllib3
from .paper import Paper
from .normalizer import Area, CompanyStatus, DateRange, normalize_company_name
from src.utils.writing_profile import audit_manuscript_text, section_profile_text

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# from markdown_pdf import MarkdownPdf, Section  # 鏀圭敤 ReportLab

# ReportLab imports for PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted, \
        Image as RLImage, HRFlowable
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("ReportLab is not installed. Run: pip install reportlab")

# 灏濊瘯瀵煎叆matplotlib鐢ㄤ簬娓叉煋鏁板鍏紡
try:
    import matplotlib

    matplotlib.use('Agg')  # 浣跨敤闈濭UI鍚庣
    import matplotlib.pyplot as plt
    from matplotlib import mathtext

    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    logger.warning("matplotlib is not installed. Math formulas will be rendered as plain text. Run: pip install matplotlib")

try:
    from config.config import get_config
except ImportError:
    import sys

    sys.path.append(str(Path(__file__).parent.parent.parent))
    from config.config import get_config

# Import the optimized Faiss-based manager (fallback to JSON if Faiss not available)
try:
    from knowledge.vector_store import auto_index_task_completion_optimized, get_optimized_knowledge_manager

    FAISS_AVAILABLE = True
except ImportError:
    try:
        from ..knowledge.vector_store import auto_index_task_completion_optimized, get_optimized_knowledge_manager

        FAISS_AVAILABLE = True
    except ImportError:
        # Knowledge module not available, provide stub implementations
        FAISS_AVAILABLE = False


        def auto_index_task_completion_optimized(config, task_summary):
            # Stub implementation - knowledge module not available
            logging.getLogger(__name__).debug("Knowledge indexing skipped - module not available")
            return True


        def auto_index_task_completion(config, task_summary):
            # Stub implementation - knowledge module not available
            logging.getLogger(__name__).debug("Knowledge indexing skipped - module not available")
            return True


        get_optimized_knowledge_manager = None

logger = logging.getLogger(__name__)

proxy = {}


@dataclass
class MCPToolResult:
    # Standard result format for MCP tools
    success: bool
    data: Any = None
    error: str = None
    metadata: Dict[str, Any] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "data": self.data,
            "error": self.error,
            "metadata": self.metadata or {}
        }


def _render_latex_to_image(latex_text: str, output_path: Path = None, fontsize: int = 12) -> Path:
    """
    浣跨敤matplotlib灏哃aTeX鍏紡娓叉煋涓哄浘鐗?

    Args:
        latex_text: LaTeX鏍煎紡鐨勬暟瀛叕寮?
        output_path: 杈撳嚭鍥剧墖璺緞,濡傛灉涓篘one鍒欒嚜鍔ㄧ敓鎴?
        fontsize: 瀛椾綋澶у皬

    Returns:
        鍥剧墖鏂囦欢璺緞
    """
    if not MATPLOTLIB_AVAILABLE:
        return None

    try:
        # 鍒涘缓涓存椂鏂囦欢璺緞
        if output_path is None:
            import hashlib
            hash_name = hashlib.md5(latex_text.encode()).hexdigest()
            output_path = Path(tempfile.gettempdir()) / f"latex_{hash_name}.png"

        # 濡傛灉鍥剧墖宸插瓨鍦?鐩存帴杩斿洖
        if output_path.exists():
            return output_path

        # 纭繚LaTeX鏂囨湰琚?鍖呰９
        if not latex_text.startswith('$'):
            latex_text = f'${latex_text}$'

        # 鍒涘缓鍥惧舰
        fig = plt.figure(figsize=(8, 1))
        fig.patch.set_facecolor('white')

        # 娓叉煋LaTeX
        text = fig.text(0.5, 0.5, latex_text,
                        fontsize=fontsize,
                        ha='center',
                        va='center',
                        usetex=False)  # 浣跨敤matplotlib鐨刴athtext鑰屼笉鏄湡姝ｇ殑LaTeX

        # 璋冩暣杈圭晫
        fig.tight_layout(pad=0.1)

        # 淇濆瓨涓篜NG
        plt.savefig(output_path,
                    dpi=150,
                    bbox_inches='tight',
                    pad_inches=0.05,
                    facecolor='white',
                    edgecolor='none')
        plt.close(fig)

        return output_path

    except Exception as e:
        logger.error(f"璀憡: LaTeX娓叉煋澶辫触: {e}")
        return None


def _wrap_special_symbol(symbol: str, fallback: str = None) -> str:
    # 涓虹壒娈婄鍙锋坊鍔犲瓧浣撳洖閫€鏀寔
    if fallback:
        return f'<font name="Arial">{symbol}</font>'
    return symbol


def _simplify_latex(latex_text: str) -> str:
    """
    绠€鍖朙aTeX鏁板鍏紡涓哄彲璇绘枃鏈?
    灏嗗父瑙佺殑LaTeX鍛戒护杞崲涓篣nicode鏁板绗彿

    Args:
        latex_text: LaTeX鏍煎紡鐨勬暟瀛叕寮?

    Returns:
        绠€鍖栧悗鐨勬枃鏈?
    """
    # 甯歌LaTeX鍛戒护鏄犲皠 - 浣跨敤鏈夊簭瀛楀吀纭繚澶勭悊椤哄簭
    # 閲嶈锛氬繀椤诲厛澶勭悊闀垮懡浠わ紝鍐嶅鐞嗙煭鍛戒护锛岄伩鍏嶉儴鍒嗗尮閰?
    replacements = [
        # 甯岃厞瀛楁瘝锛堟寜瀛楁瘝椤哄簭锛?
        (r'\\alpha', '伪'),
        (r'\\beta', '尾'),
        (r'\\gamma', '纬'),
        (r'\\Gamma', '螕'),
        (r'\\delta', '未'),
        (r'\\Delta', '螖'),
        (r'\\epsilon', '蔚'),
        (r'\\varepsilon', '蔚'),
        (r'\\zeta', '味'),
        (r'\\eta', '畏'),
        (r'\\theta', '胃'),
        (r'\\Theta', '螛'),
        (r'\\iota', '喂'),
        (r'\\kappa', '魏'),
        (r'\\lambda', '位'),
        (r'\\Lambda', '螞'),
        (r'\\mu', '渭'),
        (r'\\nu', '谓'),
        (r'\\xi', '尉'),
        (r'\\Xi', '螢'),
        (r'\\pi', '蟺'),
        (r'\\Pi', '螤'),
        (r'\\rho', '蟻'),
        (r'\\sigma', '蟽'),
        (r'\\Sigma', '危'),
        (r'\\tau', '蟿'),
        (r'\\upsilon', '蠀'),
        (r'\\Upsilon', '违'),
        (r'\\phi', '蠁'),
        (r'\\Phi', '桅'),
        (r'\\varphi', '蠁'),
        (r'\\chi', '蠂'),
        (r'\\psi', '蠄'),
        (r'\\Psi', '唯'),
        (r'\\omega', '蠅'),
        (r'\\Omega', '惟'),

        # 鏁板杩愮畻绗?
        # 娉ㄦ剰锛氭煇浜涚鍙峰湪瀹嬩綋涓彲鑳芥棤娉曟樉绀猴紝浣跨敤澶囬€夋柟妗?
        # nabla: 浣跨敤鍊掍笁瑙掑舰 鈻?(U+25BD) 浠ｆ浛 鈭?(U+2207)锛屽洜涓哄悗鑰呭湪鏌愪簺瀛椾綋涓己澶?
        (r'\\nabla', '\\u25bd'),
        (r'\\partial', 'd'),
        (r'\\infty', '\\u221e'),

        # Relation symbols
        (r'\\leq', '<='),
        (r'\\le', '<='),
        (r'\\geq', '>='),
        (r'\\ge', '>='),
        (r'\\neq', '!='),
        (r'\\ne', '!='),
        (r'\\approx', '~'),
        (r'\\equiv', '=='),
        (r'\\sim', '~'),
        (r'\\propto', '\\u221d'),

        # Binary operators
        (r'\\times', 'x'),
        (r'\\cdot', '*'),
        (r'\\div', '/'),
        (r'\\pm', '+/-'),
        (r'\\mp', '-/+'),
        (r'\\oplus', '+'),
        (r'\\otimes', 'x'),

        # Set symbols
        (r'\\notin', 'not in'),
        (r'\\subseteq', 'subseteq'),
        (r'\\supseteq', 'supseteq'),
        (r'\\subset', 'subset'),
        (r'\\supset', 'supset'),
        (r'\\in', 'in'),
        (r'\\cup', 'U'),
        (r'\\cap', 'cap'),
        (r'\\emptyset', 'empty'),

        # Logic symbols
        (r'\\forall', 'forall'),
        (r'\\exists', 'exists'),
        (r'\\neg', 'not'),
        (r'\\land', 'and'),
        (r'\\lor', 'or'),

        # Large operators
        (r'\\sum', 'sum'),
        (r'\\prod', 'prod'),
        (r'\\iiint', 'int'),
        (r'\\iint', 'int'),
        (r'\\oint', 'int'),
        (r'\\int', 'int'),

        # Arrows
        (r'\\rightarrow', '->'),
        (r'\\to', '->'),
        (r'\\leftarrow', '<-'),
        (r'\\gets', '<-'),
        (r'\\leftrightarrow', '<->'),
        (r'\\Rightarrow', '=>'),
        (r'\\Leftarrow', '<='),
        (r'\\Leftrightarrow', '<=>'),

        # Other common symbols
        (r'\\sqrt', 'sqrt'),
        (r'\\angle', 'angle'),
        (r'\\perp', 'perp'),
        (r'\\parallel', 'parallel'),

        # 涓婁笅鏍囧拰鐗规畩鏍煎紡锛堝叿浣撴暟瀛楋級
        # 娉ㄦ剰锛氫笉瑕佸湪杩欓噷杞崲涓篣nicode涓婁笅鏍囧瓧绗紙鈧€鈧佲倐绛夛級锛屽洜涓哄畫浣撳彲鑳戒笉鏀寔
        # 搴旇鍦ㄥ悗闈㈢粺涓€杞崲涓篐TML <sub>/<sup> 鏍囩
        # 鍙繚鐣欏父鐢ㄧ殑骞虫柟銆佺珛鏂圭瓑鐗规畩绗彿
        (r'\^2', '虏'),
        (r'\^3', '鲁'),

        # 鏂囨湰鍛戒护
        (r'\\text\{([^}]+)\}', r'\1'),
        (r'\\mathrm\{([^}]+)\}', r'\1'),
        (r'\\mathbf\{([^}]+)\}', r'\1'),
        (r'\\mathit\{([^}]+)\}', r'\1'),

        # 鍒嗘暟(绠€鍖栨樉绀?
        (r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)'),

        # 甯藉瓙鍜屼慨楗扮鍙?
        # 娉ㄦ剰锛氱粍鍚堝瓧绗湪 ReportLab 涓彲鑳芥棤娉曟纭樉绀猴紝浣跨敤绠€鍖栬〃绀?
        (r'\\hat\{([^}]+)\}', r'\1^'),  # 浣跨敤涓婃爣^琛ㄧず甯藉瓙
        (r'\\hat ([a-zA-Z])', r'\1^'),
        (r'\\bar\{([^}]+)\}', r'\1_bar'),
        (r'\\tilde\{([^}]+)\}', r'\1~'),  # 浣跨敤娉㈡氮鍙?
        (r'\\vec\{([^}]+)\}', r'<b>\1</b>'),  # 鍚戦噺浣跨敤绮椾綋琛ㄧず锛堟爣鍑嗘暟瀛鍙凤級

        # 鎷彿
        (r'\\left\(', '('),
        (r'\\right\)', ')'),
        (r'\\left\[', '['),
        (r'\\right\]', ']'),
        (r'\\left\{', '{'),
        (r'\\right\}', '}'),
        (r'\\left', ''),
        (r'\\right', ''),
        (r'\\big', ''),
        (r'\\Big', ''),
        (r'\\bigg', ''),
        (r'\\Bigg', ''),

        # 绌烘牸
        (r'\\\\', ' '),  # 鎹㈣
        (r'\\,', ' '),  # 灏忕┖鏍?
        (r'\\;', ' '),  # 涓┖鏍?
        (r'\\quad', '  '),  # 澶х┖鏍?
        (r'\\qquad', '    '),  # 鏇村ぇ绌烘牸
    ]

    result = latex_text

    # 鎸夐『搴忓簲鐢ㄦ墍鏈夋浛鎹?
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result)

    # 澶勭悊閫氱敤鐨勪笂鏍?^{...} 杞崲涓篐TML
    def convert_superscript(match):
        content = match.group(1)
        return f'<sup>{content}</sup>'

    result = re.sub(r'\^\{([^}]+)\}', convert_superscript, result)
    result = re.sub(r'\^([0-9a-zA-Z])', r'<sup>\1</sup>', result)

    # 澶勭悊閫氱敤鐨勪笅鏍?_{...} 杞崲涓篐TML
    def convert_subscript(match):
        content = match.group(1)
        return f'<sub>{content}</sub>'

    result = re.sub(r'\_\{([^}]+)\}', convert_subscript, result)
    result = re.sub(r'\_([0-9a-zA-Z])', r'<sub>\1</sub>', result)

    # 娓呯悊鍓╀綑鐨勫弽鏂滄潬鍜岃姳鎷彿
    # 娉ㄦ剰锛氳繖閲岃灏忓績锛屼笉瑕佹竻鐞嗘帀宸茬粡杞崲濂界殑Unicode绗彿
    result = result.replace('\\', '')
    result = re.sub(r'\{([^}]*)\}', r'\1', result)

    # 鍚庡鐞嗭細涓哄彲鑳芥棤娉曟樉绀虹殑绗彿娣诲姞瀛椾綋鏍囪鎴栨浛鎹?
    # 杩欎簺绗彿鍦ㄥ畫浣撲腑鍙兘缂哄け锛岄渶瑕佺壒娈婂鐞?
    problematic_symbols = {}

    # 妫€鏌ュ苟鏍囪鐗规畩绗彿
    for symbol, (replacement, name) in problematic_symbols.items():
        if symbol in result and replacement is not None:
            # 濡傛灉鏈夋浛浠ｇ鍙凤紝鐩存帴鏇挎崲
            result = result.replace(symbol, replacement)

    return result


def _process_inline_formatting(text: str) -> str:
    """
    灏?Markdown 琛屽唴鏍煎紡杞崲涓?ReportLab 鍙В鏋愮殑瀹夊叏 HTML锛?
    淇濊瘉鏍囩骞宠　锛岄伩鍏?PDF 鐢熸垚鏃剁殑瑙ｆ瀽閿欒銆?
    """

    text = re.sub(r'<a\s+id="([^"]+)"', r'<a name="\1"', text)

    text = re.sub(
        r'<a\s+href="([^"]+)"\s+style="[^"]*">',
        r'<a href="\1" color="#04B5BB">',
        text
    )

    # 3. 涓烘病鏈夐鑹插睘鎬х殑 href 閾炬帴娣诲姞棰滆壊
    # 鍖归厤: <a href="..."> (浣嗕笉鍖归厤宸叉湁color灞炴€х殑)
    text = re.sub(
        r'<a\s+href="([^"]+)"(?!\s+color)>',
        r'<a href="\1" color="#04B5BB">',
        text
    )

    # 鎭㈠骞跺寮篣nicode涓婁笅鏍囧鐞?
    # 杩欎竴姝ラ潪甯稿叧閿紝鍥犱负鐢ㄦ埛缁忓父鐩存帴澶嶅埗绮樿创鍖呭惈Unicode涓婃爣锛堝 鈦宦光伓锛夌殑鏂囨湰
    # 鑰岃繖浜涘瓧绗湪鏍囧噯涓枃瀛椾綋锛堝瀹嬩綋锛変腑閫氬父涓嶆敮鎸侊紝瀵艰嚧鏄剧ず涓虹┖鐧?

    superscript_map = {'¹': '1', '²': '2', '³': '3'}
    subscript_map = {}

    # 杞崲Unicode涓婃爣涓篐TML sup鏍囩
    for unicode_char, normal_char in superscript_map.items():
        if unicode_char in text:
            text = text.replace(unicode_char, f'<sup>{normal_char}</sup>')

    # 杞崲Unicode涓嬫爣涓篐TML sub鏍囩
    for unicode_char, normal_char in subscript_map.items():
        if unicode_char in text:
            text = text.replace(unicode_char, f'<sub>{normal_char}</sub>')

    # 鐗规畩瀛楃澶勭悊
    text = text.replace('渭', '碌').replace('碌', '<font name="Arial">碌</font>')
    text = text.replace('欧', '<font name="Arial">欧</font>')

    # 澶勭悊瀹嬩綋涓嶆敮鎸佺殑鏁板绗彿锛屼娇鐢?Arial 瀛椾綋鏄剧ず锛圵indows 绯荤粺鑷甫锛?
    math_symbols = []
    for sym in math_symbols:
        if sym in text:
            text = text.replace(sym, f'<font name="Arial">{sym}</font>')

    # 鍏堜繚鎶ゆ暟瀛叕寮?$$ ... $$ (閬垮厤琚悗缁鐞嗙牬鍧?
    math_formulas = []

    def protect_math(match):
        formula = match.group(1)
        placeholder = f"__MATH_FORMULA_{len(math_formulas)}__"
        math_formulas.append(formula)
        return placeholder

    # 淇濇姢琛屽唴鏁板鍏紡 $...$
    # 浣跨敤闈炶椽濠尮閰?纭繚鍙尮閰嶆垚瀵圭殑$
    text = re.sub(r'\$([^\$]+?)\$', protect_math, text)

    # 娓呯悊瀛ょ珛鐨?绗彿(娌℃湁閰嶅鐨?
    # 缁熻鍓╀綑鐨?鏁伴噺,濡傛灉鏄鏁?璇存槑鏈夊绔嬬殑$
    dollar_count = text.count('$')
    if dollar_count > 0:
        # 绉婚櫎鎵€鏈夊墿浣欑殑瀛ょ珛$绗彿
        text = text.replace('$', '')

    # 澶勭悊鏂囩尞寮曠敤鏍煎紡锛堝繀椤诲湪鏅€歁arkdown閾炬帴涔嬪墠澶勭悊锛?
    # 鏍煎紡1: [鏁板瓧] 鏍囬/鏂囦欢鍚嶏紝URL.pdf锛屾棩鏈?- PDF鏂囦欢寮曠敤锛堟牴鎹甎RL鏄惁浠?pdf缁撳熬鍒ゆ柇锛?
    def replace_pdf_reference(match):
        num = match.group(1)
        title = match.group(2)  # 鏍囬鎴栨枃浠跺悕锛堝彲浠ヤ笉鍚?pdf锛?
        url = match.group(3)  # URL蹇呴』浠?pdf缁撳熬
        date = match.group(4) if len(match.groups()) >= 4 and match.group(4) else ''
        # 浣跨敤鍥炲舰閽堝浘鏍囸煋?(U+1F4CE) 琛ㄧず鍙笅杞界殑PDF鏂囨。
        # 澶勭悊鏃犳硶纭畾鏈堜唤鐨勬儏鍐碉紝鍙樉绀哄勾浠?
        if date and date.strip():
            date_str = date.strip()
            if '鏃犳硶纭畾鏈堜唤' in date_str:
                # 鎻愬彇骞翠唤锛堝尮閰?浣嶆暟瀛?骞达級
                year_match = re.search(r'(\d{4})', date_str)
                if year_match:
                    date_str = year_match.group(1)
            date_part = f', {date_str}'
        else:
            date_part = ''
        # 浣跨敤font鏍囩鎸囧畾emoji瀛椾綋鏉ユ樉绀哄浘鏍囷紝ReportLab浼氳嚜鍔ㄥ洖閫€鍒版敮鎸佽瀛楃鐨勫瓧浣?
        return f'[{num}] <font name="EmojiFont">馃搸</font> {title}, <a href="{url}" color="#04B5BB">{url}</a>{date_part}'

    # 鏀硅繘鐨勬鍒欒〃杈惧紡1锛氭牴鎹甎RL鏄惁浠?pdf缁撳熬鏉ュ垽鏂璓DF寮曠敤锛堜笉绠℃爣棰樻槸浠€涔堬級
    # 杩欐牱鍙互姝ｇ‘璇嗗埆鏍囬涓嶅惈.pdf浣哢RL鏄疨DF鐨勬儏鍐碉紝濡傦細[1] BD CD Marker Handbook, https://example.com/file.pdf
    text = re.sub(r'\[(\d+)\]\s*(.+?)锛?https?://[^\s锛宂+?\.pdf)(?:锛?.+?))?(?=\s*\n|\s*$)',
                  replace_pdf_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 鏍煎紡2: [鏁板瓧] 鏍囬锛孶RL锛屾棩鏈?- 缃戦〉寮曠敤
    def replace_web_reference(match):
        num = match.group(1)
        title = match.group(2)
        url = match.group(3)
        date = match.group(4) if len(match.groups()) >= 4 and match.group(4) else ''
        # 浣跨敤鍦扮悆浠浘鏍囸煂?(U+1F310) 琛ㄧず缃戦〉閾炬帴
        # 澶勭悊鏃犳硶纭畾鏈堜唤鐨勬儏鍐碉紝鍙樉绀哄勾浠?
        if date and date.strip():
            date_str = date.strip()
            if '鏃犳硶纭畾鏈堜唤' in date_str:
                # 鎻愬彇骞翠唤锛堝尮閰?浣嶆暟瀛?骞达級
                year_match = re.search(r'(\d{4})', date_str)
                if year_match:
                    date_str = year_match.group(1)
            date_part = f', {date_str}'
        else:
            date_part = ''
        # 浣跨敤font鏍囩鎸囧畾emoji瀛椾綋鏉ユ樉绀哄浘鏍囷紝ReportLab浼氳嚜鍔ㄥ洖閫€鍒版敮鎸佽瀛楃鐨勫瓧浣?
        return f'[{num}] <font name="EmojiFont">馃寪</font> {title}, <a href="{url}" color="#04B5BB">{url}</a>{date_part}'

    # 鏀硅繘鐨勬鍒欒〃杈惧紡2锛氬尮閰嶉潪PDF鐨刄RL寮曠敤锛堢綉椤靛紩鐢級
    text = re.sub(r'\[(\d+)\]\s*(.+?)锛?https?://[^\s锛宂+?)(?:锛?.+?))?(?=\s*\n|\s*$)',
                  replace_web_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 澶勭悊 Markdown 閾炬帴 [text](url) (蹇呴』鍦ㄥ叾浠栨牸寮忎箣鍓嶅鐞?
    def replace_markdown_link(match):
        link_text = match.group(1)
        link_url = match.group(2)
        # ReportLab 鐨?<a> 鏍囩鏍煎紡
        return f'<a href="{link_url}" color="#04B5BB">{link_text}</a>'

    # 澶勭悊Markdown閾炬帴锛屾敮鎸乁RL涓寘鍚嫭鍙凤紙濡倃iki閾炬帴銆佽鏂嘍OI绛夛級
    # 鍖归厤妯″紡锛歎RL鍙寘鍚竴灞傚祵濂楁嫭鍙凤紝濡?https://example.com/page(1).html
    text = re.sub(r'\[([^\]]+)\]\(([^()\s]*(?:\([^)]*\)[^()\s]*)*)\)', replace_markdown_link, text)

    # 澶勭悊琛屽唴浠ｇ爜 `code` (蹇呴』鍦ㄧ矖浣撳拰鏂滀綋涔嬪墠澶勭悊,閬垮厤鍐茬獊)
    # 浣跨敤鍗犱綅绗繚鎶や唬鐮佸潡锛岄槻姝㈠悗缁殑绮椾綋/鏂滀綋姝ｅ垯璇尮閰嶄唬鐮佸潡鍐呯殑HTML鏍囩鎴栧唴瀹?
    inline_codes = []

    def protect_code(match):
        code_content = match.group(1)
        # HTML 杞箟
        code_content = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        # 妫€鏌ユ槸鍚寘鍚腑鏂囧瓧绗?
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', code_content))

        if has_chinese:
            # 鍖呭惈涓枃,浣跨敤瀹嬩綋,绋嶅皬瀛楀彿,娣诲姞鐏拌壊鑳屾櫙
            html = f'<font name="SimSun" size="9" color="#333333" backColor="#f5f5f5">{code_content}</font>'
        else:
            # 绾嫳鏂?鏁板瓧,浣跨敤 Arial 瀛椾綋
            html = f'<font name="Arial" size="9">{code_content}</font>'

        placeholder = f"__INLINE_CODE_{len(inline_codes)}__"
        inline_codes.append(html)
        return placeholder

    text = re.sub(r'`([^`]+)`', protect_code, text)

    # 绮楁枩浣擄細浼樺厛澶勭悊 ***text***
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<font name="SimHei"><b><i>\1</i></b></font>', text)
    # 绮椾綋锛氶潪璐┆鍖归厤锛屽厑璁镐腑闂村寘鍚叾浠栧瓧绗紙濡傛枩浣撶殑*锛?
    text = re.sub(r'\*\*(.+?)\*\*', r'<font name="SimHei"><b>\1</b></font>', text)
    # 鏂滀綋锛氶潪璐┆鍖归厤锛屾帓闄?*鐨勬儏鍐碉紝涓旇姹傚唴瀹逛袱渚ф棤绌虹櫧锛堥伩鍏嶅尮閰嶆暟瀛叕寮忎腑鐨?锛?
    text = re.sub(r'(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)', r'<i>\1</i>', text)

    for i, html in enumerate(inline_codes):
        text = text.replace(f"__INLINE_CODE_{i}__", html)
    for i, formula in enumerate(math_formulas):
        display = _simplify_latex(formula)
        text = text.replace(f"__MATH_FORMULA_{i}__", f'<font size="9.5"><i>{display}</i></font>')

    # 绉婚櫎涓嶆敮鎸佹垨鏃犳剰涔夌殑鏍囩
    text = re.sub(r'</?\s*nobr\b[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<hr\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<div\b[^>]*>(.*?)</div>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?div\b[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<span\b[^>]*>(.*?)</span>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?span\b[^>]*>', '', text, flags=re.IGNORECASE)
    for tag in ['section', 'article', 'header', 'footer', 'nav', 'aside', 'main']:
        text = re.sub(rf'<{tag}\b[^>]*>(.*?)</{tag}>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(rf'</?\s*{tag}\b[^>]*>', '', text, flags=re.IGNORECASE)

    text = re.sub(r'<br\s*>', '<br/>', text, flags=re.IGNORECASE)
    text = re.sub(r'<br\s*>([^<]*)</br>', r'<br/> \1', text, flags=re.IGNORECASE)
    text = re.sub(r'</br>', '', text, flags=re.IGNORECASE)
    text = text.replace('<br/>', ' <br/> ')

    for attr in ['color', 'size', 'name', 'href', 'face', 'backColor']:
        text = re.sub(rf'\b{attr}=([^"\s>]+)', rf'{attr}="\1"', text)

    text = re.sub(r'<(font|b|i|sub|sup)\b[^>]*>\s*</\1>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<a(?![^>]*\bname=)[^>]*>\s*</a>', '', text, flags=re.IGNORECASE)

    text = _apply_english_font_markup(text)

    # 澶勭悊鐗规畩绗彿锛岀‘淇濅娇鐢ㄦ敮鎸佽繖浜涚鍙风殑瀛椾綋锛屽苟璁剧疆棰滆壊
    # 鍕惧彿璁句负缁胯壊
    green_symbols = []
    for sym in green_symbols:
        text = text.replace(sym, f'<font name="SymbolFont" color="green">{sym}</font>')

    # 鍙夊彿璁句负绾㈣壊
    red_symbols = []
    for sym in red_symbols:
        text = text.replace(sym, f'<font name="SymbolFont" color="red">{sym}</font>')

    # 骞宠　鍐呰仈鏍囩锛屼慨姝ｇ己澶辨垨澶氫綑鐨勯棴鍚?
    tag_regex = re.compile(r'</?\s*(a|font|b|i|sub|sup)\b[^>]*>', re.IGNORECASE)

    def balance_inline_tags(value: str) -> str:
        parts: List[str] = []
        stack: List[str] = []
        last = 0
        for m in tag_regex.finditer(value):
            parts.append(value[last:m.start()])
            token = m.group(0)
            name = m.group(1).lower()
            closing = token.startswith('</')
            if closing:
                if name in stack:
                    while stack and stack[-1] != name:
                        parts.append(f'</{stack.pop()}>')
                    if stack and stack[-1] == name:
                        stack.pop()
                        parts.append(token)
                # 鏈尮閰嶇殑瀛ょ珛闂悎鐩存帴涓㈠純
            else:
                stack.append(name)
                parts.append(token)
            last = m.end()
        parts.append(value[last:])
        while stack:
            parts.append(f'</{stack.pop()}>')
        return ''.join(parts)

    return balance_inline_tags(text)


_EN_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9._@/+\-]*")


def _apply_english_font_markup(text: str, font_name: str = "Arial") -> str:
    parts = re.split(r'(<[^>]+>)', text)
    font_stack: List[bool] = []

    def in_locked_font() -> bool:
        return any(font_stack)

    def wrap_tokens(s: str) -> str:
        return _EN_TOKEN_RE.sub(lambda m: f'<font name="{font_name}">{m.group(0)}</font>', s)

    out: List[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith('<') and part.endswith('>'):
            if re.match(r'<\s*font\b', part, flags=re.IGNORECASE) and not part.rstrip().endswith('/>'):
                locked = bool(re.search(r'\b(name|face)\s*=', part, flags=re.IGNORECASE))
                font_stack.append(locked)
            elif re.match(r'</\s*font\s*>', part, flags=re.IGNORECASE):
                if font_stack:
                    font_stack.pop()
            out.append(part)
            continue

        out.append(part if in_locked_font() else wrap_tokens(part))

    return ''.join(out)



import subprocess as _subprocess
import tempfile as _tempfile
import os as _os
from pathlib import Path as _Path



def _build_academic_html(markdown_content: str, output_path: _Path, workspace_root: str = None) -> str:
    # 灏?Markdown 鍐呭杞崲涓哄鏈鏂囬鏍?HTML锛屼繚鐣?LaTeX 鍏紡渚?MathJax 娓叉煋
    import re as _re
    from pathlib import Path as _Path2
    from urllib.parse import quote as _url_quote

    # Pre-compute the HTML file directory for image path resolution
    _html_dir = str(output_path.parent.resolve())

    # ---- CSS for academic paper style ----
    css = """
    @page {
        size: A4;
        margin: 2.5cm 2.5cm 2.5cm 2.5cm;
        @bottom-center {
            content: counter(page);
            font-family: "SimSun", "瀹嬩綋", serif;
            font-size: 9pt;
        }
    }
    * { box-sizing: border-box; }
    body {
        font-family: "SimSun", "瀹嬩綋", "Noto Serif CJK SC", serif;
        font-size: 11pt;
        line-height: 1.8;
        color: #1a1a1a;
        text-align: justify;
        word-spacing: 0;
        letter-spacing: 0;
        max-width: 100%;
        padding: 0;
        margin: 0;
    }
    h1 {
        font-family: "SimHei", "榛戜綋", "Noto Sans CJK SC", sans-serif;
        font-size: 18pt;
        font-weight: bold;
        text-align: center;
        margin: 24pt 0 16pt 0;
        line-height: 1.4;
        page-break-before: avoid;
        page-break-after: avoid;
    }
    h2 {
        font-family: "SimHei", "榛戜綋", "Noto Sans CJK SC", sans-serif;
        font-size: 14pt;
        font-weight: bold;
        margin: 18pt 0 10pt 0;
        line-height: 1.4;
        page-break-after: avoid;
        border-bottom: 1px solid #333;
        padding-bottom: 4pt;
    }
    h3 {
        font-family: "SimHei", "榛戜綋", "Noto Sans CJK SC", sans-serif;
        font-size: 12pt;
        font-weight: bold;
        margin: 14pt 0 8pt 0;
        line-height: 1.4;
        page-break-after: avoid;
    }
    h4 {
        font-family: "SimHei", "榛戜綋", "Noto Sans CJK SC", sans-serif;
        font-size: 11pt;
        font-weight: bold;
        margin: 12pt 0 6pt 0;
        page-break-after: avoid;
    }
    p {
        margin: 4pt 0 6pt 0;
        text-indent: 2em;
        orphans: 3;
        widows: 3;
    }
    p.no-indent, p.caption, p.table-caption, .abstract p, .keywords p {
        text-indent: 0;
    }
    .abstract-label, .keywords-label {
        font-family: "SimHei", "榛戜綋", "Noto Sans CJK SC", sans-serif;
        font-weight: bold;
    }
    .abstract {
        margin: 8pt 0 12pt 0;
        font-size: 10.5pt;
    }
    .keywords {
        margin: 4pt 0 8pt 0;
        font-size: 10.5pt;
    }
    table {
        border-collapse: collapse;
        width: 100%;
        margin: 10pt 0 6pt 0;
        font-size: 9.5pt;
        page-break-inside: avoid;
    }
    table th {
        background-color: #f0f0f0;
        font-weight: bold;
        border: 1px solid #555;
        padding: 5pt 8pt;
        text-align: center;
    }
    table td {
        border: 1px solid #555;
        padding: 4pt 8pt;
        text-align: center;
    }
    .table-caption {
        font-size: 9pt;
        text-align: center;
        margin: 4pt 0 10pt 0;
        color: #333;
    }
    .figure {
        text-align: center;
        margin: 14pt 0 8pt 0;
        page-break-inside: avoid;
    }
    .figure img {
        max-width: 100%;
        height: auto;
        max-height: 18cm;
    }
    .figure-caption {
        font-size: 9pt;
        margin-top: 4pt;
        color: #333;
        text-align: center;
    }
    pre {
        background: #f5f5f5;
        border: 1px solid #ddd;
        padding: 8pt 12pt;
        font-size: 8.5pt;
        line-height: 1.4;
        overflow-x: auto;
        white-space: pre-wrap;
        word-wrap: break-word;
        margin: 8pt 0;
        font-family: "Consolas", "Courier New", monospace;
    }
    code {
        font-family: "Consolas", "Courier New", monospace;
        background: #f0f0f0;
        padding: 1pt 3pt;
        font-size: 9.5pt;
    }
    pre code {
        background: transparent;
        padding: 0;
    }
    blockquote {
        border-left: 3px solid #ccc;
        margin: 8pt 0;
        padding: 4pt 12pt;
        color: #555;
        background: #fafafa;
    }
    .references {
        font-size: 9.5pt;
        line-height: 1.6;
        text-align: left;
        word-break: break-word;
        overflow-wrap: anywhere;
    }
    .references p {
        text-indent: -2.5em;
        padding-left: 2.5em;
        margin: 2pt 0;
        text-align: left;
    }
    .math-block {
        display: block;
        text-align: center;
        margin: 10pt 0;
        font-size: 9.6pt;
        font-weight: 400;
        line-height: 1.45;
    }
    mjx-container {
        font-weight: 400 !important;
        color: #4a4a4a !important;
    }
    mjx-container[display="true"] {
        margin: 0.45em 0 !important;
        overflow-x: auto;
        overflow-y: hidden;
    }
    mjx-container svg {
        max-width: 100%;
        height: auto;
        transform: scale(0.9);
        transform-origin: center center;
        color: #4a4a4a !important;
        fill: #4a4a4a !important;
        stroke: none !important;
        opacity: 0.88;
    }
    mjx-container svg g,
    mjx-container svg path,
    mjx-container svg use {
        fill: #4a4a4a !important;
        stroke: none !important;
        stroke-width: 0 !important;
    }
    hr {
        border: none;
        border-top: 1px solid #ccc;
        margin: 16pt 0;
    }
    sup { font-size: 0.8em; }
    sub { font-size: 0.8em; }
    .citation {
        font-size: 0.85em;
        vertical-align: super;
    }
    .section-break {
        page-break-before: always;
    }
    """

    # ---- Convert markdown to HTML preserving math ----
    lines = markdown_content.split('\n')
    html_body = []
    in_code_block = False
    in_table = False
    table_rows = []
    in_refs = False
    in_abstract = False
    in_keywords = False
    code_lang = ''

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                html_body.append('<pre><code>')
            else:
                in_code_block = False
                html_body.append('</code></pre>')
            i += 1
            continue

        if in_code_block:
            html_body.append(_escape_html(line))
            i += 1
            continue

        # Display math: $$...$$ (multi-line or single-line)
        stripped = line.strip()
        if stripped == '$$':
            html_body.append('<div class="math-block">$$')
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                html_body.append(lines[i])
                i += 1
            html_body.append('$$</div>')
            i += 1
            continue

        # Single-line display math: $$ ... $$
        if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            math_content = stripped[2:-2].strip()
            html_body.append(f'<div class="math-block">$${math_content}$$</div>')
            i += 1
            continue

        # 整行就是一个 $...$ 公式(LLM 常把独立公式写成单 $),按块级居中显示而非行内
        if (stripped.startswith('$') and stripped.endswith('$')
                and not stripped.startswith('$$') and len(stripped) > 2
                and stripped.count('$') == 2):
            math_content = stripped[1:-1].strip()
            if math_content:
                html_body.append(f'<div class="math-block">$${math_content}$$</div>')
                i += 1
                continue

        # Tables
        if '|' in line and line.strip().startswith('|') and not in_table:
            if i + 1 < len(lines) and _re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
                in_table = True
                table_rows = []
                header_cells = [c.strip() for c in line.strip().split('|')[1:-1]]
                table_rows.append(('header', header_cells))
                i += 2
                continue

        if in_table:
            if line.strip().startswith('|'):
                cells = [c.strip() for c in line.strip().split('|')[1:-1]]
                table_rows.append(('row', cells))
                i += 1
                continue
            else:
                # End table, render it
                html_body.append('<table>')
                for row_type, cells in table_rows:
                    if row_type == 'header':
                        html_body.append('<thead><tr>' + ''.join(f'<th>{_process_inline_markdown(c)}</th>' for c in cells) + '</tr></thead>')
                    else:
                        if row_type == 'row' and table_rows and table_rows[0][0] == 'header' and table_rows.index((row_type, cells)) == 1:
                            html_body.append('<tbody>')
                        html_body.append('<tr>' + ''.join(f'<td>{_process_inline_markdown(c)}</td>' for c in cells) + '</tr>')
                html_body.append('</tbody></table>')

                # Check if next non-empty line is a table caption
                table_caption = None
                ci = i
                while ci < len(lines) and not lines[ci].strip():
                    ci += 1
                if ci < len(lines):
                    cap_match = _re.match(
                        r'^\*?\s*((?:琛▅Table)\s*\d+[\.\uff0e]?.*)\*?\s*$',
                        lines[ci].strip(),
                        flags=_re.IGNORECASE
                    )
                    if cap_match:
                        table_caption = cap_match.group(1).strip()
                        ci += 1
                i = ci

                if table_caption:
                    html_body.append(f'<p class="table-caption">{_escape_html(table_caption)}</p>')

                table_rows = []
                in_table = False
                continue

        # Images
        img_match = _re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            alt = img_match.group(1)
            src = img_match.group(2)
            # Resolve relative image paths to absolute file:// URLs for Chrome headless
            img_abs_src = src
            if src.startswith("./") or src.startswith("../"):
                try:
                    img_abs_src = "file:///" + _Path2(_html_dir).joinpath(src).resolve().as_posix()
                except Exception:
                    pass
            figure_caption = alt
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                cap_match = _re.match(
                    r'^\*?\s*((?:鍥緗鍦東Figure)\s*\d+[\.\uff0e]?.*)\*?\s*$',
                    next_line,
                    flags=_re.IGNORECASE
                )
                if cap_match:
                    figure_caption = cap_match.group(1).strip()
                    i += 1
            html_body.append(
                f'<div class="figure"><img src="{img_abs_src}" alt="{_escape_html(alt)}">'
                f'<p class="figure-caption">{_escape_html(figure_caption)}</p></div>'
            )

            i += 1
            continue

        # Headers
        if line.startswith('# '):
            text = _process_inline_markdown(line[2:].strip())
            html_body.append(f'<h1>{text}</h1>')
            if any(kw in text for kw in ['References', 'Bibliography', 'Citation', '\u53c2\u8003\u6587\u732e']):
                in_refs = True
                in_abstract = False
        elif line.startswith('## '):
            text = _process_inline_markdown(line[3:].strip())
            html_body.append(f'<h2>{text}</h2>')
            if any(kw in text for kw in ['References', 'Bibliography', 'Citation', '\u53c2\u8003\u6587\u732e']):
                in_refs = True
                in_abstract = False
                in_keywords = False
            elif any(kw in text for kw in ['Keywords']):
                in_keywords = True
                in_abstract = False
                html_body.append('<!-- DEBUG: in_keywords set to True -->')
            elif any(kw in text for kw in ['鎽樿', 'Abstract']):
                in_abstract = True
                in_keywords = False
            else:
                in_abstract = False
                in_keywords = False
        elif line.startswith('### '):
            text = _process_inline_markdown(line[4:].strip())
            html_body.append(f'<h3>{text}</h3>')
        elif line.startswith('#### '):
            text = _process_inline_markdown(line[5:].strip())
            html_body.append(f'<h4>{text}</h4>')
        elif line.startswith('##### '):
            text = _process_inline_markdown(line[6:].strip())
            html_body.append(f'<h5>{text}</h5>')
        elif line.startswith('###### '):
            text = _process_inline_markdown(line[7:].strip())
            html_body.append(f'<h6>{text}</h6>')
        elif line.startswith('---') or line.startswith('***'):
            html_body.append('<hr>')
        elif not line.strip():
            html_body.append('')
        else:
            # Regular paragraph
            text = _process_inline_markdown(line.strip())

            # Assign CSS class based on context
            if in_refs:
                cls = 'references'
            elif in_abstract:
                cls = 'abstract'
            elif in_keywords:
                cls = 'keywords'
                in_keywords = False
            else:
                cls = ''
            html_body.append(f'<p class="{cls}">{text}</p>')

        i += 1

    # Build full HTML
    body_str = '\n'.join(html_body)

    # MathJax 来源:优先本地自包含的 tex-svg.js(离线可用、无需额外字体文件).
    # 远程 jsdelivr CDN 的异步脚本在 headless Chrome 的 virtual-time-budget 窗口内
    # 往往来不及加载执行,导致公式显示为原始 LaTeX 源码,因此本地优先.
    _local_mathjax = _Path(__file__).resolve().parent / "assets" / "tex-svg.js"
    if _local_mathjax.exists():
        _mathjax_src = "file:///" + str(_local_mathjax).replace("\\", "/")
    else:
        _mathjax_src = "https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js"

    full_html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>{css}</style>
<script>
window.MathJax = {{
    tex: {{
        inlineMath: [['$', '$']],
        displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
        processEscapes: false,
        packages: {{'[+]': ['ams', 'newcommand', 'configMacros']}},
    }},
    options: {{
        skipHtmlTags: ['script', 'noscript', 'style', 'textarea', 'pre', 'code'],
        ignoreHtmlClass: 'no-mathjax',
        processHtmlClass: 'mathjax-process',
    }},
    svg: {{
        fontCache: 'local',
        scale: 0.9,
        minScale: 0.5,
    }},
    startup: {{
        ready() {{
            MathJax.startup.defaultReady();
            MathJax.startup.promise.then(() => {{
                console.log('MathJax rendering complete');
            }});
        }}
    }}
}};
</script>
<script id="MathJax-script" src="{_mathjax_src}"></script>
</head>
<body>
{body_str}
</body>
</html>'''

    return full_html


def _process_inline_markdown(text: str) -> str:
    """Process inline Markdown while preserving math expressions."""
    import re as _re

    # Step 1: Protect display math $$...$$ first (if any appear inline)
    display_placeholders = []
    def _protect_display(m):
        display_placeholders.append(m.group(1))
        return f'__DMATH_{len(display_placeholders)-1}__'
    text = _re.sub(r'\$\$(.+?)\$\$', _protect_display, text)

    # Step 2: Protect inline math $...$ (but not $$)
    math_placeholders = []
    def _protect_math(m):
        math_placeholders.append(m.group(1))
        return f'__MATH_{len(math_placeholders)-1}__'
    # Match $...$ where content cannot be empty and does not contain $
    text = _re.sub(r'(?<!\$)\$(?!\$)([^\$]+?)\$(?!\$)', _protect_math, text)

    # Step 3: Process formatting (bold, italic, inline code)
    text = _re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = _re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = _re.sub(r'`([^`]+)`', r'<code>\1</code>', text)

    # Step 4: Links
    text = _re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)

    # Step 5: Restore display math
    for idx, formula in enumerate(display_placeholders):
        text = text.replace(f'__DMATH_{idx}__', f'$${formula}$$')

    # Step 6: Restore inline math
    for idx, formula in enumerate(math_placeholders):
        text = text.replace(f'__MATH_{idx}__', f'${formula}$')

    return text


def _escape_html(text: str) -> str:
    # 杞箟 HTML 鐗规畩瀛楃
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def _generate_pdf_with_chrome(markdown_content: str, output_path: _Path, workspace_root: str = None) -> bool:
    # 浣跨敤 Chrome headless + MathJax 灏?Markdown 杞崲涓洪珮璐ㄩ噺瀛湳 PDF
    import subprocess as _sp
    import os as _os_module

    try:
        # Find Chrome
        chrome_paths = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files\Chromium\Application\chrome.exe",
            "/usr/bin/google-chrome",
            "/usr/bin/chromium-browser",
            "/usr/bin/chromium",
        ]
        if sys.platform == 'darwin':
            chrome_paths.insert(0, "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
            chrome_paths.insert(1, "/Applications/Chromium.app/Contents/MacOS/Chromium")

        chrome_exe = None
        for p in chrome_paths:
            if _os_module.path.exists(p):
                chrome_exe = p
                break

        if not chrome_exe:
            logger.error("Chrome/Chromium not found for PDF generation")
            return False

        # Build HTML
        html_content = _build_academic_html(markdown_content, output_path, workspace_root)
        if not html_content:
            return False

        # Image paths are relative to the HTML file location (same dir as PDF). Chrome resolves them correctly.
        # Write HTML file next to PDF
        html_path = output_path.with_suffix('.html')
        with open(str(html_path), 'w', encoding='utf-8') as f:
            f.write(html_content)

        # Prepare Chrome command for headless PDF
        pdf_path_abs = str(output_path.absolute())
        html_path_abs = str(html_path.absolute()).replace('\\', '/')
        html_url = 'file:///' + html_path_abs

        cmd = [
            chrome_exe,
            '--headless=new',
            '--disable-gpu',
            '--no-sandbox',
            '--disable-software-rasterizer',
            '--disable-extensions',
            '--print-to-pdf=' + pdf_path_abs,
            '--no-pdf-header-footer',
            '--print-to-pdf-no-header',
            '--run-all-compositor-stages-before-draw',
            '--virtual-time-budget=20000',
            '--disable-dev-shm-usage',
            '--allow-file-access-from-files',
            html_url
        ]

        logger.info(f"Running Chrome headless for PDF generation: {output_path.name}")

        env = _os_module.environ.copy()
        env['LANGUAGE'] = 'zh_CN:zh'
        env['LANG'] = 'zh_CN.UTF-8'

        result = _sp.run(
            cmd,
            capture_output=True,
            timeout=180,
            env=env
        )

        # Clean up temp HTML
        try:
            html_path.unlink()
        except Exception:
            pass

        if output_path.exists() and output_path.stat().st_size > 1000:
            logger.info(f"Chrome PDF generated successfully: {pdf_path_abs} ({output_path.stat().st_size} bytes)")
            return True
        else:
            stderr_msg = result.stderr.decode('utf-8', errors='ignore')[:500] if result.stderr else '(none)'
            logger.error(f"Chrome PDF generation failed: empty or missing output. stderr: {stderr_msg}")
            if output_path.exists():
                try:
                    output_path.unlink()
                except Exception:
                    pass
            return False

    except _sp.TimeoutExpired:
        logger.error("Chrome PDF generation timed out (180s)")
        return False
    except FileNotFoundError:
        logger.error(f"Chrome executable not found: {chrome_exe}")
        return False
    except Exception as e:
        logger.error(f"Chrome PDF generation failed: {e}")
        import traceback
        traceback.print_exc()
        return False


class MCPTools:
    # Multi Agent System MCP Tools Implementation

    def __init__(self, workspace_path: str = None):
        self.config = get_config()
        self.workspace_path = Path(workspace_path) if workspace_path else Path.cwd()
        self.workspace_path.mkdir(exist_ok=True, parents=True)

        # Session context for workspace-aware operations
        self.session_id = None
        self.session_workspace_path = None
        self.full_workspace_path = os.path.realpath(self.workspace_path)
        if not self.full_workspace_path.endswith(os.sep):
            self.full_workspace_path += os.sep

        # 鍒濆鍖?username锛屽皾璇曚粠 workspace 閰嶇疆鏂囦欢璇诲彇锛屽惁鍒欎娇鐢ㄩ粯璁ゅ€?
        self.username = self._get_username_from_workspace()
        self._pubmed_client = None
        self._pubmed_client_workspace = None

    def _get_username_from_workspace(self) -> str:
        # Read username from the workspace config file if available.
        try:
            username_file = self.workspace_path / '.username'
            if username_file.exists():
                with open(username_file, 'r', encoding='utf-8') as f:
                    username = f.read().strip()
                    if username:
                        return username
        except Exception as e:
            logger.debug(f"Failed to read username from workspace: {e}")

        return "user"

    def set_session_context(self, session_id: str, session_workspace_path: str):
        # Set session context for workspace-aware operations
        self.session_id = session_id
        self.session_workspace_path = Path(session_workspace_path)
        # Update workspace path to session-specific path
        self.workspace_path = self.session_workspace_path
        self.full_workspace_path = os.path.realpath(self.workspace_path)
        if not self.full_workspace_path.endswith(os.sep):
            self.full_workspace_path += os.sep
        # 鏇存柊 username
        self.username = self._get_username_from_workspace()
        # MCPTools can serve different sessions. Do not reuse one session's
        # literature cache path for another session.
        if self._pubmed_client_workspace != str(self.workspace_path):
            if self._pubmed_client is not None:
                self._pubmed_client.close()
            self._pubmed_client = None
            self._pubmed_client_workspace = None
        logger.info(
            f"Set session context - ID: {session_id}, Workspace: {session_workspace_path}, Username: {self.username}")

    def get_session_context(self) -> Dict[str, Any]:
        # Get current session context
        return {
            "session_id": self.session_id,
            "session_workspace_path": str(self.session_workspace_path) if self.session_workspace_path else None,
            "workspace_path": str(self.workspace_path)
        }

    def _get_pubmed_client(self):
        from src.tools.pubmed_client import PubMedClient

        workspace = str(self.workspace_path)
        if self._pubmed_client is None or self._pubmed_client_workspace != workspace:
            cache_dir = self.workspace_path / "research" / "literature_cache" / "pubmed"
            self._pubmed_client = PubMedClient(cache_dir=cache_dir)
            self._pubmed_client_workspace = workspace
        return self._pubmed_client

    def _safe_join(self, path: str) -> Path:
        if os.path.isabs(path):
            raise Exception(f"Path '{path}' is absolute. Only relative paths are allowed.")
        joined_path = os.path.join(self.workspace_path, path)
        full_joined_path = os.path.realpath(joined_path)
        if not full_joined_path.startswith(self.full_workspace_path):
            raise Exception(f"Path '{path}' is outside workspace directory.")
        return Path(full_joined_path)

    # ================ WEB SEARCH TOOLS ================

    def academic_search(
            self,
            queries: List[str],
            sources: Optional[List[str]] = None,
            max_results_per_query: int = 5,
            max_workers: int = 6,
    ) -> MCPToolResult:
        """Search structured scholarly sources while preserving the MCP tool interface."""
        try:
            from src.tools.academic_search import search_academic_sources

            blocks, warnings = search_academic_sources(
                queries,
                sources=sources,
                max_results_per_query=max_results_per_query,
                max_workers=max_workers,
            )
            return MCPToolResult(
                success=any(block.get("success") for block in blocks) or (not blocks and not warnings),
                data={"results": blocks, "warnings": warnings},
                error="; ".join(warnings) if warnings and not any(block.get("success") for block in blocks) else None,
            )
        except Exception as exc:
            return MCPToolResult(success=False, error=f"Academic search failed: {exc}")

    def batch_web_search(
            self,
            queries: List[str],
            max_results_per_query: int = 15,
            max_workers: int = 5
    ) -> MCPToolResult:
        """
        Batch web search using configurable search provider with concurrent processing.
        
        Users need to implement their own search provider. Below is an example available:
        [
            {
                "query": "search query",
                "search_results": [
                    {
                        "title": "Page title",
                        "link": "https://example.com",
                        "snippet": "Description snippet",
                        "date": "Feb 8, 2022",
                    },
                    ...
                ]
            },
            ...
        ]
        
        Args:
            queries: List of search queries
            max_results_per_query: Maximum search results per query
            max_workers: Maximum number of concurrent search requests
        """
        try:
            from config.config import get_search_engine_config
            search_config = get_search_engine_config()

            if not search_config:
                return MCPToolResult(
                    success=False,
                    error="Search engine not configured"
                )

            # Ensure we never return more than 15 results per query
            actual_max_results = min(max_results_per_query, 15)

            def search_single_query(query: str) -> Dict[str, Any]:
                # Search a single query
                try:
                    search_results = self._generic_search(query, actual_max_results, search_config)

                    if not search_results.success:
                        return {
                            'query': query,
                            'success': False,
                            'error': search_results.error,
                            'results': []
                        }

                    # Process search results
                    search_data = search_results.data
                    search_data["organic"] = search_data["organic"][:actual_max_results]

                    return {
                        'query': query,
                        'success': True,
                        'results': search_data,
                        'timestamp': time.time()
                    }

                except Exception as e:
                    logger.error(f"Error searching query '{query}': {e}")
                    return {
                        'query': query,
                        'success': False,
                        'error': str(e),
                        'results': []
                    }

            # Execute searches concurrently
            all_results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(queries))) as executor:
                # Submit all search tasks
                future_to_query = {executor.submit(search_single_query, query): query for query in queries}

                # Collect results as they complete
                for future in as_completed(future_to_query):
                    try:
                        result = future.result()
                        all_results.append(result)
                    except Exception as e:
                        query = future_to_query[future]
                        logger.error(f"Error processing search for '{query}': {e}")
                        all_results.append({
                            'query': query,
                            'success': False,
                            'error': str(e),
                            'results': []
                        })

            # Sort results to maintain original query order
            query_order = {query: i for i, query in enumerate(queries)}
            all_results.sort(key=lambda x: query_order.get(x['query'], float('inf')))

            return MCPToolResult(
                success=True,
                data=all_results,
                metadata={
                    'total_queries': len(queries),
                    'successful_queries': len([r for r in all_results if r.get('success', False)]),
                    'concurrent_workers': min(max_workers, len(queries))
                }
            )

        except Exception as e:
            logger.error(f"Batch web search failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _generic_search(self, query: str, max_results: int, config: Dict[str, Any]) -> MCPToolResult:
        """Run a configured web search and return normalized results."""
        try:
            # This is a placeholder - users should implement their own search logic
            # raise NotImplementedError(
            #     "Generic search provider not implemented. Please implement your own search logic in _generic_search method. "
            #     "The return format should match the standard format with 'organic' results containing title, link, snippet, and date fields."
            # )

            # Example implementation for serper (commented out):
            url = config['base_url']

            payload = json.dumps({
                "q": query,
                "num": 10
            })

            headers = {
                'X-API-KEY': config["api_keys"][0],
                'Content-Type': 'application/json'
            }

            response = requests.request("POST", url, headers=headers, data=payload, timeout=config["timeout"],
                                        proxies=proxy)
            response.raise_for_status()

            return MCPToolResult(success=True, data=response.json())
        except Exception as e:
            return MCPToolResult(success=False, error=f"Generic search failed: {e}")

    def _extract_google_search_date(self, search_item: Dict[str, Any]) -> Optional[str]:
        # Extract publication date from Google Search result
        try:
            # Check pagemap metatags for various date formats
            pagemap = search_item.get('pagemap', {})
            metatags = pagemap.get('metatags', [{}])

            if metatags:
                meta = metatags[0]
                # Common date meta tags
                date_fields = [
                    'article:published_time',
                    'article:modified_time',
                    'date',
                    'pubdate',
                    'published',
                    'datePublished',
                    'dateModified',
                    'dc.date',
                    'dc.date.created',
                    'creation_date'
                ]

                for field in date_fields:
                    if field in meta and meta[field]:
                        return meta[field]

            # Check for news articles with publish date
            newsarticle = pagemap.get('newsarticle', [{}])
            if newsarticle and newsarticle[0].get('datepublished'):
                return newsarticle[0]['datepublished']

            # Check article schema
            article = pagemap.get('article', [{}])
            if article and article[0].get('datepublished'):
                return article[0]['datepublished']

            return None

        except Exception as e:
            logger.info(f"Error extracting date from search result: {e}")
            return None

    @staticmethod
    def _extract_publication_date_from_html(url: str) -> Optional[str]:
        # Extract publication date directly from webpage HTML
        try:
            # Fetch HTML content
            response = requests.get(url, timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # Common meta tags for publication date
            meta_selectors = [
                'meta[property="article:published_time"]',
                'meta[property="article:modified_time"]',
                'meta[name="date"]',
                'meta[name="pubdate"]',
                'meta[name="published"]',
                'meta[name="datePublished"]',
                'meta[name="publication-date"]',
                'meta[property="og:published_time"]',
                'meta[name="DC.date"]',
                'meta[name="DC.date.created"]',
                'meta[itemprop="datePublished"]',
                'meta[itemprop="dateModified"]'
            ]

            for selector in meta_selectors:
                meta_tag = soup.select_one(selector)
                if meta_tag:
                    content = meta_tag.get('content') or meta_tag.get('datetime')
                    if content:
                        try:
                            # Parse and standardize the date
                            parsed_date = dateutil.parser.parse(content)
                            return parsed_date.isoformat()
                        except ValueError:
                            continue

            # Check for time tags with datetime attribute
            time_tags = soup.find_all('time', {'datetime': True})
            for time_tag in time_tags:
                try:
                    parsed_date = dateutil.parser.parse(time_tag['datetime'])
                    return parsed_date.isoformat()
                except ValueError:
                    continue

            # JSON-LD structured data
            json_ld_scripts = soup.find_all('script', {'type': 'application/ld+json'})
            for script in json_ld_scripts:
                try:
                    data = json.loads(script.string)
                    if isinstance(data, dict):
                        date_fields = ['datePublished', 'dateCreated', 'dateModified']
                        for field in date_fields:
                            if field in data:
                                parsed_date = dateutil.parser.parse(data[field])
                                return parsed_date.isoformat()
                    elif isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict):
                                for field in date_fields:
                                    if field in item:
                                        parsed_date = dateutil.parser.parse(item[field])
                                        return parsed_date.isoformat()
                except ValueError:
                    continue

            return None

        except Exception as e:
            logger.debug(f"Error extracting publication date from {url}: {e}")
            return None

    def _content_extractor(self, url: str, max_tokens: int, config: Dict[str, Any]) -> MCPToolResult:
        # Get content using Direct BeautifulSoup (Bypass placeholder API)
        try:
            import requests
            from bs4 import BeautifulSoup

            # 浼鎴愮湡瀹炵殑楂樼骇娴忚鍣紝缁曡繃闃茬伀澧?
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7',
                'Accept-Encoding': 'gzip, deflate, br',
                'Referer': 'https://www.google.com/',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }

            # 绂佺敤璇佷功楠岃瘉锛岄槻姝?SSL 鎶ラ敊
            import urllib3
            urllib3.disable_warnings()

            response = requests.get(url, headers=headers, timeout=30, verify=False)
            response.raise_for_status()

            content_type = response.headers.get('Content-Type', '')

            # 濡傛灉鐩存帴杩斿洖鐨勫氨鏄?PDF
            if 'application/pdf' in content_type:
                return MCPToolResult(success=True, data=response.content, metadata={'is_pdf': True})

            # 鏄?HTML 椤甸潰锛屽皾璇曟壘 PDF 閾炬帴
            from bs4 import BeautifulSoup
            from urllib.parse import urljoin
            soup = BeautifulSoup(response.content, 'html.parser')

            pdf_url = None
            for a in soup.find_all('a', href=True):
                href = a['href']
                if href.lower().endswith('.pdf') or '/pdf/' in href.lower():
                    pdf_url = urljoin(url, href) if not href.startswith('http') else href
                    break

            if pdf_url:
                try:
                    pdf_resp = requests.get(pdf_url, headers=headers, timeout=30, verify=False)
                    pdf_resp.raise_for_status()
                    if 'application/pdf' in pdf_resp.headers.get('Content-Type', ''):
                        return MCPToolResult(success=True, data=pdf_resp.content, metadata={'is_pdf': True})
                except Exception as e:
                    logger.warning(f"PDF涓嬭浇澶辫触({e})锛岄檷绾т繚瀛楬TML")

            # 娌℃壘鍒?PDF 鎴栦笅杞藉け璐ワ紝淇濆瓨鍘熷 HTML
            return MCPToolResult(success=True, data=response.text, metadata={'is_pdf': False})

        except Exception as e:
            logger.warning(f"Webpage crawling failed: {url}; reason: {str(e)}")
            return MCPToolResult(success=False, error=f"Direct crawling failed: {str(e)}")

    def url_crawler(
            self,
            documents: List[Dict],
            max_tokens_per_url: int = 100000,
            include_metadata: bool = True,
            max_workers: int = 10
    ) -> MCPToolResult:
        """
        Extract LLM-friendly content from URLs using configurable crawler service.
        Content is saved to specified file paths.

        Users need to implement their own URL crawler. The return format should include:
        - Extracted text content from the URL
        - Metadata like title, publication date, word count
        - Success/error status for each URL

        Args:
            documents: List of document dictionaries containing:
                - url: Web page URL to extract
                - file_path: Local path to save extracted content
                - title: (Optional) Web page title
                - time: (Optional) Web page publication time
            max_tokens_per_url: Maximum tokens per URL result
            include_metadata: Whether to include metadata about extraction
            max_workers: Maximum number of concurrent extraction requests
        """
        try:
            from config.config import get_url_crawler_config
            crawler_config = get_url_crawler_config()

            if not crawler_config:
                return MCPToolResult(
                    success=False,
                    error="URL crawler not configured"
                )

            def process_single_document(doc: Dict) -> Dict[str, Any]:
                # Process a single document: extract content and save to file
                url = doc['url']
                file_path = doc['file_path']
                title = doc.get('title')
                doc_time = doc.get('time')

                # Skip domains that commonly block automated crawling; keep URL-only references.
                FORBIDDEN_DOMAINS = [
                    'mdpi.com',
                    'ieeexplore.ieee.org',
                    'sciencedirect.com',
                    'springer.com',
                    'wiley.com'
                ]

                if any(domain in url.lower() for domain in FORBIDDEN_DOMAINS):
                    logger.warning(f"Skipping restricted domain; saving URL reference only: {url}")
                    try:
                        url_save_path = file_path.rsplit('.', 1)[0] + '_url_ref.txt'
                        url_info = f"Title: {title}\nURL: {url}\nDate: {doc_time}\nStatus: URL-only reference (domain blocked from crawling)\n"
                        self.file_write(file_path=url_save_path, content=url_info, create_dirs=True)
                        logger.info(f"Saved URL reference to: {url_save_path}")
                    except Exception as save_err:
                        logger.warning(f"Failed to save URL reference: {save_err}")
                    return {
                        'url': url,
                        'file_path': file_path,
                        'title': title,
                        'success': True,
                        'saved_as_url_ref': True,
                        'error': "Website crawling blocked; saved URL reference information for citation use.",
                        'write_success': True
                    }
                # ========================================================================

                result_base = {
                    'url': url,
                    'file_path': file_path,
                    'title': title,
                    'time': doc_time,
                    'success': False,
                    'error': None,
                    'content_length': 0,
                    'word_count': 0,
                    'publication_date': None,
                    'extraction_timestamp': time.time(),
                    'write_success': False
                }

                try:
                    # Extract publication date from the webpage
                    publication_date = self._extract_publication_date_from_html(url)
                    result_base["publication_date"] = publication_date

                    # Extract content using content extractor
                    content_result = self._content_extractor(url, max_tokens_per_url, crawler_config)

                    if not content_result.success:
                        try:
                            url_save_path = file_path.rsplit('.', 1)[0] + '_url_ref.txt'
                            url_info = f"Title: {title}\nURL: {url}\nDate: {doc_time}\nStatus: Extraction failed - {content_result.error}\n"
                            self.file_write(file_path=url_save_path, content=url_info, create_dirs=True)
                            logger.info(f"Extraction failed; saved URL reference to: {url_save_path}")
                        except Exception as save_err:
                            logger.warning(f"Failed to save URL reference: {save_err}")
                        result_base['error'] = content_result.error
                        result_base['saved_as_url_ref'] = True
                        return result_base

                    is_pdf = content_result.metadata and content_result.metadata.get('is_pdf')
                    raw_data = content_result.data

                    if not raw_data:
                        result_base['error'] = "Extracted content is empty"
                        return result_base

                    # 鏋勫缓鍏冩暟鎹ご锛堜粎鐢ㄤ簬 HTML锛?
                    metadata_header = ""
                    if title:
                        metadata_header += f"Title: {title}\n"
                    metadata_header += f"URL: {url}\n"
                    if doc_time:
                        metadata_header += f"Date: {doc_time}\n"
                    if publication_date:
                        metadata_header += f"Publication Date: {publication_date}\n"
                    metadata_header += "\n"

                    if is_pdf:
                        # PDF 浜岃繘鍒跺啓鍏ワ紝鍚庣紑鏀逛负 .pdf
                        file_path = file_path.rsplit('.', 1)[0] + '.pdf'
                        full_path = self.workspace_path / file_path
                        full_path.parent.mkdir(parents=True, exist_ok=True)
                        with open(full_path, 'wb') as f:
                            f.write(raw_data)
                        write_success = True
                    else:
                        # HTML 淇濆瓨锛屽悗缂€鏀逛负 .html
                        file_path = file_path.rsplit('.', 1)[0] + '.html'
                        write_result = self.file_write(
                            file_path=file_path,
                            content=metadata_header + raw_data,
                            create_dirs=True
                        )
                        write_success = write_result.success

                    if not write_success:
                        result_base['error'] = "File write failed"
                        return result_base

                    # Extract and persist document facts immediately after saving crawled content.
                    try:
                        self.document_extract(
                            tasks=[{"file_path": file_path,
                                    "task": "Extract core information and references for paper writing"}],
                            max_workers=1
                        )
                        logger.info(f"Background extraction wrote JSONL successfully: {file_path}")
                    except Exception as ext_err:
                        logger.warning(f"Background extraction warning: {ext_err}")

                    # Build success result
                    content_len = len(raw_data) if is_pdf else len(raw_data.split())
                    result = {
                        **result_base,
                        'file_path': file_path,
                        'success': True,
                        'content_length': content_len,
                        'word_count': 0 if is_pdf else len(raw_data.split()),
                        'publication_date': publication_date,
                        'write_success': True
                    }

                    if include_metadata:
                        result['metadata'] = {
                            'is_pdf': is_pdf,
                            'has_publication_date': publication_date is not None,
                            'date_extraction_method': 'html_parsing' if publication_date else None,
                            'file_size': len(raw_data) if is_pdf else len(raw_data.encode('utf-8'))
                        }

                    return result

                except Exception as e:
                    logger.error(f"Error processing document {url}: {e}")
                    # Try to extract publication date even if processing failed
                    try:
                        publication_date = self._extract_publication_date_from_html(url)
                    except:
                        publication_date = None

                    return {
                        **result_base,
                        'error': str(e),
                        'publication_date': publication_date
                    }

            # Execute processing concurrently
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(documents))) as executor:
                # Submit all processing tasks
                future_to_doc = {executor.submit(process_single_document, doc): doc for doc in documents}

                # Collect results as they complete
                for future in as_completed(future_to_doc):
                    doc = future_to_doc[future]
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        url = doc['url']
                        logger.error(f"Error processing extraction for '{url}': {e}")

                        # Try to extract publication date even if processing failed
                        try:
                            publication_date = self._extract_publication_date_from_html(url)
                        except:
                            publication_date = None

                        results.append({
                            'url': url,
                            'file_path': doc['file_path'],
                            'title': doc.get('title'),
                            'time': doc.get('time'),
                            'success': False,
                            'error': str(e),
                            'publication_date': publication_date,
                            'extraction_timestamp': time.time(),
                            'write_success': False
                        })

            # Sort results to maintain original document order
            url_order = {doc['url']: i for i, doc in enumerate(documents)}
            results.sort(key=lambda x: url_order.get(x['url'], float('inf')))

            successful_extractions = len([r for r in results if r.get('success', False)])
            successful_writes = len([r for r in results if r.get('write_success', False)])

            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_documents': len(documents),
                    'successful_extractions': successful_extractions,
                    'successful_writes': successful_writes,
                    'failed_processing': len(documents) - successful_extractions,
                    'concurrent_workers': min(max_workers, len(documents))
                }
            )

        except Exception as e:
            logger.error(f"URL crawler batch processing failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _extract_original_filename(self, filename: str) -> str:
        """
        浠庢枃浠跺悕涓彁鍙栧師濮嬫枃浠跺悕锛堝幓鎺塮ile_id鍓嶇紑锛?

        Args:
            filename: 鍙兘鍖呭惈file_id鍓嶇紑鐨勬枃浠跺悕锛屾牸寮忓 'file_id_filename.ext' 鎴?'filename.ext'

        Returns:
            鍘熷鏂囦欢鍚嶏紙鍘绘帀file_id鍓嶇紑锛?
        """
        if '_' in filename:
            parts = filename.split('_', 1)
            # 濡傛灉绗竴閮ㄥ垎鏄痜ile_id锛?浣嶄互涓婂崄鍏繘鍒讹級锛屽垯浣跨敤绗簩閮ㄥ垎
            if len(parts) > 1 and len(parts[0]) >= 8 and re.match(r'^[a-f0-9]{8,}', parts[0].lower()):
                return parts[1]
        return filename

    def _extract_title_from_filename(self, filename: str) -> str:
        """
        浠庢枃浠跺悕涓彁鍙栨爣棰橈紙鍘绘帀file_id鍓嶇紑锛屼繚鐣欐枃浠舵墿灞曞悕濡?pdf/.txt/.doc锛?

        Args:
            filename: 鍙兘鍖呭惈file_id鍓嶇紑鐨勬枃浠跺悕

        Returns:
            鏍囬锛堜繚鐣欏師濮嬫枃浠舵墿灞曞悕锛屽彧鍘绘帀缂撳瓨鏂囦欢鐨?txt鍚庣紑锛?
        """
        # 鍏堟彁鍙栧師濮嬫枃浠跺悕锛堝幓鎺塮ile_id鍓嶇紑锛?
        original_filename = self._extract_original_filename(filename)
        # 鍙幓鎺夌紦瀛樻枃浠剁殑.txt鍚庣紑锛堝 .doc.txt, .docx.txt, .pdf.txt锛?
        # 浣嗕繚鐣欏師鐢?txt鏂囦欢鐨勬墿灞曞悕
        if (original_filename.endswith('.doc.txt') or
                original_filename.endswith('.docx.txt') or
                original_filename.endswith('.pdf.txt')):
            original_filename = original_filename[:-4]
        return original_filename

    def _extract_title_from_file_content(self, file_path: Path) -> tuple:
        """
        浠庢枃浠跺唴瀹逛腑鎻愬彇鏍囬鍜孶RL

        Args:
            file_path: 鏂囦欢璺緞

        Returns:
            (title, url_source) 鍏冪粍
        """
        title = "Unknown Title"
        url_source = "Unknown URL"

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                lines = content.split('\n')

                # 浼樺厛绾?: 鏌ユ壘鍏冩暟鎹锛堝鏋滅埇铏繚瀛樹簡鏍囬淇℃伅锛?
                # 鏍煎紡: "Title: xxx" 鎴?"# 鏍囬"
                metadata_title = None
                for i, line in enumerate(lines[:50]):
                    line_stripped = line.strip()
                    # 鍖归厤 "Title: xxx" 鏍煎紡
                    title_match = re.match(r'^Title:\s*(.+)$', line_stripped, re.IGNORECASE)
                    if title_match:
                        metadata_title = title_match.group(1).strip()
                        logger.info(f"Extracted title from metadata at line {i + 1}: {metadata_title[:80]}")
                        break
                    
                    # 鍖归厤 Markdown 涓€绾ф爣棰?
                    md_title_match = re.match(r'^#\s+(.+)$', line_stripped)
                    if md_title_match:
                        metadata_title = md_title_match.group(1).strip()
                        # 鍘绘帀鍙兘鐨凪arkdown鏍煎紡鏍囪
                        metadata_title = re.sub(r'\*\*|__', '', metadata_title).strip()
                        logger.info(f"Extracted title from Markdown at line {i + 1}: {metadata_title[:80]}")
                        break
                
                if metadata_title and len(metadata_title) > 5:
                    title = metadata_title

                # 浼樺厛绾?: 濡傛灉娌℃湁鍏冩暟鎹爣棰橈紝灏濊瘯浠庡唴瀹逛腑鎻愬彇锛堝鏈鏂囬€氬父绗竴娈靛寘鍚爣棰橈級
                if title == "Unknown Title":
                    # 鏌ユ壘闈炵┖琛岋紝鎺掗櫎URL鍜岀壒娈婃爣璁拌
                    for i, line in enumerate(lines[:30]):
                        line_stripped = line.strip()
                        if (line_stripped and 
                            10 <= len(line_stripped) <= 300 and
                            'http' not in line_stripped and 
                            not line_stripped.startswith('[') and
                            not line_stripped.startswith('#') and
                            not line_stripped.startswith('Title:') and
                            # 鎺掗櫎鏄庢樉鐨勯潪鏍囬鍐呭
                            not line_stripped.lower().startswith('abstract') and
                            not line_stripped.lower().startswith('keywords')):
                            if line.startswith('<!DOCTYPE') or line.startswith('<html'):
                                continue
                            title = line_stripped[:300]
                            logger.info(f"Extracted title from content at line {i + 1}: {title[:80]}")
                            break

                # 鏀硅繘URL鎻愬彇閫昏緫锛氭帓闄や腑鏂囨爣鐐圭鍙凤紝纭繚URL涓嶅寘鍚棩鏈?
                for line in lines[:50]:
                    # 鍖归厤URL锛屼絾鎺掗櫎涓枃鏍囩偣绗彿锛堬紝銆傦紱锛氾紒锛燂級鍜屽彸鏂规嫭鍙穄
                    url_match = re.search(r'https?://[^\s\]锛屻€傦紱锛氾紒锛焆+', line)
                    if url_match:
                        url_source = url_match.group(0)
                        logger.info(f"鎻愬彇鍒癠RL: {url_source[:80]}")
                        break
                        
                # 濡傛灉娌℃壘鍒癠RL锛屽皾璇曚粠鏂囦欢鍚嶆帹鏂?
                if url_source == "Unknown URL":
                    filename = os.path.basename(file_path)
                    # 濡傛灉鏂囦欢鍚嶅寘鍚玌RL缂栫爜鎴栧師濮婾RL
                    if 'http' in filename.lower():
                        url_match = re.search(r'(https?://[^\s\[\]]+)', filename)
                        if url_match:
                            url_source = url_match.group(1)
                            
        except Exception as e:
            logger.warning(f"璀憡: 鏃犳硶璇诲彇鐮旂┒鏂囦欢 {file_path} - {str(e)}")

        return title, url_source

    def _extract_title_from_research_filename(self, file_path: str) -> str:
        """
        浠巖esearch鏂囦欢鍚嶄腑鎻愬彇鏍囬锛堝鐢ㄦ柟妗堬級

        Args:
            file_path: 鏂囦欢璺緞

        Returns:
            鏍囬锛堝鏋滄枃浠跺悕鏈夋晥锛?
        """
        filename = os.path.basename(file_path)
        if filename and filename != file_path:
            title_candidate = os.path.splitext(filename)[0]
            # 濡傛灉鏂囦欢鍚嶇湅璧锋潵鏈夋剰涔夛紙涓嶆槸闅忔満瀛楃涓诧級锛屼娇鐢ㄥ畠
            if len(title_candidate) > 3 and not re.match(r'^[a-f0-9]{32}', title_candidate):
                logger.info(f"浠庢枃浠跺悕鎻愬彇鏍囬: {title_candidate}")
                return title_candidate
        return "Unknown Title"

    # new
    def generate_abstract_and_keywords(self, article_content: str, user_query: str = "") -> Dict[str, str]:
        """Generate title, abstract, and keywords for an article."""
        try:
            from src.utils.llm_client import chat_completion_response

            config = get_config()
            model_config = config.get_custom_llm_config()
            model_name = model_config.get('model') or config.model_name
            # 璇█妫€娴嬶細鍐冲畾鐢熸垚璇█
            sample_text = article_content[:5000]
            zh_count = len(re.findall(r'[\u4e00-\u9fff]', sample_text))
            en_count = len(re.findall(r'[a-zA-Z]', sample_text))

            is_english_content = False
            if zh_count < 50 and en_count > 200:
                is_english_content = True
            elif en_count > 0 and (zh_count / en_count) < 0.05:
                is_english_content = True

            # 鏍规嵁鍐呭璇█瀹氬埗 Prompt
            if is_english_content:
                lang_instruction = "The article is in English. You MUST generate the Title, Abstract, and Keywords in English."
                format_instruction = """
                Please strictly follow this format:

                Title:
                [Write Title Here]

                Abstract:
                [Write Abstract Here]

                Keywords:
                [Keyword1; Keyword2; Keyword3]
                """
            else:
                lang_instruction = "The article contains Chinese; generate the requested metadata in Chinese."
                format_instruction = """
                璇蜂弗鏍兼寜鐓т互涓嬫牸寮忚緭鍑猴細

                鏍囬锛?
                [鍦ㄨ繖閲屽啓鏍囬]

                鎽樿锛?
                [鍦ㄨ繖閲屽啓鎽樿鍐呭]

                鍏抽敭璇嶏細
                [鍏抽敭璇?; 鍏抽敭璇?; 鍏抽敭璇?]
                """

            prompt = f"""Generate a concise title, abstract, and 5-8 keywords for the following article.

Article content:
{article_content}

Requirements:
1. Language requirement: {lang_instruction}
2. The title must be concise and reflect the article topic.
3. The abstract should summarize the main content, methods, findings, and conclusions.
4. Keywords should reflect the topic and core methods.
5. Follow this output format exactly:

{format_instruction}
"""

            logger.info("Generating title, abstract, and keywords with configured model...")
            response = chat_completion_response(
                {
                    "model": model_name,
                    "messages": [
                        {"role": "system", "content": "You are a professional academic editor. Generate the title, abstract, and keywords in the requested format."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3
                },
                model_config=model_config,
                agent_name="generate_abstract_and_keywords",
                request_logger=logger,
            )

            # 瑙ｆ瀽鍝嶅簲
            result_data = response.json()
            result_text = result_data['choices'][0]['message']['content']

            # 鎵撳嵃鍘熷杩斿洖缁撴灉鐢ㄤ簬璋冭瘯
            logger.info("\n" + "=" * 60)
            logger.info(f"{model_name} raw model response:")
            logger.info(result_text)
            logger.info("=" * 60 + "\n")

            # 瑙ｆ瀽缁撴灉 - 浣跨敤澶氱绛栫暐
            title = ""
            abstract = ""
            keywords = ""

            # 绛栫暐1锛氬皾璇曟爣鍑嗘牸寮忔彁鍙栵紙鏍囬/Title锛?..鎽樿/Abstract锛?..鍏抽敭璇?Keywords锛?..锛?
            title_match = re.search(r'(?:Title)[:?]\s*(.*?)(?=(?:Abstract)|(?:Keywords)|$)', result_text,
                                    re.DOTALL | re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()
                logger.info(f"strategy 1 extracted title ({len(title)} chars)")

            abstract_match = re.search(r'(?:Abstract)[:?]\s*(.*?)(?=(?:Keywords)|$)', result_text,
                                       re.DOTALL | re.IGNORECASE)
            if abstract_match:
                abstract = abstract_match.group(1).strip()
                logger.info(f"strategy 1 extracted abstract ({len(abstract)} chars)")

            keywords_match = re.search(r'(?:Keywords)[:?]\s*(.*?)$', result_text, re.DOTALL | re.IGNORECASE)
            if keywords_match:
                keywords = keywords_match.group(1).strip()
                logger.info("strategy 1 extracted keywords")

            # 绛栫暐2锛氬鏋滅瓥鐣?澶辫触锛屽皾璇曟洿瀹芥澗鐨勫尮閰?
            if not title:
                # 鏌ユ壘 "鏍囬" 鍚庨潰鐨勫唴瀹?
                title_match2 = re.search(r'(?:title)[:?\s]*(.*?)(?=abstract|keywords|$)', result_text,
                                         re.IGNORECASE | re.DOTALL)
                if title_match2:
                    title = title_match2.group(1).strip()
                    logger.info(f"strategy 2 extracted title ({len(title)} chars)")

            if not abstract:
                # 鏌ユ壘 "鎽樿" 鍚庨潰鐨勫唴瀹癸紝鐩村埌閬囧埌 "鍏抽敭璇? 鎴栨枃鏈粨鏉?
                abstract_match2 = re.search(r'(?:abstract)[:?\s]*(.*?)(?=keywords|$)', result_text,
                                            re.IGNORECASE | re.DOTALL)
                if abstract_match2:
                    abstract = abstract_match2.group(1).strip()
                    logger.info(f"strategy 2 extracted abstract ({len(abstract)} chars)")

            if not keywords:
                # 鏌ユ壘 "鍏抽敭璇? 鍚庨潰鐨勫唴瀹?
                keywords_match2 = re.search(r'(?:keywords)[:?\s]*(.*?)$', result_text,
                                            re.IGNORECASE | re.DOTALL)
                if keywords_match2:
                    keywords = keywords_match2.group(1).strip()
                    logger.info("strategy 2 extracted keywords")

            # 绛栫暐3锛氬鏋滀粛鐒跺け璐ワ紝灏濊瘯鎸夎鍒嗗壊
            if not title or not abstract or not keywords:
                lines = result_text.split('\n')
                in_title = False
                in_abstract = False
                in_keywords = False
                title_lines = []
                abstract_lines = []
                keywords_lines = []

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # 妫€鏌ユ槸鍚槸鏍囬鏍囬
                    if re.match(r'(?:Title)[:?]?', line, re.IGNORECASE):
                        in_title = True
                        in_abstract = False
                        in_keywords = False
                        # 濡傛灉鏍囬鍚庨潰鐩存帴鏈夊唴瀹癸紝鎻愬彇瀹?
                        content = re.sub(r'^(?:Title)[:?]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            title_lines.append(content)
                        continue

                    # 妫€鏌ユ槸鍚槸鎽樿鏍囬
                    if re.match(r'(?:Abstract)[:?]?', line, re.IGNORECASE):
                        in_title = False
                        in_abstract = True
                        in_keywords = False
                        # 濡傛灉鏍囬鍚庨潰鐩存帴鏈夊唴瀹癸紝鎻愬彇瀹?
                        content = re.sub(r'^(?:Abstract)[:?]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            abstract_lines.append(content)
                        continue

                    # 妫€鏌ユ槸鍚槸鍏抽敭璇嶆爣棰?
                    if re.match(r'(?:Keywords)[:?]?', line, re.IGNORECASE):
                        in_title = False
                        in_keywords = True
                        in_abstract = False
                        # 濡傛灉鏍囬鍚庨潰鐩存帴鏈夊唴瀹癸紝鎻愬彇瀹?
                        content = re.sub(r'^(?:Keywords)[:?]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            keywords_lines.append(content)
                        continue

                    # 鏀堕泦鍐呭
                    if in_title:
                        title_lines.append(line)
                    elif in_abstract:
                        abstract_lines.append(line)
                    elif in_keywords:
                        keywords_lines.append(line)

                if not title and title_lines:
                    title = ' '.join(title_lines)
                    logger.info(f"strategy 3 extracted title ({len(title)} chars)")

                if not abstract and abstract_lines:
                    abstract = ' '.join(abstract_lines)
                    logger.info(f"strategy 3 extracted abstract ({len(abstract)} chars)")

                if not keywords and keywords_lines:
                    keywords = ' '.join(keywords_lines)
                    logger.info("strategy 3 extracted keywords")

            # 娓呯悊鎻愬彇鐨勫唴瀹?
            if title:
                # 绉婚櫎澶氫綑鐨勭┖鐧藉拰鎹㈣
                title = re.sub(r'\s+', ' ', title).strip()
                # 绉婚櫎寮€澶寸殑鏍囩偣绗彿
                title = re.sub(r'^[锛?\-\s]+', '', title)

            if abstract:
                # 绉婚櫎澶氫綑鐨勭┖鐧藉拰鎹㈣
                abstract = re.sub(r'\s+', ' ', abstract).strip()
                # 绉婚櫎寮€澶寸殑鏍囩偣绗彿
                abstract = re.sub(r'^[锛?\-\s]+', '', abstract)

            if keywords:
                # 绉婚櫎澶氫綑鐨勭┖鐧藉拰鎹㈣
                keywords = re.sub(r'\s+', ' ', keywords).strip()
                # 绉婚櫎寮€澶寸殑鏍囩偣绗彿
                keywords = re.sub(r'^[锛?\-\s]+', '', keywords)

            # 楠岃瘉缁撴灉
            if title and abstract and keywords:
                logger.info(f"\nSuccessfully generated title, abstract, and keywords ({len(abstract)} chars in abstract)")
                logger.info(f"title: {title}")
                logger.info(f"abstract preview: {abstract[:100]}...")
                logger.info(f"keywords: {keywords}")
            else:
                logger.info("metadata extraction incomplete")
                logger.info(f"  title present: {bool(title)} ({len(title) if title else 0} chars)")
                logger.info(f"  abstract present: {bool(abstract)} ({len(abstract) if abstract else 0} chars)")
                logger.info(f"  keywords present: {bool(keywords)}")

                # 濡傛灉鎻愬彇澶辫触锛屼娇鐢ㄥ鐢ㄦ柟妗?
                if not title:
                    # 灏濊瘯浠庢枃绔犲唴瀹逛腑鎻愬彇绗竴涓竴绾ф爣棰?
                    title_from_content = re.search(r'^#\s+(.+)$', article_content, re.MULTILINE)
                    if title_from_content:
                        title = title_from_content.group(1).strip()
                        logger.info("  fallback: extracted title from content")
                    else:
                        title = "Research Report"
                        logger.info("  fallback: using default title")

                if not abstract and len(result_text) > 50:
                    abstract = result_text[:300].strip()
                    logger.info("  fallback: using first 300 chars as abstract")

                if not keywords:
                    keywords = "keywords unavailable"
                    logger.info("  fallback: using default keywords")

            return {
                "title": title if title else "Research Report",
                "abstract": abstract if abstract else "abstract unavailable",
                "keywords": keywords if keywords else "keywords unavailable"
            }

        except Exception as e:
            import traceback
            logger.error(f"Failed to generate title, abstract, and keywords: {str(e)}")
            logger.error(f"traceback:\n{traceback.format_exc()}")
            return {
                "title": "Academic Paper",
                "abstract": "abstract generation failed",
                "keywords": "keyword generation failed"
            }

    def file_read_dq(self, file_path: str, encoding: str = 'utf-8') -> MCPToolResult:
        # Read file content
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            content = full_path.read_text(encoding=encoding)
            if len(content) > 40000:
                content = (
                    "Due to the content being too long, only the first 30,000 and last 10,000 characters are returned.\n"
                    "Below is the returned portion of the file content:\n\n"
                    f"First 30,000 characters:\n\n{content[:30000]}\n\n"
                    f"Last 10,000 characters:\n\n{content[-10000:]}"
                )

            return MCPToolResult(
                success=True,
                data=content,
                metadata={
                    'file_size': len(content),
                    'line_count': len(content.splitlines()),
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def load_json(self, file_path: str, encoding: str = 'utf-8') -> MCPToolResult:
        """
        Read JSON format file 
        """
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            res = []

            with open(full_path, "r", encoding=encoding, errors='ignore') as f:
                for idx, line in enumerate(f):
                    try:
                        ele = json.loads(line.strip())
                        res.append(ele)
                    except Exception as e:
                        logger.warning(f"Failed to process file: {e}")
                        continue

            return MCPToolResult(
                success=True,
                data=res,
                metadata={
                    'line_count': len(res),
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def extract_author_and_title_for_reference(self, article_content: str, source_info: str = "") -> Dict[str, str]:
        """Extract author, title, and source metadata for references."""
        try:
            from src.utils.llm_client import chat_completion_response

            config = get_config()
            model_config = get_config().get_custom_llm_config()
            model_name = model_config.get('model') or config.model_name
            # 鎻愬彇鏂囩珷寮€澶寸殑500涓瓧绗?
            content_excerpt = article_content[:500] if len(article_content) > 500 else article_content

            source_value = source_info if source_info else "Not provided"
            prompt = f"""Extract only the author, title, and source from the following article excerpt.

Rules:
1. Do not invent any missing information.
2. If the author is missing, output [Unknown author].
3. If the title is missing, output [Unknown title].
4. Use the provided source exactly as given.

Article excerpt:
{content_excerpt}

Output format:
Author:
[author names only]

Title:
[article title only]

Source:
{source_value}
"""

            logger.info("姝ｅ湪璋冪敤 PANGU 妯″瀷鎻愬彇浣滆€呭拰鏍囬淇℃伅...")
            response = chat_completion_response(
                {
                    "model": model_name,
                    "messages": [
                        {"role": "system",
                         "content": "You are a professional reference manager. Extract author and title information and format references as requested."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3
                },
                model_config=model_config,
                agent_name="extract_author_and_title_for_reference",
                request_logger=logger,
            )

            # 瑙ｆ瀽鍝嶅簲
            result_data = response.json()
            result_text = result_data['choices'][0]['message']['content']

            # 鎵撳嵃鍘熷杩斿洖缁撴灉鐢ㄤ簬璋冭瘯
            logger.info("\n" + "=" * 60)
            logger.info(f"{model_name} raw model response for reference metadata:")
            logger.info(result_text)
            logger.info("=" * 60 + "\n")

            # 瑙ｆ瀽缁撴灉
            author = ""
            title = ""
            source = ""

            # 绛栫暐1锛氬皾璇曟爣鍑嗘牸寮忔彁鍙?
            author_match = re.search(r'Author[:?]\s*(.*?)(?=Title|Source|$)', result_text, re.DOTALL | re.IGNORECASE)
            if author_match:
                author = author_match.group(1).strip()
                logger.info(f"extracted author: {author}")

            title_match = re.search(r'Title[:?]\s*(.*?)(?=Source|$)', result_text, re.DOTALL | re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()
                logger.info(f"extracted title: {title}")

            source_match = re.search(r'Source[:?]\s*(.*?)$', result_text, re.DOTALL | re.IGNORECASE)
            if source_match:
                source = source_match.group(1).strip()
                logger.info(f"extracted source: {source}")

            # 绛栫暐2锛氬鏋滅瓥鐣?澶辫触锛屽皾璇曟洿瀹芥澗鐨勫尮閰?
            if not author:
                author_match2 = re.search(r'(?:author)[:?\s]*(.*?)(?=title|source|$)', result_text,
                                          re.IGNORECASE | re.DOTALL)
                if author_match2:
                    author = author_match2.group(1).strip()
                    logger.info(f"strategy 2 extracted author: {author}")

            if not title:
                title_match2 = re.search(r'(?:title)[:?\s]*(.*?)(?=source|$)', result_text,
                                         re.IGNORECASE | re.DOTALL)
                if title_match2:
                    title = title_match2.group(1).strip()
                    logger.info(f"strategy 2 extracted title: {title}")

            if not source:
                source_match2 = re.search(r'(?:source)[:?\s]*(.*?)$', result_text, re.IGNORECASE | re.DOTALL)
                if source_match2:
                    source = source_match2.group(1).strip()
                    logger.info(f"strategy 2 extracted source: {source}")

            # 娓呯悊鎻愬彇鐨勫唴瀹?
            if author:
                author = re.sub(r'\s+', ' ', author).strip()
                author = re.sub(r'^[锛?\-\s]+', '', author)
                # 绉婚櫎澶氫綑鐨勬崲琛?
                author = author.replace('\n', ' ')

            if title:
                title = re.sub(r'\s+', ' ', title).strip()
                title = re.sub(r'^[锛?\-\s]+', '', title)
                title = title.replace('\n', ' ')

            if source:
                source = re.sub(r'\s+', ' ', source).strip()
                source = re.sub(r'^[锛?\-\s]+', '', source)
                source = source.replace('\n', ' ')

            # 浣跨敤榛樿鍊硷紙濡傛灉鎻愬彇澶辫触锛?
            if not author or author == "":
                author = "[Unknown author]"
                logger.info("  fallback: using [Unknown author]")

            if not title or title == "":
                title = "[Unknown title]"
                logger.info("  fallback: using [Unknown title]")

            if not source or source == "":
                source = source_info if source_info else "[Unknown source]"
                logger.info(f"  fallback: using source: {source}")

            logger.info("successfully extracted reference metadata")
            logger.info(f"author: {author}")
            logger.info(f"title: {title}")
            logger.info(f"source: {source}")

            return {
                "author": author,
                "title": title,
                "source": source
            }

        except Exception as e:
            import traceback
            logger.error(f"Failed to extract author/title/source metadata: {str(e)}")
            logger.error(f"traceback:\n{traceback.format_exc()}")
            return {
                "author": "[Unknown author]",
                "title": "[Unknown title]",
                "source": source_info if source_info else "[Unknown source]"
            }

    # new
    def insert_abstract_and_keywords_to_file(self, file_path: str, title: str = "", abstract: str = "",
                                             keywords: str = "", username: str = "鐢ㄦ埛"):
        """Insert title, abstract, and keywords at the start of a file."""

        # try:
        #     # 璇诲彇鍘熷鏂囦欢鍐呭
        #     with open(file_path, 'r', encoding='utf-8') as f:
        #         original_content = f.read()
        #     # ----- 鏂板锛氬己鍒舵竻鐞嗗紑澶翠换浣曢潪鏍囬鍐呭锛岀‘淇濇爣棰樻槸鏂囦欢绗竴琛?-----
        #     lines = original_content.split('\n')
        #     while lines and (not lines[0].strip() or re.match(r'^#+\s+', lines[0].strip())):
        #         # 濡傛灉绗竴琛屽凡缁忔槸鏍囬鎴栫┖琛岋紝璺冲嚭
        #         if re.match(r'^#+\s+', lines[0].strip()):
        #             break
        #         lines.pop(0)
        #     # 閲嶆柊缁勫悎
        #     original_content = '\n'.join(lines)
        #     # ----- 娓呯悊缁撴潫 -----
        #     # 銆愭柊澧炪€戞鏌ュ苟绉婚櫎閲嶅鐨勬爣棰橈紙瑙ｅ喅鍋跺彂鎬у弻鏍囬Bug锛?
        #     if title:
        #         from difflib import SequenceMatcher
        #         lines = original_content.split('\n')
        #         first_content_idx = -1
        #
        #         # 鎵惧埌绗竴涓潪绌鸿
        #         for i, line in enumerate(lines):
        #             if line.strip():
        #                 first_content_idx = i
        #                 break
        #
        #         if first_content_idx != -1:
        #             first_line = lines[first_content_idx].strip()
        #             # 妫€鏌ユ槸鍚槸鏍囬鏍煎紡锛? 寮€澶存垨 **绮椾綋**锛?
        #             heading_match = re.match(r'^(#+\s*|\*\*)(.+)', first_line)
        #             if heading_match:
        #                 # 鎻愬彇鏍囬鏂囨湰锛堝幓闄arkdown鏍囪锛?
        #                 existing_title = heading_match.group(2)
        #                 # 濡傛灉鏄矖浣撶粨灏撅紝涔熻鍘绘帀
        #                 if existing_title.endswith('**'):
        #                     existing_title = existing_title[:-2]
        #                 existing_title = existing_title.strip()
        #
        #                 # 璁＄畻鐩镐技搴?
        #                 similarity = SequenceMatcher(None, existing_title, title).ratio()
        #
        #                 # 濡傛灉鐩镐技搴珮锛屾垨鑰呭寘鍚叧绯伙紝鍒欒涓烘槸閲嶅鏍囬
        #                 # 闄嶄綆鍖呭惈鍏崇郴鐨勮鍒ら闄╋細鍙湁褰撶幇鏈夋爣棰橀暱搴帴杩戞柊鏍囬鏃舵墠鑰冭檻鍖呭惈鍏崇郴
        #                 is_contained = (title in existing_title or existing_title in title)
        #                 len_ratio = min(len(title), len(existing_title)) / max(len(title), len(existing_title))
        #
        #                 if similarity > 0.7 or (is_contained and len_ratio > 0.6):
        #                     print(f"妫€娴嬪埌閲嶅鏍囬锛屽凡绉婚櫎鍘熸枃浠跺紑澶寸殑鏍囬: {first_line}")
        #                     # 绉婚櫎璇ヨ
        #                     lines.pop(first_content_idx)
        #                     # 绉婚櫎绱ч殢鍏跺悗鐨勭┖琛?
        #                     while first_content_idx < len(lines) and not lines[first_content_idx].strip():
        #                         lines.pop(first_content_idx)
        #                     # 鏇存柊original_content
        #                     original_content = '\n'.join(lines)
        #
        #     # 鏋勫缓鏍囬銆佹憳瑕佸拰鍏抽敭璇嶉儴鍒?
        #     header_section = ""
        #
        #     # 鍒ゆ柇璇█鏄惁涓鸿嫳鏂囷紙鍩轰簬鏍囬鍜屾憳瑕佷腑鐨勪腑鏂囧瓧绗暟閲忥級
        #     # 濡傛灉涓枃瀛楃灏戜簬5涓紝璁や负鏄嫳鏂囧唴瀹?
        #     chinese_chars = re.findall(r'[\u4e00-\u9fff]', (title or "") + (abstract or ""))
        #     is_english = len(chinese_chars) < 5
        #
        #     abstract_label = "Abstract" if is_english else "鎽樿"
        #     keywords_label = "Keywords" if is_english else "鍏抽敭璇?
        #
        #     # 娣诲姞鏍囬锛堜綔涓轰竴绾ф爣棰橈級
        #     if title:
        #         header_section += f"# {title}\n\n"
        #
        #     # 娣诲姞鎽樿锛堜綔涓轰簩绾ф爣棰橈級
        #     if abstract:
        #         header_section += f"## {abstract_label}\n\n{abstract}\n\n"
        #
        #     # 娣诲姞鍏抽敭璇嶏紙浣滀负浜岀骇鏍囬锛?
        #     if keywords:
        #         header_section += f"## {keywords_label}\n\n{keywords}\n\n"
        #
        #     # 娣诲姞鐢熸垚淇℃伅鏂囨湰
        #     if is_english:
        #         footer_text = f'Generated by {username} and SciAssistant'
        #     else:
        #         footer_text = f'鏈枃绔犵敱{username}鍜孲ciAssistant鐢熸垚'
        #
        #     # 浣跨敤 font 鏍囩璁剧疆棰滆壊 (#808080 鐏拌壊)锛宒iv 鏍囩鎺у埗瀵归綈
        #     header_section += f' <div style="text-align: right;"> <font color="#808080">鈥斺€攞footer_text}</font> </div> \n\n'
        #
        #     # 灏嗘爣棰樸€佹憳瑕佸拰鍏抽敭璇嶆彃鍏ュ埌鏂囦欢寮€澶?
        #     new_content = header_section + original_content
        #
        #     # 鍐欏洖鏂囦欢
        #     with open(file_path, 'w', encoding='utf-8') as f:
        #         f.write(new_content)
        #
        #     logger.info(f"鎴愬姛灏嗘爣棰樸€佹憳瑕佸拰鍏抽敭璇嶆彃鍏ュ埌鏂囦欢寮€澶? {file_path}")
        #
        # except Exception as e:
        #     logger.error(f"璀憡: 鎻掑叆鏍囬銆佹憳瑕佸拰鍏抽敭璇嶅け璐?- {str(e)}")
        return

    def _normalize_heading_levels(self, content: str) -> str:
        """Normalize Markdown heading levels inside a section."""
        lines = content.split('\n')
        normalized_lines = []
        first_heading_found = False
        current_chapter_level = 0
        first_content_line = True

        for line in lines:
            stripped_line = line.strip()

            # 璺宠繃绌鸿锛岀洿鍒版壘鍒扮涓€涓唴瀹硅
            if first_content_line and not stripped_line:
                normalized_lines.append(line)
                continue

            # 妫€娴嬫槸鍚槸Markdown鏍囬锛堜互#寮€澶达級
            heading_match = re.match(r'^(#+)\s+(.+)$', stripped_line)

            # 妫€娴嬫槸鍚槸绮椾綋鏍囬锛?*xxx**鏍煎紡锛屽彲鑳界敤浣滄爣棰橈級
            bold_match = re.match(r'^\*\*(.+?)\*\*\s*$', stripped_line)

            if heading_match:
                hash_symbols = heading_match.group(1)
                heading_text = heading_match.group(2)
                current_level = len(hash_symbols)

                # 濡傛灉鏄涓€涓爣棰橈紝灏嗗叾璁句负浜岀骇鏍囬锛堢珷鑺傛爣棰橈級
                if not first_heading_found:
                    first_heading_found = True
                    first_content_line = False
                    current_chapter_level = current_level
                    # 纭繚绔犺妭鏍囬涓轰簩绾ф爣棰橈紝鍘婚櫎鍙兘鐨勭矖浣撴爣璁?
                    heading_text_clean = re.sub(r'^\*\*(.+?)\*\*$', r'\1', heading_text)
                    normalized_lines.append(f"## {heading_text_clean}")
                else:
                    # 璁＄畻鐩稿浜庣珷鑺傛爣棰樼殑灞傜骇宸?
                    level_diff = current_level - current_chapter_level
                    # 鏂扮殑鏍囬灞傜骇 = 3锛堝洜涓虹珷鑺傛槸2绾э級+ 灞傜骇宸?
                    new_level = max(3, 3 + level_diff)
                    # 闄愬埗鏈€澶ф爣棰樺眰绾т负6
                    new_level = min(new_level, 6)
                    # 鍘婚櫎鏍囬涓殑绮椾綋鏍囪
                    heading_text_clean = re.sub(r'^\*\*(.+?)\*\*$', r'\1', heading_text)
                    normalized_lines.append(f"{'#' * new_level} {heading_text_clean}")
            elif bold_match and not first_heading_found:
                # 濡傛灉绗竴涓唴瀹规槸绮椾綋鏂囨湰涓旇繕娌℃湁鎵惧埌鏍囬锛屽皢鍏惰浆鎹负浜岀骇鏍囬
                first_heading_found = True
                first_content_line = False
                current_chapter_level = 2
                heading_text = bold_match.group(1)
                normalized_lines.append(f"## {heading_text}")
            elif first_content_line and stripped_line and not first_heading_found:
                # 濡傛灉绗竴涓潪绌鸿涓嶆槸鏍囬鏍煎紡锛屼篃涓嶆槸绮椾綋锛屼絾鐪嬭捣鏉ュ儚鏍囬锛堢煭鏂囨湰锛屼笉浠ユ爣鐐圭粨灏撅級
                # 妫€鏌ユ槸鍚儚鏍囬锛氶暱搴€備腑锛?100瀛楃锛変笖涓嶄互鍙ュ彿銆侀棶鍙风瓑缁撳熬
                if len(stripped_line) < 100 and not re.search(r'[銆傦紵锛??!]$', stripped_line):
                    first_heading_found = True
                    first_content_line = False
                    current_chapter_level = 2
                    # 鍘婚櫎鍙兘鐨勭矖浣撴爣璁?
                    heading_text = re.sub(r'^\*\*(.+?)\*\*$', r'\1', stripped_line)
                    normalized_lines.append(f"## {heading_text}")
                else:
                    # 涓嶅儚鏍囬锛屼繚鎸佸師鏍凤紝浣嗘爣璁板凡缁忔壘鍒扮涓€涓唴瀹?
                    first_content_line = False
                    normalized_lines.append(line)
            else:
                # 闈炴爣棰樿淇濇寔涓嶅彉
                first_content_line = False
                normalized_lines.append(line)

        return '\n'.join(normalized_lines)

    def _is_failed_section_content(self, content: str) -> bool:
        if not content:
            return False
        failure_markers = [
            "Section generation failed during the model call",
            "Section generation failed after retries",
            "no valid model response, invalid format, or timeout",
            "section writer failed",
        ]
        return any(marker in content for marker in failure_markers)

    def _is_first_section(self, file_path: Union[str, Path] = "", chapter_outline: str = "") -> bool:
        text = f"{file_path or ''}\n{chapter_outline or ''}".lower()
        return (
            "part_1" in text
            or "abstract" in text
            or "\u6458\u8981" in text
            or "\u5173\u952e\u8bcd" in text
            or "keyword" in text
        )

    def _is_reference_section(self, content: str = "", file_path: Union[str, Path] = "", chapter_outline: str = "") -> bool:
        text = f"{content or ''}\n{file_path or ''}\n{chapter_outline or ''}".lower()
        return (
            "reference" in text
            or "bibliography" in text
            or "\u53c2\u8003\u6587\u732e" in text
            or "\u53c2\u8003\u6587\u732b" in text
            or "part_8" in text
        )

    def _markdown_relative_path(self, asset_path: Path, markdown_path: Union[str, Path]) -> str:
        try:
            start_dir = Path(markdown_path).parent
            rel_path = os.path.relpath(asset_path, start=start_dir)
        except Exception:
            try:
                rel_path = str(asset_path.relative_to(self.workspace_path))
            except Exception:
                rel_path = str(asset_path)
        return rel_path.replace("\\", "/")

    def _available_report_images(self) -> List[Path]:
        images = []
        for dirname in ("experiment_results", "user_uploads"):
            img_dir = self.workspace_path / dirname
            if not img_dir.exists():
                continue
            for ext in ("*.png", "*.jpg", "*.jpeg", "*.webp"):
                images.extend(sorted(img_dir.glob(ext)))

        priority_names = []
        results_md = self.workspace_path / "experiment_results" / "experiment_results.md"
        if results_md.exists():
            try:
                text = results_md.read_text(encoding="utf-8", errors="ignore")
                for name in re.findall(r"[-*]\s+([A-Za-z0-9_.+\-]+?\.(?:png|jpg|jpeg|webp))", text, flags=re.I):
                    priority_names.append(name.lower())
            except Exception:
                pass

        def sort_key(path: Path):
            name = path.name.lower()
            if name in priority_names:
                return (priority_names.index(name), name)
            return (len(priority_names) + 1, name)

        return sorted(dict((p.resolve(), p) for p in images).values(), key=sort_key)

    def _select_image_for_caption(self, caption: str, images: List[Path], used_names: set) -> Optional[Path]:
        if not images:
            return None

        caption_l = (caption or "").lower()
        aliases = [
            (["radar", "\u96f7\u8fbe"], ["radar"]),
            (["precision-recall", "p-r", "pr", "\u7cbe\u786e\u7387", "\u53ec\u56de"], ["precision_recall", "precision", "recall"]),
            (["bubble", "\u6c14\u6ce1", "scatter", "\u6563\u70b9", "fps"], ["fps_vs", "params_gflops", "gflops", "fps"]),
            (["gflops", "params", "\u53c2\u6570"], ["params_gflops", "params", "gflops"]),
            (["map50-95", "map50_95"], ["map50_95"]),
            (["map50"], ["map50"]),
            (["loss", "\u635f\u5931"], ["loss"]),
            (["lr", "\u5b66\u4e60\u7387"], ["lr"]),
            (["comparison", "bar", "\u5bf9\u6bd4", "\u67f1\u72b6"], ["comparison", "compare", "yolo_version"]),
        ]
        caption_words = set(re.findall(r"[a-z0-9_]+", caption_l))

        best = None
        best_score = -1
        for image in images:
            name = image.name.lower()
            if name in used_names:
                continue
            score = 0
            for word in caption_words:
                if len(word) >= 3 and word in name:
                    score += 2
            for triggers, targets in aliases:
                if any(trigger in caption_l for trigger in triggers):
                    score += sum(4 for target in targets if target in name)
            if score > best_score:
                best = image
                best_score = score

        if best is not None:
            return best
        for image in images:
            if image.name.lower() not in used_names:
                return image
        return None

    def _repair_image_placeholders(self, content: str, markdown_path: Union[str, Path]) -> str:
        if not content:
            return content

        images = self._available_report_images()
        if not images:
            return content

        used_names = set()
        for match in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", content):
            used_names.add(Path(match.split("#", 1)[0].split("?", 1)[0]).name.lower())

        placeholder_re = re.compile(
            r"(?m)^\s*\*?\s*[\(\uff08]\s*(?:\u6b64\u5904\u63d2\u5165)?\s*"
            r"(?:\u56fe|figure)\s*(\d*)\s*[:\uff1a]?\s*(.+?)\s*[\)\uff09]\s*\*?\s*$",
            flags=re.I,
        )
        next_number = 1

        def replace_placeholder(match):
            nonlocal next_number
            number = match.group(1).strip()
            caption = match.group(2).strip(" :\uff1a") or "Generated figure"
            image = self._select_image_for_caption(caption, images, used_names)
            if image is None:
                return match.group(0)
            used_names.add(image.name.lower())
            fig_number = number or str(next_number)
            next_number = max(next_number + 1, int(fig_number) + 1 if fig_number.isdigit() else next_number + 1)
            rel_path = self._markdown_relative_path(image, markdown_path)
            alt = re.sub(r"\s+", " ", caption).strip() or image.stem
            return f"![{alt}]({rel_path})\n*\u56fe {fig_number}. {caption}*"

        repaired = placeholder_re.sub(replace_placeholder, content)

        if placeholder_re.search(repaired):
            logger.warning("Some image placeholders could not be repaired because available images were exhausted.")
        return repaired

    def _extract_reference_field(self, block: str, labels: List[str]) -> str:
        for label in labels:
            patterns = [
                rf"\*\*\s*{label}\s*[\uff1a:]\s*\*\*\s*(.+)",
                rf"[-*]?\s*\*\*\s*{label}\s*\*\*\s*[\uff1a:]\s*(.+)",
                rf"[-*]?\s*{label}\s*[\uff1a:]\s*(.+)",
            ]
            for pattern in patterns:
                match = re.search(pattern, block, flags=re.I)
                if match:
                    return match.group(1).strip().strip("*` ")
        return ""

    def _looks_like_bibliographic_reference(self, entry: Dict[str, str]) -> bool:
        raw = (entry.get("raw") or "").strip()
        title = (entry.get("title") or "").strip()
        author = (entry.get("author") or "").strip()
        source = (entry.get("source") or "").strip()
        doi = (entry.get("doi") or "").strip()
        url = (entry.get("url") or "").strip()

        if raw:
            raw_lower = raw.lower()
            if any(marker in raw_lower for marker in ("researchgate", "ads abstract", "url-only")):
                return False
            return bool(len(raw) >= 40 and re.search(r"\b(20\d{2}|19\d{2})\b", raw) and (title or doi or '"' in raw))

        haystack = " ".join([title, author, source, url]).lower()
        rejected_markers = [
            "ads abstract",
            "researchgate",
            "pmc.ncbi.nlm.nih.gov",
            "ui.adsabs.harvard.edu",
            "abstract.",
            "unknown",
            "cannot extract",
            "captcha",
        ]
        if any(marker in haystack for marker in rejected_markers):
            return False
        if not title or len(title) < 12:
            return False
        if not (author or doi or url):
            return False
        if url and not doi and not source and not author:
            return False
        if title.lower().startswith(("online detection", "deep learning-enabled", "ccformer:")) and not author:
            return False
        return True

    def _extract_numbered_reference_entries(self, text: str) -> List[Dict[str, str]]:
        entries = []
        numbered_blocks = re.finditer(
            r"(?ms)^\s*\[(\d+)\]\s+(.+?)(?=^\s*\[\d+\]\s+|\Z)",
            text or "",
        )
        for match in numbered_blocks:
            raw = re.sub(r"\s+", " ", match.group(2)).strip()
            raw = raw.rstrip()
            if len(raw) < 40:
                continue

            title = ""
            title_match = re.search(r'"([^"]{8,220})"', raw)
            if title_match:
                title = title_match.group(1).strip()

            author = ""
            if title_match:
                author = raw[:title_match.start()].strip(" ,.;")
            elif "," in raw:
                author = raw.split(",", 1)[0].strip(" ,.;")

            year_match = re.search(r"\b(20\d{2}|19\d{2})\b", raw)
            doi_match = re.search(r"\b10\.\d{4,9}/[^\s,;，；)）]+", raw, flags=re.I)
            url_match = re.search(r"https?://[^\s)）]+", raw)

            entries.append({
                "author": author,
                "title": title,
                "source": "",
                "year": year_match.group(1) if year_match else "",
                "doi": doi_match.group(0).rstrip(".") if doi_match else "",
                "url": url_match.group(0).rstrip(".,;") if url_match else "",
                "raw": raw,
            })
        return entries

    def _extract_reference_entries_from_jsonl(self, text: str) -> List[Dict[str, str]]:
        entries = []
        for line in (text or "").splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                data = json.loads(line)
            except Exception:
                continue

            core = data.get("core_content") or ""
            if not core or "failed" in str(data.get("doc_time", "")).lower():
                continue
            if re.search(r"captcha|blocked|unknown|cannot extract", core, flags=re.I):
                continue

            title = ""
            author = ""
            title_author_match = re.search(
                r"Title/Authors\s*:\s*(.+?)(?:\.\s+Method:|\n|$)",
                core,
                flags=re.I | re.S,
            )
            if title_author_match:
                title_author = re.sub(r"\s+", " ", title_author_match.group(1)).strip()
                if ";" in title_author:
                    title, author = [part.strip(" .;") for part in title_author.split(";", 1)]
                elif " / " in title_author:
                    title, author = [part.strip(" .;") for part in title_author.split(" / ", 1)]
                else:
                    title = title_author.strip(" .;")
            if not title:
                title_match = re.search(r"Title\s*:\s*(.+?)(?:\.\s+(?:Authors?|Method|Key Findings|Figures):|\n|$)", core, flags=re.I | re.S)
                if title_match:
                    title = re.sub(r"\s+", " ", title_match.group(1)).strip(" .;")
            if not author:
                author_match = re.search(r"Authors?\s*:\s*(.+?)(?:\.\s+(?:Method|Key Findings|Figures):|\n|$)", core, flags=re.I | re.S)
                if author_match:
                    author = re.sub(r"\s+", " ", author_match.group(1)).strip(" .;")
            if not title:
                continue

            source_path = data.get("file_path") or ""
            year = ""
            year_match = re.search(r"\b(20\d{2}|19\d{2})\b", str(data.get("doc_time", "")) + " " + core)
            if year_match:
                year = year_match.group(1)
            doi_match = re.search(r"\b10\.\d{4,9}/[^\s,;，；)）]+", core, flags=re.I)

            entries.append({
                "author": author,
                "title": title,
                "source": Path(source_path).name if source_path else "",
                "year": year,
                "doi": doi_match.group(0).rstrip(".") if doi_match else "",
                "url": "",
            })
        return entries

    def _build_reference_section(self, target_file_path: Union[str, Path] = "report/references.md") -> str:
        structured_path = self.workspace_path / "research" / "references.json"
        if structured_path.is_file():
            try:
                payload = json.loads(structured_path.read_text(encoding="utf-8"))
                entries = []
                seen = set()
                for position, item in enumerate(payload if isinstance(payload, list) else [], 1):
                    if not isinstance(item, dict):
                        continue
                    title = re.sub(r"\s+", " ", str(item.get("title") or "")).strip()
                    if not title:
                        continue
                    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", str(item.get("doi") or ""), flags=re.I).strip()
                    key = ("doi", doi.lower()) if doi else ("title", re.sub(r"\W+", "", title.lower()))
                    if key in seen:
                        continue
                    seen.add(key)
                    entries.append((position, item, title, doi))
                if not entries:
                    return "# 参考文献\n\n> Reference generation failed: research/references.json contains no valid records.\n"
                lines = ["# 参考文献", ""]
                for position, item, title, doi in entries:
                    parts = [
                        re.sub(r"\s+", " ", str(item.get("authors") or "")).strip(),
                        title,
                        re.sub(r"\s+", " ", str(item.get("venue") or "")).strip(),
                        str(item.get("year") or "").strip(),
                        f"DOI: {doi}" if doi else str(item.get("url") or "").strip(),
                    ]
                    lines.append(f"[{position}] " + ". ".join(part.rstrip(".") for part in parts if part) + ".")
                return "\n".join(lines) + "\n"
            except Exception as exc:
                logger.error("Structured bibliography could not be built from %s: %s", structured_path, exc)
                return f"# 参考文献\n\n> Reference generation failed: {exc}\n"

        # Compatibility fallback for old workspaces created before the
        # structured catalogue existed. New runs must use references.json.
        source_files = []
        for pattern in (
            "*literature*.md",
            "*Literature*.md",
            "*review*.md",
            "*Review*.md",
            "*reference*.md",
            "*Reference*.md",
            "*citation*.md",
            "*Citation*.md",
            "doc_analysis/*.jsonl",
            "research_output/*.md",
            "research/**/*.md",
            "url_crawler_save_files/research/*_url_ref.txt",
            "url_crawler_save_files/research/*.jsonl",
        ):
            source_files.extend(sorted(self.workspace_path.glob(pattern)))
        source_files = [
            path for path in dict.fromkeys(source_files)
            if "report" not in path.relative_to(self.workspace_path).parts
        ]

        entries = []
        seen = set()
        for source in source_files:
            try:
                text = source.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                continue

            if source.name.lower().endswith("_url_ref.txt"):
                title = self._extract_reference_field(text, ["Title", "\u6807\u9898", "\u9898\u540d"])
                url = self._extract_reference_field(text, ["URL", "Link", "\u94fe\u63a5"])
                year = ""
                text_without_urls = re.sub(r"https?://\S+", " ", text)
                year_match = re.search(r"\b(20\d{2}|19\d{2})\b", text_without_urls)
                if year_match:
                    year = year_match.group(1)
                entry = {
                    "author": "",
                    "title": title,
                    "source": source.stem.replace("_url_ref", ""),
                    "year": year,
                    "doi": "",
                    "url": url,
                }
                if self._looks_like_bibliographic_reference(entry):
                    key = (title or url).lower()
                    if key not in seen:
                        seen.add(key)
                        entries.append(entry)
                continue

            if source.suffix.lower() == ".jsonl":
                for entry in self._extract_reference_entries_from_jsonl(text):
                    if not self._looks_like_bibliographic_reference(entry):
                        continue
                    key = (entry.get("doi") or entry.get("title") or entry.get("url")).lower()
                    if key in seen:
                        continue
                    seen.add(key)
                    entries.append(entry)
                continue

            for entry in self._extract_numbered_reference_entries(text):
                if not self._looks_like_bibliographic_reference(entry):
                    continue
                key = (entry.get("doi") or entry.get("title") or entry.get("raw") or entry.get("url")).lower()
                if key in seen:
                    continue
                seen.add(key)
                entries.append(entry)

            blocks = re.split(r"(?m)^#{2,4}\s+", text)
            if len(blocks) == 1:
                blocks = re.split(r"\n\s*---\s*\n", text)

            for block in blocks:
                block = block.strip()
                if len(block) < 20:
                    continue
                title = self._extract_reference_field(block, ["Title", "\u6807\u9898", "\u9898\u540d"])
                author = self._extract_reference_field(block, ["Authors", "Author", "\u4f5c\u8005"])
                source_name = self._extract_reference_field(
                    block,
                    ["Journal/Conference", "Journal", "Source", "\u671f\u520a/\u4f1a\u8bae", "\u671f\u520a", "\u6765\u6e90"],
                )
                year = self._extract_reference_field(block, ["Year", "\u5e74\u4efd"])
                doi = self._extract_reference_field(block, ["DOI", "doi"])
                url = self._extract_reference_field(block, ["URL", "Link", "\u94fe\u63a5"])

                if not title:
                    first_line = block.splitlines()[0].strip()
                    first_line = re.sub(r"^\s*\d+[\.\uff0e]\s*", "", first_line)
                    first_line = re.sub(r"^\s*\u53c2\u8003\u6587\u732e\s*\d*\s*[\uff1a:]?\s*", "", first_line)
                    if 8 <= len(first_line) <= 220 and not first_line.startswith("#"):
                        title = first_line.strip()
                if title and any(
                    bad in title.lower()
                    for bad in (
                        "summary",
                        "conclusion",
                        "search note",
                        "retrieval",
                        "\u6587\u732e\u603b\u7ed3",
                        "\u7efc\u8ff0\u603b\u7ed3",
                        "\u6587\u732e\u8986\u76d6",
                        "\u68c0\u7d22\u8bf4\u660e",
                        "\u4e3b\u9898\u5206\u6790",
                        "\u8986\u76d6\u4e3b\u9898",
                    )
                ):
                    continue

                if not year:
                    block_without_urls = re.sub(r"https?://\S+", " ", block)
                    year_match = re.search(r"\b(20\d{2}|19\d{2})\b", block_without_urls)
                    year = year_match.group(1) if year_match else ""
                if not doi:
                    doi_match = re.search(r"\b10\.\d{4,9}/[^\s,;，；)）]+", block, flags=re.I)
                    doi = doi_match.group(0).rstrip(".") if doi_match else ""
                if not url:
                    url_match = re.search(r"https?://[^\s)）]+", block)
                    url = url_match.group(0).rstrip(".,;") if url_match else ""

                if not title and not doi and not url:
                    continue

                key = (doi or title or url).lower()
                if key in seen:
                    continue
                seen.add(key)
                entries.append({
                    "author": author,
                    "title": title,
                    "source": source_name,
                    "year": year,
                    "doi": doi,
                    "url": url,
                })
                if not self._looks_like_bibliographic_reference(entries[-1]):
                    entries.pop()
                    seen.discard(key)
                    continue

        if not entries:
            return (
                "# \u53c2\u8003\u6587\u732e\n\n"
                "> Reference generation failed: no parseable reference metadata was found in collected research files.\n"
            )

        lines = ["# \u53c2\u8003\u6587\u732e", ""]
        for idx, entry in enumerate(entries, 1):
            if entry.get("raw"):
                lines.append(f"[{idx}] {entry['raw'].rstrip('.')}.")
                continue
            parts = []
            if entry["author"]:
                parts.append(entry["author"].rstrip("."))
            if entry["title"]:
                parts.append(entry["title"].rstrip("."))
            if entry["source"]:
                source_part = entry["source"].rstrip(".")
                if entry["year"] and entry["year"] not in source_part:
                    source_part = f"{source_part}, {entry['year']}"
                parts.append(source_part)
            elif entry["year"]:
                parts.append(entry["year"])
            if entry["doi"]:
                parts.append(f"DOI: {entry['doi']}")
            elif entry["url"] and entry["source"]:
                parts.append(entry["url"])
            lines.append(f"[{idx}] " + ". ".join(part for part in parts if part) + ".")
        return "\n".join(lines) + "\n"

    def _extract_user_fact_value(self, notes: str, labels: List[str]) -> str:
        if not notes:
            return ""
        label_pattern = "|".join(re.escape(label) for label in labels)
        patterns = [
            rf"(?:{label_pattern})\s*[:：/]\s*([^\n\r；;。]+)",
            rf"(?:{label_pattern})\s*为\s*([^\n\r；;。]+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, notes, flags=re.I)
            if match:
                return re.sub(r"\s+", " ", match.group(1)).strip()
        return ""

    def _resolve_clean_title(self, user_query: str = "", outline: str = "", notes: str = "") -> str:
        """\u4ece\u591a\u4e2a\u6765\u6e90\u89e3\u6790\u8bba\u6587\u771f\u5b9e\u6807\u9898,\u8df3\u8fc7 '\u8bba\u6587\u6807\u9898' \u8fd9\u7c7b\u5360\u4f4d\u7b26.
        \u4f18\u5148\u7ea7:user_query/notes \u4e2d\u663e\u5f0f\u7684"\u9898\u76ee/\u6807\u9898" > \u5927\u7eb2\u9996\u6761\u975e\u5360\u4f4d\u884c > user_query \u9996\u884c > \u672a\u547d\u540d\u7814\u7a76."""
        placeholder = re.compile(r"^\s*(\u8bba\u6587\u6807\u9898|\u8bba\u6587\u9898\u76ee|paper\s*title|\u672a\u547d\u540d\u7814\u7a76|title|\u6807\u9898)\s*$", re.I)
        # 1. \u4ece user_query / notes \u4e2d\u663e\u5f0f\u63d0\u53d6"\u9898\u76ee/\u6807\u9898:X"\u6216"\u9898\u76ee\u662f\u300cX\u300d"
        for src in (user_query or "", notes or ""):
            m = re.search(
                r"(?:\u8bba\u6587\u9898\u76ee|\u9898\u76ee|\u6807\u9898|title)\s*[:\uff1a\u662f\u4e3a]*\s*[\u201c\u300c\u300e\"'\u300a]?([^\u201c\u201d\u300c\u300d\u300e\u300f\"'\u300a\u300b\n]{4,90})",
                src, flags=re.I)
            if m:
                cand = m.group(1).strip().strip("\u201c\u201d\u300c\u300d\u300e\u300f\"'\u300a\u300b\uff08\uff09()\u3010\u3011 \uff1a:")
                if cand and not placeholder.match(cand):
                    return cand[:120]
        # 2. \u5927\u7eb2\u4e2d\u7b2c\u4e00\u6761\u975e\u5360\u4f4d\u3001\u975e\u6458\u8981/\u5173\u952e\u8bcd\u7684\u6807\u9898\u884c
        for line in (outline or "").splitlines():
            cleaned = line.strip().lstrip("#").strip()
            if cleaned and not placeholder.match(cleaned) and not re.search(
                    r"abstract|keyword|\u6458\u8981|\u5173\u952e\u8bcd", cleaned, flags=re.I):
                return cleaned[:120]
        # 3. user_query \u9996\u4e2a\u975e\u7a7a\u884c(\u622a\u65ad)
        first = next((l.strip() for l in (user_query or "").splitlines() if l.strip()), "")
        if first and not placeholder.match(first):
            return first[:120]
        return "\u672a\u547d\u540d\u7814\u7a76"

    def _build_first_section_fallback(
        self,
        target_file_path: Union[str, Path],
        current_chapter_outline: str = "",
        user_query: str = "",
        authoritative_user_notes: str = "",
    ) -> str:
        title = self._resolve_clean_title(user_query, current_chapter_outline, authoritative_user_notes)

        method = self._extract_user_fact_value(
            authoritative_user_notes,
            ["\u7814\u7a76\u76ee\u6807", "\u65b9\u6cd5", "\u6a21\u578b", "method", "model"],
        )
        dataset = self._extract_user_fact_value(
            authoritative_user_notes,
            ["\u6570\u636e\u96c6", "\u6837\u672c", "\u56fe\u50cf", "dataset", "images"],
        )
        result = self._extract_user_fact_value(
            authoritative_user_notes,
            ["\u7ed3\u679c", "\u7cbe\u5ea6", "mAP", "precision", "recall", "result"],
        )

        fact_sentences = []
        if method:
            fact_sentences.append(f"\u65b9\u6cd5\u4e0a\uff0c\u7814\u7a76\u56f4\u7ed5 {method} \u5c55\u5f00\u3002")
        if dataset:
            fact_sentences.append(f"\u6570\u636e\u4e0e\u5b9e\u9a8c\u8bbe\u7f6e\u4f9d\u636e\u7528\u6237\u63d0\u4f9b\u6750\u6599\uff1a{dataset}\u3002")
        if result:
            fact_sentences.append(f"\u53ef\u6838\u9a8c\u7684\u7ed3\u679c\u4fe1\u606f\u5305\u62ec\uff1a{result}\u3002")
        if not fact_sentences:
            fact_sentences.append(
                "\u7531\u4e8e\u7b2c\u4e00\u7ae0\u6a21\u578b\u751f\u6210\u672a\u8fd4\u56de\u6709\u6548\u5185\u5bb9\uff0c"
                "\u672c\u6458\u8981\u4ec5\u4fdd\u5b88\u603b\u7ed3\u7528\u6237\u5df2\u63d0\u4f9b\u7684\u7814\u7a76\u4e3b\u9898\uff0c\u4e0d\u8865\u5199\u672a\u6838\u9a8c\u6570\u636e\u3002"
            )

        abstract = (
            f"\u672c\u7814\u7a76\u9762\u5411\u201c{title}\u201d\u4efb\u52a1\uff0c"
            "\u76ee\u6807\u662f\u57fa\u4e8e\u7528\u6237\u63d0\u4f9b\u7684\u771f\u5b9e\u6750\u6599\u6574\u7406\u8bba\u6587\u7684\u7814\u7a76\u95ee\u9898\u3001\u65b9\u6cd5\u548c\u5b9e\u9a8c\u4f9d\u636e\u3002"
            + "".join(fact_sentences)
            + "\u4e3a\u907f\u514d\u5f15\u5165\u4e0d\u53ef\u6838\u9a8c\u7ed3\u8bba\uff0c\u672c\u6587\u540e\u7eed\u7ae0\u8282\u4ec5\u5bf9\u5df2\u6709\u5b9e\u9a8c\u8bb0\u5f55\u548c\u53ef\u8ffd\u6eaf\u6587\u732e\u8fdb\u884c\u5206\u6790\u3002"
        )
        keywords = "\uff1b".join([kw for kw in [title.split()[0] if title else "", "\u65e0\u635f\u68c0\u6d4b", "\u6df1\u5ea6\u5b66\u4e60", "\u591a\u6a21\u6001\u878d\u5408"] if kw])
        return f"# {title}\n\n## \u6458\u8981\n\n{abstract}\n\n## \u5173\u952e\u8bcd\n\n{keywords}\n"

    def _postprocess_merged_report(self, content: str, output_file: Union[str, Path]) -> str:
        content = re.sub(r"(?m)^(#{1,6})\s+#+\s+", r"\1 ", content or "")
        content = self._repair_image_placeholders(content, output_file)
        return content

    def _audit_merged_report(self, content: str, output_file: Union[str, Path]) -> List[str]:
        issues = audit_manuscript_text(content)
        if issues:
            logger.warning(
                "Final manuscript audit found possible quality issues in %s: %s",
                output_file,
                ", ".join(issues),
            )
        else:
            logger.info("Final manuscript audit passed for %s", output_file)
        return issues

    ##
    def merge_reports(self, section_contents, output_file, workspace_root: Path = None, unique_id=None):
        import os, re
        from pathlib import Path
        import logging
        logger = logging.getLogger(__name__)

        report_files = []
        for section_content in section_contents:
            if isinstance(section_content, dict):
                file_path = section_content.get('file_path')
            else:
                file_path = section_content
            if file_path:
                full_path = self.workspace_path / file_path
                report_files.append(full_path)

        def extract_index(file_path):
            filename = os.path.basename(file_path)
            match = re.search(r'part_(\d+)\.md', filename)
            if match:
                return int(match.group(1))
            return None

        indexed_files = []
        for file_path in report_files:
            idx = extract_index(file_path)
            if idx is not None:
                indexed_files.append((idx, file_path))

        indexed_files.sort(key=lambda x: x[0])

        if not indexed_files:
            logger.warning("No part_*.md files found; cannot merge.")
            return None

        try:
            merged_content = ""
            # 1. 鏋佸叾鏆村姏鐨勭函鍑€鎷兼帴锛屽師姹佸師鍛?
            for idx, file_path in indexed_files:
                filename = os.path.basename(file_path)
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as infile:
                        content = infile.read().strip()
                        if self._is_failed_section_content(content) and self._is_reference_section(content, file_path):
                            logger.warning(f"Reference section failed in {filename}; rebuilding from collected research files.")
                            content = self._build_reference_section(file_path)
                            try:
                                file_path.write_text(content, encoding="utf-8")
                            except Exception as write_err:
                                logger.warning(f"Failed to overwrite repaired reference section {filename}: {write_err}")
                        elif self._is_failed_section_content(content) and self._is_first_section(file_path):
                            logger.warning(f"First section failed in {filename}; rebuilding a conservative fallback.")
                            content = self._build_first_section_fallback(file_path)
                            try:
                                file_path.write_text(content, encoding="utf-8")
                            except Exception as write_err:
                                logger.warning(f"Failed to overwrite repaired first section {filename}: {write_err}")
                        elif self._is_failed_section_content(content):
                            error_msg = (
                                f"Section file {filename} contains a generation failure placeholder. "
                                "Regenerate the failed section before merging the final report."
                            )
                            logger.error(error_msg)
                            return None
                        if content:
                            merged_content += content + "\n\n"
                            logger.info(f"Merged file content: {filename}")
                except Exception as e:
                    logger.warning(f"Failed to read file {filename}: {e}")

            # 2. 寮鸿鍒涘缓鐖剁洰褰曪紙闃叉鐩綍涓嶅瓨鍦級
            Path(output_file).parent.mkdir(parents=True, exist_ok=True)

            merged_content = self._postprocess_merged_report(merged_content, output_file)
            audit_issues = self._audit_merged_report(merged_content, output_file)

            # 3. 灏嗘嫾鎺ュソ鐨勫唴瀹逛竴娆℃€у啓鍏?final_report.md
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(merged_content)

            logger.info(f"Final report merge completed: {output_file}")

            # 4. 鐢熸垚 PDF
            try:
                pdf_path = Path(output_file).with_suffix('.pdf')
                success = _generate_pdf_with_chrome(merged_content, pdf_path, str(self.workspace_path))
                if success:
                    logger.info(f"PDF generated successfully: {pdf_path}")
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")

            return {
                "title": "Academic Paper",
                "abstract": "Report generated",
                "keywords": "",
                "audit_issues": audit_issues,
            }

        except Exception as e:
            logger.error(f"Fatal error while merging report: {e}")
            raise
    def _extract_authoritative_fact_pairs(self, notes: str) -> List[Dict[str, str]]:
        # Extract conservative explicit key-value facts from user notes.
        if not notes:
            return []

        field_aliases = {
            "model": ["model", "model name", "architecture", "network", "method", "baseline"],
            "dataset": ["dataset", "dataset name", "data", "data source", "sample size", "image count"],
            "classes": ["class", "classes", "category", "categories", "label", "labels"],
            "hardware": ["hardware", "device", "gpu", "cpu", "platform"],
            "optimizer": ["optimizer"],
            "epochs": ["epoch", "epochs"],
            "batch": ["batch", "batch size", "batch_size"],
            "learning_rate": ["learning rate", "lr"],
            "split": ["split", "train val test"],
            "metric": ["metric", "metrics", "result", "results", "accuracy", "precision", "recall"],
        }
        alias_to_field = {alias: field for field, aliases in field_aliases.items() for alias in aliases}

        facts = []
        allowed_explicit_keys = {
            "数据集名称",
            "来源 / 是否公开",
            "类别数与类别名（对应 data.yaml）",
            "类别数类别名（对应 data.yaml）",
            "图像总数",
            "训练/验证/测试划分",
            "图像分辨率 / 标注方式",
            "是否有数据增强",
            "硬件平台",
            "优化器",
            "训练轮数",
            "batch",
            "batch size",
            "learning rate",
            "模型名称",
            "方法名称",
        }
        for raw_line in notes.splitlines():
            line = raw_line.strip().strip("|")
            if not line or len(line) > 320:
                continue
            line = re.sub(r"^\s*[-*+]\s*", "", line)
            match = re.match(r"^(.{1,50}?)(?:[:=]|：)\s*(.{1,220})$", line)
            if not match:
                continue
            key = match.group(1).strip().strip("*#` ")
            value = match.group(2).strip().strip("*#` ;；")
            if not key or not value:
                continue
            if re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", key) and re.search(r"\b(int|float|bool|str|list|dict)\b", value):
                continue
            key_norm = key.lower().replace("_", " ").strip()
            canonical = "explicit"
            for alias, field in alias_to_field.items():
                if alias == key_norm or alias in key_norm:
                    canonical = field
                    break
            if canonical == "explicit" and key not in allowed_explicit_keys:
                continue
            facts.append({"field": canonical, "key": key, "value": value})
        return facts

    def _normalized_fact_value(self, value: str) -> str:
        value = (value or "").lower()
        value = re.sub(r"[\s,;:。.!?()\[\]{}<>\"'`*_/-]+", "", value)
        return value

    def _extract_numbers(self, value: str) -> List[str]:
        return re.findall(r"\d+(?:\.\d+)?", value or "")

    def _version_family(self, token: str) -> str:
        letters = "".join(re.findall(r"[A-Za-z]+", token or "")).lower()
        return letters[:-1] if letters.endswith("v") and len(letters) > 1 else letters

    def _find_authoritative_conflicts(self, text: str, notes: str) -> List[str]:
        # Conservatively detect generic conflicts with explicit user-provided facts.
        if not notes or not text:
            return []

        facts = self._extract_authoritative_fact_pairs(notes)
        if not facts:
            return []

        issues = []
        text_norm = self._normalized_fact_value(text)
        text_lower = text.lower()
        # Keep validation conservative. Broad fields such as "metric" or
        # "learning_rate" often contain lists of names or several unrelated
        # numbers; treating every nearby number as a conflict causes repeated
        # retries and long-context timeouts.
        numeric_fields = {"epochs", "batch", "split"}

        field_aliases = {
            "model": ["model", "architecture", "network", "method", "baseline"],
            "dataset": ["dataset", "data", "sample", "image"],
            "classes": ["class", "category", "label"],
            "hardware": ["hardware", "device", "gpu", "platform"],
            "optimizer": ["optimizer"],
            "epochs": ["epoch", "epochs"],
            "batch": ["batch"],
            "learning_rate": ["learning rate", "lr"],
            "split": ["训练/验证/测试", "train/val/test", "train validation test", "train-val-test"],
            "metric": ["metric", "result", "accuracy", "precision", "recall"],
        }

        for fact in facts:
            field = fact["field"]
            key = fact["key"]
            value = fact["value"]
            value_norm = self._normalized_fact_value(value)
            if not value_norm:
                continue

            if field == "classes":
                class_names = re.findall(r"[A-Za-z]+_Apple", value)
                mentions_class_fact = any(term in text_lower for term in ["类别", "class", "classes", "category", "categories", "label"])
                if mentions_class_fact and class_names:
                    missing = [name for name in class_names if name not in text]
                    generated_class_names = re.findall(r"[A-Za-z]+_Apple", text)
                    if missing and generated_class_names:
                        issues.append(
                            f"Generated class labels conflict with user fact '{key}: {value}'."
                        )
                continue

            if field in numeric_fields:
                if field == "split" and re.search(r"\blambda\b|\bdef\b|\.view\(|\.transpose\(|self\.", value, flags=re.I):
                    continue
                expected_numbers = set(self._extract_numbers(value))
                if expected_numbers:
                    aliases = field_aliases.get(field, [])
                    escaped_aliases = "|".join(re.escape(alias) for alias in aliases)
                    if escaped_aliases:
                        nearby_numbers = set()
                        for match in re.finditer(escaped_aliases, text, flags=re.IGNORECASE):
                            window = text[max(0, match.start() - 80): match.end() + 120]
                            nearby_numbers.update(self._extract_numbers(window))
                        if nearby_numbers and not (nearby_numbers & expected_numbers):
                            issues.append(
                                f"Generated {field} value conflicts with user fact '{key}: {value}'."
                            )
                            continue

            key_pattern = re.escape(key)
            for match in re.finditer(rf"{key_pattern}\s*[:锛?]\s*(.{{1,180}})", text, flags=re.IGNORECASE):
                generated_value = match.group(1).splitlines()[0].strip()
                generated_norm = self._normalized_fact_value(generated_value)
                if generated_norm and value_norm not in generated_norm and generated_norm not in value_norm:
                    issues.append(
                        f"Generated value for '{key}' conflicts with user fact '{key}: {value}'."
                    )
                    break

        return issues

    # Implement concat_section_files.
    def concat_section_files(self, section_files, final_file_path):
        try:
            logger.info(f"鎴戠幇鍦ㄥ紑濮嬭皟鐢╟oncat_section_files浜嗭細{section_files}, {final_file_path}")

            # 馃毃 鏍稿績淇锛氬繀椤诲厛杞崲涓哄伐浣滃尯鐨勭粷瀵硅矾寰勶紝鐒跺悗鍐嶅垱寤烘枃浠跺す锛?
            if not os.path.isabs(final_file_path):
                final_file_path = self.workspace_path / final_file_path
            else:
                final_file_path = Path(final_file_path)

            output_dir = final_file_path.parent
            if output_dir:
                output_dir.mkdir(parents=True, exist_ok=True)

            # ========== 断章检测:防止某一章被漏写却静默拼接出残缺论文 ==========
            # 磁盘扫描 report 目录下实际写出的 part_*.md,若编号不连续(中间缺章),
            # 直接返回可操作的失败信息,让上层 writer 的 LLM 循环补写缺失章节后重新拼接.
            try:
                present_parts = []
                for p in output_dir.glob("part_*.md"):
                    m = re.match(r'part_(\d+)\.md$', p.name)
                    if m:
                        present_parts.append(int(m.group(1)))
                if present_parts:
                    present_parts = sorted(set(present_parts))
                    expected = list(range(present_parts[0], present_parts[-1] + 1))
                    missing = [n for n in expected if n not in present_parts]
                    if missing:
                        missing_files = ", ".join(f"part_{n}.md" for n in missing)
                        # part_N 对应正文第 (N-1) 章(part_1 是标题/摘要页)
                        missing_chapters = ", ".join(f"第{n - 1}章" for n in missing)
                        error_msg = (
                            f"检测到论文章节缺失:{missing_files}({missing_chapters})尚未写出,"
                            f"当前仅有 {', '.join('part_%d.md' % n for n in present_parts)}.\n"
                            f"请先调用 section_writer 把缺失的章节写出来(target_file_path 分别设为 {missing_files}),"
                            f"然后再调用 concat_section_files 合并.禁止跳过任何章节直接合并."
                        )
                        logger.error(f"concat_section_files 断章检测失败: {error_msg}")
                        return MCPToolResult(
                            success=False,
                            error=error_msg,
                            data={"merged_files": 0, "output_path": str(final_file_path)},
                            metadata={
                                'final_file_path': str(final_file_path),
                                'missing_parts': missing_files,
                                'missing_chapter_detected': True,
                            }
                        )
            except Exception as gap_err:
                logger.warning(f"断章检测异常(跳过检测,继续合并): {gap_err}")
            # ===================================================================

            unique_id = f"msg_{int(time.time() * 1000)}"

            # 1. 鍚堝苟绔犺妭锛屽苟鎻愬彇鐢熸垚濂界殑 鏍囬銆佹憳瑕併€佸叧閿瘝
            abstract_keywords = self.merge_reports(section_files, final_file_path, self.workspace_path, unique_id)

            if abstract_keywords is None:
                return MCPToolResult(
                    success=False,
                    error="Failed to merge section files because one or more required sections were not generated.",
                    data={
                        "merged_files": 0,
                        "output_path": str(final_file_path),
                    },
                    metadata={
                        'final_file_path': str(final_file_path),
                        'section_count': len(section_files),
                        'merge_failed': True,
                    }
                )

            try:
                final_text = final_file_path.read_text(encoding="utf-8", errors="ignore")
                user_notes = []
                uploads_dir = self.workspace_path / "user_uploads"
                if uploads_dir.exists():
                    for path in sorted(uploads_dir.iterdir()):
                        if path.suffix.lower() in {".txt", ".md"}:
                            user_notes.append(path.read_text(encoding="utf-8", errors="ignore"))
                user_notes_text = "\n\n".join(user_notes)
                conflicts = self._find_authoritative_conflicts(final_text, user_notes_text)
                if conflicts:
                    conflict_msg = "; ".join(conflicts)
                    logger.error(f"Authoritative fact validation failed: {conflict_msg}")
                    return MCPToolResult(
                        success=False,
                        error=f"Authoritative fact validation failed: {conflict_msg}",
                        data={
                            "merged_files": len(section_files),
                            "output_path": str(final_file_path),
                            "conflicts": conflicts,
                        },
                        metadata={
                            'final_file_path': str(final_file_path),
                            'section_count': len(section_files),
                            'validation_failed': True,
                        }
                    )
            except Exception as validation_err:
                logger.warning(f"Authoritative fact validation skipped: {validation_err}")

            pdf_path = Path(final_file_path).with_suffix('.pdf')
            pdf_success = pdf_path.exists()

            return MCPToolResult(
                success=True,
                data={
                    "merged_files": len(section_files),
                    "output_path": str(final_file_path),
                    "pdf_path": str(pdf_path) if pdf_success else None,
                    "abstract": abstract_keywords.get("abstract", ""),
                    "keywords": abstract_keywords.get("keywords", "")
                },
                metadata={
                    'final_file_path': str(final_file_path),
                    'pdf_path': str(pdf_path) if pdf_success else None,
                    'section_count': len(section_files),
                }
            )
        except Exception as e:
            logger.error(f"Concatenate section files failed: {e}", exc_info=True)
            return MCPToolResult(success=False, error=str(e))

    def _validate_file_allocation(
            self,
            classification_result: str,
            user_file_count: int,
            research_file_count: int,
            has_user_files: bool
    ) -> Dict:
        """
        楠岃瘉LLM鐨勬枃浠跺垎閰嶇粨鏋滄槸鍚鍚堣姹?

        Args:
            classification_result: LLM杩斿洖鐨勫垎绫荤粨鏋?
            user_file_count: 鐢ㄦ埛鏂囦欢鏁伴噺
            research_file_count: research鏂囦欢鏁伴噺
            has_user_files: 鏄惁鏈夌敤鎴锋枃浠?

        Returns:
            楠岃瘉缁撴灉瀛楀吀锛屽寘鍚玽alid, message, expected, actual瀛楁
        """
        try:
            # 缁熻鍒嗛厤鐨勬枃浠舵暟閲?
            user_files_assigned = classification_result.count('user_uploads')
            research_files_assigned = classification_result.count('research/')

            # 璁＄畻棰勬湡鍊?
            if has_user_files and research_file_count > 0:
                # 娣峰悎鍦烘櫙
                expected_user = user_file_count  # 100%
                if user_file_count >= 8:
                    expected_research_pct = 70
                elif user_file_count >= 5:
                    expected_research_pct = 80
                else:
                    expected_research_pct = 90
                expected_research_min = int(research_file_count * expected_research_pct / 100)

                # 楠岃瘉
                user_valid = user_files_assigned >= expected_user * 0.8  # 鍏佽80%瀹瑰樊
                research_valid = research_files_assigned >= expected_research_min * 0.5  # 鍏佽50%瀹瑰樊

                if not user_valid or not research_valid:
                    return {
                        "valid": False,
                        "message": f"Mixed scenario: User files {user_files_assigned}/{expected_user}, Research files {research_files_assigned}/{expected_research_min}",
                        "expected": f"User: {expected_user} (100%), Research: {expected_research_min} ({expected_research_pct}%)",
                        "actual": f"User: {user_files_assigned}, Research: {research_files_assigned}"
                    }

            elif research_file_count > 0:
                # 鍙湁research鏂囦欢鍦烘櫙
                if research_file_count >= 30:
                    expected_research_pct = 85
                elif research_file_count >= 20:
                    expected_research_pct = 90
                elif research_file_count >= 10:
                    expected_research_pct = 95
                else:
                    expected_research_pct = 100

                expected_research_min = int(research_file_count * expected_research_pct / 100)
                research_valid = research_files_assigned >= expected_research_min * 0.7  # 鍏佽70%瀹瑰樊

                if not research_valid:
                    return {
                        "valid": False,
                        "message": f"Research-only scenario: {research_files_assigned}/{expected_research_min} files assigned",
                        "expected": f"Research: {expected_research_min} ({expected_research_pct}%)",
                        "actual": f"Research: {research_files_assigned}"
                    }

            return {
                "valid": True,
                "message": "File allocation meets requirements",
                "expected": "N/A",
                "actual": f"User: {user_files_assigned}, Research: {research_files_assigned}"
            }

        except Exception as e:
            logger.error(f"Validation error: {e}")
            return {
                "valid": True,  # 楠岃瘉澶辫触鏃朵笉闃绘娴佺▼
                "message": f"Validation error: {e}",
                "expected": "N/A",
                "actual": "N/A"
            }

    def search_result_classifier(
            self,
            outline: str,
            key_files: List[Dict],
            model: str = None,
            temperature: float = 0.3,
            max_tokens: int = 4000
    ) -> MCPToolResult:
        """
        if not model or model == "pangu_auto":
            model = get_config().model_name
        Classify and organize search result files according to a structured outline for comprehensive long-form content generation.

        Args:
            outline: Structured outline defining the sections and subsections for organizing the long-form content
            key_files: List of key files to classify
            model: AI model to use for classification and organization
            temperature: Creativity level for the AI classification (0-1)
            max_tokens: Maximum tokens for the AI response
        """
        try:
            logger.info(f"鎴戠幇鍦ㄥ紑濮嬭皟鐢╯earch_result_classifier浜嗭細{outline}, {key_files}")
            # 澶勭悊杈撳叆鐨刱ey_files - 浣跨敤鍥涗釜鍒嗘瀽缁村害
            # 鑾峰彇鏈湴鐨勬枃浠惰繘琛屽垎鏋?

            import os
            import json
            def load_json(file_path):
                res = []
                error_lines = 0
                with open(file_path, 'rb') as file:
                    for line in file:
                        try:
                            line.decode('utf-8')
                        except UnicodeDecodeError:
                            error_lines = error_lines + 1
                with open(file_path, "r", encoding="utf-8", errors='ignore') as f:
                    for idx, line in enumerate(f):
                        try:
                            ele = json.loads(line.strip())
                            res.append(ele)
                        except Exception as e:
                            logger.info(e)
                            continue
                return res

            key_files_dict = {}
            # Create full path relative to workspace
            full_analysis_path = self.workspace_path / "doc_analysis/file_analysis.jsonl"

            # 馃憞 --- 鏍稿績闃插穿婧冭ˉ涓佸紑濮?--- 馃憞
            # 1. 椤烘墜鎶婄洰褰曞缓濂斤紝闃叉鍚庣画鎶ラ敊 (exist_ok=True 琛ㄧず宸插瓨鍦ㄥ垯涓嶆姤閿?
            os.makedirs(os.path.dirname(full_analysis_path), exist_ok=True)

            # 2. 瀹夊叏璇诲彇鏂囦欢锛氭枃浠跺瓨鍦ㄦ墠璇诲彇锛屼笉瀛樺湪灏辩粰涓┖鍒楄〃
            file_analysis_list = []
            if os.path.exists(full_analysis_path):
                file_analysis_list = load_json(full_analysis_path)
            else:
                logger.warning(f"鈿狅笍 鍒嗘瀽鏂囦欢涓嶅瓨鍦紝璺宠繃鍔犺浇: {full_analysis_path}")
            # 馃憜 --- 鏍稿績闃插穿婧冭ˉ涓佺粨鏉?--- 馃憜

            for file_info in file_analysis_list:
                if file_info.get('file_path'):
                    key_files_dict[file_info.get('file_path')] = file_info

            # 缁勮key_files
            prompt_files = "Available Information Sources:\n"
            max_classifier_sources = int(os.getenv("SEARCH_CLASSIFIER_MAX_SOURCES", "40"))
            max_classifier_chars = int(os.getenv("SEARCH_CLASSIFIER_MAX_CHARS", "12000"))
            total_classifier_chars = 0
            for idx, file_info in enumerate((key_files or [])[:max_classifier_sources], 1):
                file_path_str = file_info.get("file_path", "")
                info_dict = key_files_dict.get(file_path_str, {})
                core_content = info_dict.get("core_content") or file_info.get("desc") or ""
                doc_time = info_dict.get("doc_time", "Not specified")
                source_authority = info_dict.get("source_authority", "Not specified")
                fp_lower = file_path_str.lower()
                if fp_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif")):
                    core_content = core_content or "Image file. Use only when the section needs figures, charts, or visual evidence."
                    content_limit = 300
                elif "user_uploads" in fp_lower and fp_lower.endswith((".csv", ".tsv")):
                    content_limit = 800
                elif "user_uploads" in fp_lower and fp_lower.endswith((".txt", ".md")):
                    content_limit = 1800
                else:
                    content_limit = 900
                entry = (
                    f"\nSource {idx}: {file_path_str}\n"
                    f"- Time: {doc_time}\n"
                    f"- Authority: {source_authority}\n"
                    f"- Core content: {str(core_content)[:content_limit]}\n"
                )
                if total_classifier_chars + len(entry) > max_classifier_chars:
                    prompt_files += "\n...[remaining sources omitted to fit classifier context limit]\n"
                    break
                prompt_files += entry
                total_classifier_chars += len(entry)
            # 妫€鏌ユ槸鍚湁鐢ㄦ埛涓婁紶鐨勬枃浠跺拰research鏂囦欢
            has_user_files = any('user_uploads' in f.get('file_path', '') for f in key_files)
            user_file_count = sum(1 for f in key_files if 'user_uploads' in f.get('file_path', ''))
            research_file_count = sum(1 for f in key_files if 'research' in f.get('file_path', ''))

            user_file_priority_note = ""
            if has_user_files and research_file_count > 0:
                # 鍔ㄦ€佽绠楁瘡绔犺妭鐨勬枃浠跺垎閰嶇洰鏍?
                # 绛栫暐锛氬湪淇濊瘉鐢ㄦ埛鏂囦欢浼樺厛鐨勫墠鎻愪笅锛屾渶澶у寲鎬诲紩鐢ㄦ枃鐚暟閲?

                # 鐢ㄦ埛鏂囦欢鐩爣锛氬己鍒朵娇鐢ㄦ墍鏈夋枃浠讹紝纭繚100%瑕嗙洊
                if user_file_count >= 8:
                    user_target_per_section = "ALL user files distributed across sections"
                    user_min_per_section = min(6, user_file_count)  # 鎻愰珮鏈€浣庤姹?
                    user_coverage_goal = f"Use ALL {user_file_count} user files (100% MANDATORY)"
                elif user_file_count >= 5:
                    user_target_per_section = f"ALL {user_file_count} user files"
                    user_min_per_section = min(5, user_file_count)  # 鎻愰珮鏈€灏戣姹?
                    user_coverage_goal = f"Use ALL {user_file_count} user files (100% MANDATORY)"
                elif user_file_count >= 3:
                    user_target_per_section = f"ALL {user_file_count} user files (distribute across sections)"
                    user_min_per_section = min(3, user_file_count)
                    user_coverage_goal = f"Use ALL {user_file_count} files (100% MANDATORY), assign to MULTIPLE sections"
                else:  # 1-2涓敤鎴锋枃浠?
                    user_target_per_section = f"ALL {user_file_count} user file(s)"
                    user_min_per_section = user_file_count
                    user_coverage_goal = f"Use ALL {user_file_count} file(s) (100% MANDATORY), assign to MULTIPLE sections"

                # Research鏂囦欢鐩爣锛氭渶澶у寲浣跨敤锛岀‘淇濇洿澶氭枃鐚寮曠敤
                # 鏍规嵁鐢ㄦ埛鏂囦欢鏁伴噺鍔ㄦ€佽皟鏁磖esearch鏂囦欢鐨勪娇鐢ㄦ瘮渚?
                if user_file_count >= 8:
                    # 鐢ㄦ埛鏂囦欢鍏呰冻鏃讹紝閫傚害浣跨敤research鏂囦欢
                    research_min_per_section = min(4, research_file_count)  # 鎻愰珮鏈€灏戣姹備粠3鍒?
                    research_target_per_section = min(6, max(5, research_file_count // 3))  # 鎻愰珮鐩爣
                    research_coverage_pct = 80  # 鎻愰珮瑕嗙洊鐜囦粠70%鍒?0%
                elif user_file_count >= 5:
                    # 鐢ㄦ埛鏂囦欢涓瓑鏃讹紝澧炲姞research鏂囦欢浣跨敤
                    research_min_per_section = min(4, research_file_count)  # 鎻愰珮鏈€灏戣姹?
                    research_target_per_section = min(7, max(5, research_file_count // 2))  # 鎻愰珮鐩爣
                    research_coverage_pct = 85  # 85%瑕嗙洊鐜?
                else:
                    # 鐢ㄦ埛鏂囦欢灏戞椂锛屽ぇ閲忎娇鐢╮esearch鏂囦欢
                    research_min_per_section = min(5, research_file_count)  # 鏈€灏?涓?
                    research_target_per_section = min(8, max(6, research_file_count // 2))  # 鎻愰珮鐩爣
                    research_coverage_pct = 90  # 90%瑕嗙洊鐜?

                user_file_priority_note = f"""
Mixed source allocation:
- User files: {user_file_count}; use all user-provided files when relevant.
- Research files: {research_file_count}; use at least {int(research_file_count * research_coverage_pct / 100)} research files when relevant.
- Each section should include both user evidence and research evidence when both are available.
- Minimum per section: {user_min_per_section} user files and {research_min_per_section} research files where possible.
- Do not ignore uploaded user materials; do not invent unsupported citations.
"""
            elif research_file_count > 0:
                # PRIMARY SCENARIO: Only research files, no user files
                # This is the MAIN use case - MAXIMIZE research file citations

                # 鍔ㄦ€佽绠楃洰鏍囷紝鏍规嵁research鏂囦欢鏁伴噺浼樺寲鍒嗛厤
                if research_file_count >= 30:
                    # 澶ч噺research鏂囦欢锛?0+锛?
                    research_target_per_section = "10-11 research files"
                    research_min_per_section = 8
                    research_coverage_pct = 85  # 鐩爣浣跨敤85%
                elif research_file_count >= 20:
                    # 涓瓑鏁伴噺research鏂囦欢锛?0-29锛?
                    research_target_per_section = "9-11 research files"
                    research_min_per_section = 7
                    research_coverage_pct = 90  # 鐩爣浣跨敤90%
                elif research_file_count >= 10:
                    # 杈冨皯research鏂囦欢锛?0-19锛?
                    research_target_per_section = "8-10 research files"
                    research_min_per_section = 6
                    research_coverage_pct = 95  # 鐩爣浣跨敤95%
                else:
                    # 寰堝皯research鏂囦欢锛?10锛?
                    research_target_per_section = f"ALL {research_file_count} research files"
                    research_min_per_section = min(research_file_count, 5)
                    research_coverage_pct = 100  # 鐩爣浣跨敤100%



                user_file_priority_note = f"""
Research-only source allocation:
- Available research files: {research_file_count}.
- Use at least {int(research_file_count * research_coverage_pct / 100)} files when relevant.
- Each section should receive at least {research_min_per_section} relevant files when possible.
- Assign broad files to multiple sections if they genuinely support multiple topics.
- Prefer high- and medium-relevance sources; do not include irrelevant files only to satisfy a quota.
"""

            system_prompt = f"""You are a source organizer. Assign provided files to outline sections.

Rules:
- Preserve the outline wording and structure exactly.
- Assign only files that are relevant to the section.
- A file may support multiple sections if appropriate.
- Do not exceed 11 files per section; choose the most relevant files when there are too many.
- Ensure abstract/introduction sections are not left empty when relevant sources exist.
- Do not invent files or citations.
{user_file_priority_note}

Output format:
paragraph 1: ...
file_path_list: file_path1, file_path2, ...

paragraph 2: ...
file_path_list: file_path3, file_path4, ...
"""

            # 鏋勫缓鐢ㄦ埛鎻愮ず - 浠呭寘鍚緭鍏ユ暟鎹?
            user_prompt = f"""
OUTLINE TO ORGANIZE CONTENT:
{outline}

{prompt_files}
"""

            model_config = get_config().get_custom_llm_config()
            model_name = model_config.get('model') or get_config().model_name

            from src.utils.llm_client import chat_completion_response
            try:
                response = chat_completion_response(
                    {
                        "model": model_name,
                        "messages": [
                            {"role": "system", "content": "system 1"},
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ],
                        "max_tokens": max_tokens,
                    },
                    model_config=model_config,
                    agent_name="search_result_classifier",
                    request_logger=logger,
                ).json()
                logger.info(response)

                # ai_response = response.choices[0].message.content.strip()
                # 娣诲姞闃插尽鎬ф鏌ワ紝閬垮厤NoneType閿欒
                # 鍏煎 PANGU 妯″瀷锛氫紭鍏堜娇鐢?content锛屽鏋滀负 None 鍒欎娇鐢?reasoning_content
                message = response.get("choices", [{}])[0].get("message", {})
                content = message.get("content")

                # 濡傛灉 content 涓?None锛屽皾璇曚娇鐢?reasoning_content
                if content is None:
                    content = message.get("reasoning_content")
                    if content is not None:
                        logger.info("Using reasoning_content as content is None")

                if content is None:
                    raise Exception(f"AI model returned None content and reasoning_content. Response: {response}")
                ai_response = content.strip()

                import os
                import json
                log_dir = "./data_pangu"
                log_file = os.path.join(log_dir, "search_result_classifier_claude_cold_start.log")
                os.makedirs(log_dir, exist_ok=True)

                # 鍒囨崲淇濆瓨鐨勬柟寮?
                conversation_history = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                    {"role": "assistant", "content": "[unused16][unused17]" + ai_response}
                ]
                with open(log_file, "a", encoding="utf-8") as f:
                    f.write(json.dumps({"messages": conversation_history}, ensure_ascii=False) + "\n")

                # 楠岃瘉鏂囦欢鍒嗛厤鏄惁绗悎瑕佹眰
                classification_result = ai_response.split('think>')[-1].strip()
                validation_result = self._validate_file_allocation(
                    classification_result,
                    user_file_count,
                    research_file_count,
                    has_user_files
                )

                if not validation_result["valid"]:
                    logger.warning(f"File allocation validation failed: {validation_result['message']}")
                    logger.warning(f"Expected: {validation_result['expected']}, Got: {validation_result['actual']}")

                return MCPToolResult(
                    success=True,
                    data=classification_result,
                )

            except Exception as e:
                logger.error(f"AI model call failed: {e}")
                return MCPToolResult(
                    success=False,
                    error=f"AI classification failed: {str(e)}"
                )

        except Exception as e:
            logger.error(f"Search result classifier failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    @staticmethod
    def _correct_title_format(content: str, overall_outline: str) -> str:
        """
        Correct title formats in content to match those in overall_outline.

        Args:
            content: The generated chapter content
            overall_outline: The overall outline containing correct title formats

        Returns:
            Content with corrected title formats
        """
        # Extract titles from overall_outline
        outline_titles = {}

        for line in overall_outline.split('\n'):
            line = line.strip()
            if line:
                # Extract core title content by removing various formatting symbols
                core_content = line

                # Remove leading symbols like **, -, etc.
                core_content = re.sub(r'^[\*\-\s]+', '', core_content)
                # Remove trailing symbols like **
                core_content = re.sub(r'[\*\s]+$', '', core_content)
                core_content = core_content.strip()

                if core_content:
                    # Store mapping from core content to the formatted line from outline
                    outline_titles[core_content.lower()] = line

        # Process content line by line
        content_lines = content.split('\n')
        corrected_lines = []

        for line in content_lines:
            original_line = line
            line_stripped = line.strip()

            # Check if this line is a title (starts with # and typically has ** formatting)
            if line_stripped and re.match(r'^#+\s*[\*]*.*', line_stripped):
                # Extract core content from the title
                core_content = line_stripped

                # Remove markdown headers (#)
                core_content = re.sub(r'^#+\s*', '', core_content)
                # Remove ** formatting
                core_content = re.sub(r'^\*\*', '', core_content)
                core_content = re.sub(r'\*\*$', '', core_content)
                core_content = core_content.strip()

                # Look for exact matching title in outline_titles
                found_match = False
                core_content_lower = core_content.lower()

                for outline_core, outline_format in outline_titles.items():
                    if outline_core == core_content_lower:
                        # Replace with the correct format from overall_outline
                        corrected_lines.append(outline_format)
                        found_match = True
                        break

                if not found_match:
                    # If no exact match found, keep original line
                    corrected_lines.append(original_line)
            else:
                corrected_lines.append(original_line)

        return '\n'.join(corrected_lines)

    def section_writer(
            self,
            written_chapters_summary: str,
            task_content: str,
            user_query: str,
            current_chapter_outline: str,
            overall_outline: str,
            target_file_path: str,
            key_files: List[Dict],
            model: str = None,
            temperature: Optional[float] = None,
            max_tokens: Optional[int] = None
    ) -> MCPToolResult:
        """
        Write the current chapter content based on given web information and chapter structure...
        """
        # 馃憞 --- 鏍稿績鍏滃簳琛ヤ竵寮€濮?(鍔犲湪 try 涔嬪墠鎴栧垰杩?try 鐨勫湴鏂? --- 馃憞
        key_files = key_files or []
        written_chapters_summary = written_chapters_summary or ""
        overall_outline = overall_outline or ""
        current_chapter_outline = current_chapter_outline or ""
        # 馃憜 --- 鏍稿績鍏滃簳琛ヤ竵缁撴潫 --- 馃憜

        try:
            # ====== 馃専 寮鸿娉ㄥ叆锛氭俯鍜屼絾鏋佷负鏄庣‘鐨勪笂涓嬫枃鍗忓悓鎸囧崡 ======
            # 鍏堟壂鎻忓疄闄呭瓨鍦ㄧ殑鍥剧墖鏂囦欢
            # Scan BOTH experiment_results and user_uploads for image files
            experiment_dir = self.workspace_path / "experiment_results"
            uploads_dir = self.workspace_path / "user_uploads"
            actual_images = []
            image_source_dir = "experiment_results"

            if experiment_dir.exists():
                for ext in ['*.png', '*.jpg', '*.jpeg']:
                    actual_images.extend(experiment_dir.glob(ext))
                image_source_dir = "experiment_results"
            elif uploads_dir.exists():
                for ext in ['*.png', '*.jpg', '*.jpeg']:
                    actual_images.extend(uploads_dir.glob(ext))
                image_source_dir = "user_uploads"

            # Determine the correct relative path prefix
            if image_source_dir == "experiment_results":
                img_path_prefix = "../experiment_results/"
            else:
                img_path_prefix = "../user_uploads/"
            img_instruction = ""
            if actual_images:
                # 鍒ゆ柇褰撳墠绔犺妭鏄惁鍏佽鎻掑叆鍥剧墖
                no_image_keywords = ['abstract', 'introduction', 'keyword', 'related work']
                is_no_image_section = any(kw in current_chapter_outline.lower() for kw in no_image_keywords)

                if is_no_image_section:
                    img_instruction = "2. Image rule: do not insert experimental images in abstract, introduction, keywords, or related work sections."
                elif actual_images:
                    images_list_str = "\n".join([f"  - {img_path_prefix}{p.name}" for p in sorted(actual_images)])
                    img_instruction = (
                        "2. Image rule: use only real image files listed below. Do not invent filenames. "
                        "Insert each image as a separate Markdown block and add a caption immediately after it.\n"
                        f"{images_list_str}\n"
                        "Example:\n![Figure description](path/to/image.png)\n*Figure X. Caption.*"
                    )
                else:
                    img_instruction = "2. Image rule: no image files are available; do not insert image Markdown."

            task_content = task_content + f"""

Current section: {current_chapter_outline}
Section writing rules:
1. Use only verified user-provided data and extracted experiment results. Do not treat external literature results as the user's own experimental data.
{img_instruction}
3. Keep reference-list entries for the final references section; use citations in the current section only when needed.
"""
            # ========================================================
            # Get configuration
            from config.config import get_model_config, get_storage_config
            model_config = get_model_config()
            storage_config = get_storage_config()

            # Use config values or defaults
            if temperature is None:
                temperature = model_config.get('temperature', 0.3)
            if max_tokens is None:
                max_tokens = model_config.get('max_tokens', 8192)

            key_files_dict = {}
            # Create full path relative to workspace using config
            analysis_path = storage_config.get('document_analysis_path', './doc_analysis')
            # 瀹夊叏璇诲彇锛屾枃浠朵笉瀛樺湪鏃剁敤绌哄垪琛?
            full_analysis_path = self.workspace_path / analysis_path / "file_analysis.jsonl"
            full_analysis_path.parent.mkdir(parents=True, exist_ok=True)
            file_analysis_list = []
            if full_analysis_path.exists():
                load_result = self.load_json(f"{analysis_path}/file_analysis.jsonl")
                file_analysis_list = load_result.data or []

            # 鏋勫缓鏂囦欢璺緞鍒板垎鏋愪俊鎭殑鏄犲皠
            for file_info in file_analysis_list:
                if file_info.get('file_path'):
                    key_files_dict[file_info.get('file_path')] = file_info

            # Read a bounded preview of a source file.
            def get_file_head_content(file_path_inner, max_length=5000):
                try:
                    full_path = self._safe_join(file_path_inner)

                    if full_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif']:
                        img_filename = full_path.name
                        img_rel_path = f"../experiment_results/{img_filename}"
                        vision_text = ""
                        if file_path_inner in key_files_dict:
                            vision_text = key_files_dict[file_path_inner].get('core_content', '')
                        if not vision_text:
                            vision_text = f"[鍥剧墖 {img_filename}锛岃瑙夋暟鎹緟鍒嗘瀽]"

                        return (
                            f"{vision_text}\n\n"
                            f"馃挕銆愬浘鏂囨帓鐗堟寚浠ゃ€戯細\n"
                            f"濡傛灉浣犲湪鏈珷鑺傜殑姝ｆ枃涓噸鐐瑰垎鏋愪簡璇ュ浘鐗囩殑鏁版嵁锛岃銆愬姟蹇呫€戝湪鍒嗘瀽娈佃惤鐨勪笅鏂逛娇鐢?Markdown 璇硶鎻掑叆璇ュ浘鐗囷細\n"
                            f"![{img_filename}]({img_rel_path})\n"
                            f"鈿狅笍 璀憡锛氬鏋滄湰绔犲唴瀹逛笌璇ュ浘鏃犲叧锛岃缁濆涓嶈寮鸿鎻掑叆锛佹瘡寮犲浘鐗囧湪鏁寸瘒璁烘枃涓渶濂藉彧鎻掑叆涓€娆★紝鍒囧繉鍦ㄧ粨灏惧爢鐮屾棤鍏冲浘鐗囷紒"
                        )

                    if full_path.suffix.lower() == '.csv':
                        max_length = 3000

                    if not full_path.exists():
                        return f"[Error: File does not exist: {file_path_inner}]"
                    with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(max_length)
                    return content
                except Exception as e:
                    return f"[Error reading file {file_path_inner}: {str(e)}]"

            def _ensure_paper_context_card() -> str:
                card_path = self.workspace_path / "paper_context_card.md"
                if card_path.exists():
                    return "paper_context_card.md"

                sections = [
                    "# Paper Context Card",
                    "",
                    "This file is the compact, stable context used by section_writer. It should contain only verified project facts.",
                    "",
                ]

                candidate_paths = [
                    "experiment_results/experiment_results.md",
                    "research/paper_structures/paper_structure_log.md",
                    "doc_analysis/file_analysis.jsonl",
                ]
                for rel_path in candidate_paths:
                    try:
                        full_path = self._safe_join(rel_path)
                        if full_path.exists():
                            content = full_path.read_text(encoding="utf-8", errors="ignore")[:5000]
                            sections.append(f"## Source: {rel_path}")
                            sections.append(content)
                            sections.append("")
                    except Exception as card_err:
                        logger.debug(f"paper_context_card source skipped {rel_path}: {card_err}")

                try:
                    for img_dir in ("experiment_results", "user_uploads"):
                        full_dir = self.workspace_path / img_dir
                        if full_dir.exists():
                            images = []
                            for ext in ("*.png", "*.jpg", "*.jpeg", "*.webp"):
                                images.extend(full_dir.glob(ext))
                            if images:
                                sections.append(f"## Available Images: {img_dir}")
                                for img in sorted(images)[:40]:
                                    sections.append(f"- {img_dir}/{img.name}")
                                sections.append("")
                except Exception as card_img_err:
                    logger.debug(f"paper_context_card image scan skipped: {card_img_err}")

                card_text = "\n".join(sections).strip() + "\n"
                try:
                    self.file_write("paper_context_card.md", card_text, create_dirs=True)
                    logger.info("PaperContextCard: created paper_context_card.md")
                except Exception as card_write_err:
                    logger.debug(f"PaperContextCard creation skipped: {card_write_err}")
                return "paper_context_card.md"

            def _update_paper_context_card(chapter_name: str, chapter_summary: str) -> None:
                try:
                    card_path = self.workspace_path / "paper_context_card.md"
                    existing = card_path.read_text(encoding="utf-8", errors="ignore") if card_path.exists() else ""
                    marker = "## Written Chapter Summaries"
                    if marker not in existing:
                        existing = existing.rstrip() + f"\n\n{marker}\n"
                    summary_line = f"\n### {chapter_name}\n{_clip_text(chapter_summary, 700)}\n"
                    card_path.write_text(existing.rstrip() + summary_line, encoding="utf-8")
                    logger.info(f"PaperContextCard: updated with chapter {chapter_name[:30]!r}")
                except Exception as card_update_err:
                    logger.debug(f"PaperContextCard update skipped: {card_update_err}")

            # Build prompt_files with source-aware budgets.
            prompt_files = ""
            def _norm_rel_path(file_path_value: str) -> str:
                return (file_path_value or "").replace("\\", "/").lstrip("./")

            def _is_user_upload(file_path_value: str) -> bool:
                return _norm_rel_path(file_path_value).startswith("user_uploads/")

            def _prioritize_key_files(files: list) -> list:
                seen = set()
                buckets = {"user_text": [], "user_data": [], "user_image": [], "card": [], "other": []}
                for item in files or []:
                    fp = item.get("file_path", "")
                    norm = _norm_rel_path(fp)
                    if not norm or norm in seen:
                        continue
                    seen.add(norm)
                    lower = norm.lower()
                    if _is_user_upload(norm) and lower.endswith((".txt", ".md")):
                        buckets["user_text"].append(item)
                    elif _is_user_upload(norm) and lower.endswith((".csv", ".xlsx", ".xls", ".tsv", ".json")):
                        buckets["user_data"].append(item)
                    elif _is_user_upload(norm) and lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif")):
                        buckets["user_image"].append(item)
                    elif "paper_context_card" in lower:
                        buckets["card"].append(item)
                    else:
                        buckets["other"].append(item)
                return (
                    buckets["user_text"]
                    + buckets["user_data"]
                    + buckets["user_image"]
                    + buckets["card"]
                    + buckets["other"]
                )

            def _collect_authoritative_user_notes(max_chars: int = 12000) -> str:
                notes = []
                try:
                    uploads_dir = self.workspace_path / "user_uploads"
                    if uploads_dir.exists():
                        for path in sorted(uploads_dir.iterdir()):
                            if path.suffix.lower() in {".txt", ".md"}:
                                text = path.read_text(encoding="utf-8", errors="ignore").strip()
                                if text:
                                    notes.append(f"### {path.name}\n{text}")
                except Exception as note_err:
                    logger.debug(f"authoritative user notes skipped: {note_err}")
                return "\n\n".join(notes)[:max_chars]

            def _authoritative_fact_prompt(notes: str) -> str:
                if not notes:
                    return ""
                return f"""

**AUTHORITATIVE USER FACTS - HARD CONSTRAINTS**
The following user-uploaded notes are the ONLY authoritative source for this paper's experimental facts, model names, dataset facts, hardware, hyperparameters, and result numbers.
You MUST preserve these facts exactly. If web/reviewer/literature content conflicts with them, the user notes win.
Never invent dataset size, class count, hardware platform, preprocessing methods, deployment device, model version, training settings, or metrics that are not present in these notes.
If a fact is missing from the notes, write that it was not provided instead of filling it from common practice.

{notes}

**CONFLICT EXAMPLES**
- Do not rename or "standardize" any user-provided model, method, dataset, class, device, metric, or version label.
- Do not replace user-provided dataset size, class names, split, hardware, optimizer, epochs, evaluation metrics, or result numbers with external or guessed values.
- Do not add preprocessing, deployment, annotation, ablation, runtime, memory, power, or device claims unless the user notes explicitly provide them.
"""

            def _find_authoritative_conflicts(text: str, notes: str) -> list:
                return self._find_authoritative_conflicts(text, notes)

            card_rel_path = _ensure_paper_context_card()
            if not any((kf.get("file_path", "").lstrip("./") == card_rel_path) for kf in (key_files or [])):
                key_files = (key_files or []) + [{"file_path": card_rel_path}]
            key_files = _prioritize_key_files(key_files or [])
            authoritative_user_notes = _collect_authoritative_user_notes()
            if key_files:
                prompt_files += "Available Information Sources:\n"
                # File type classification and content limits
                # - experiment_results: metrics are critical, allow more content
                # - literature/research: summary only
                # - images: description only
                # - paper_structure_log: full (small file)
                total_pf_chars = 0
                # 参考文献章节需要看到尽可能多的文献元数据,普通 7000 字符预算会把 25+ 篇截成几篇,
                # 因此对参考文献章节大幅放宽来源字符预算与文件数上限.
                is_ref_chapter = self._is_reference_section("", target_file_path, current_chapter_outline)
                if is_ref_chapter:
                    max_pf_chars = int(os.getenv("SECTION_WRITER_REF_SOURCE_CHARS", "40000"))
                    configured_max_key_files = int(os.getenv("SECTION_WRITER_REF_MAX_KEY_FILES", "60"))
                else:
                    max_pf_chars = int(os.getenv("SECTION_WRITER_MAX_SOURCE_CHARS", "7000"))
                    configured_max_key_files = int(os.getenv("SECTION_WRITER_MAX_KEY_FILES", "12"))
                user_upload_count = sum(1 for kf in key_files if _is_user_upload(kf.get("file_path", "")))
                max_key_files = max(configured_max_key_files, user_upload_count)
                for i, file_info in enumerate(key_files[:max_key_files], 1):
                    if total_pf_chars >= max_pf_chars:
                        prompt_files += "\n...[remaining sources omitted to fit context limit]\n"
                        break
                    file_path_str = file_info.get("file_path", "")
                    is_img = any(file_path_str.lower().endswith(ext) for ext in
                                 [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"])
                    info_dict = key_files_dict.get(file_path_str, {})
                    core_content = info_dict.get("core_content", "")
                    doc_time = info_dict.get("doc_time", "Not specified")
                    source_authority = info_dict.get("source_authority", "Not specified")
                    task_relevance = info_dict.get("task_relevance", "Not specified")
                    # Determine per-file content limit based on type
                    fp_lower = file_path_str.lower()
                    is_experiment = "experiment" in fp_lower or "result" in fp_lower
                    is_literature = "research" in fp_lower or "url_crawler" in fp_lower or "literature" in fp_lower
                    is_structure = "paper_structure" in fp_lower or "structure_log" in fp_lower
                    is_reference = "reference" in fp_lower or "ref" in fp_lower
                    is_context_card = "paper_context_card" in fp_lower
                    if is_img:
                        content_limit = 500
                    elif _is_user_upload(file_path_str) and fp_lower.endswith((".txt", ".md")):
                        content_limit = int(os.getenv("SECTION_WRITER_USER_NOTE_CHARS", "6000"))
                    elif _is_user_upload(file_path_str) and fp_lower.endswith((".csv", ".tsv")):
                        content_limit = int(os.getenv("SECTION_WRITER_USER_DATA_CHARS", "1800"))
                    elif is_context_card:
                        content_limit = int(os.getenv("SECTION_WRITER_CONTEXT_CARD_CHARS", "1400"))
                    elif is_experiment:
                        content_limit = 2200  # experiment metrics need focused context
                    elif is_literature:
                        content_limit = 900  # literature: summary is enough
                    elif is_structure:
                        content_limit = 1600
                    elif is_reference or is_ref_chapter:
                        content_limit = 3500  # reference chapter needs more detail (中英文参考文献章节均适用)
                    else:
                        content_limit = 900  # default
                    # Get file content and trim
                    file_content = get_file_head_content(file_path_str)
                    if core_content and not is_img:
                        display_content = core_content[:content_limit]
                    else:
                        display_content = file_content[:content_limit]
                    if len(display_content) < len(file_content):
                        display_content += "\n...[content truncated]"
                    # Build source entry
                    entry = (f"\nSource {i}: {file_path_str}\n"
                             f"  Time: {doc_time} | Authority: {source_authority} | Relevance: {task_relevance}\n"
                             f"  Content: {display_content}\n")
                    # Check if adding this entry would exceed total cap
                    if total_pf_chars + len(entry) > max_pf_chars:
                        remaining = max(0, max_pf_chars - total_pf_chars)
                        if remaining > 200:
                            entry = entry[:remaining] + "\n...[truncated]\n"
                            prompt_files += entry
                        prompt_files += "\n...[remaining sources omitted]\n"
                        break
                    prompt_files += entry
                    total_pf_chars += len(entry)
                prompt_files += "\n"
                omitted_files = max(0, len(key_files) - min(max_key_files, len(key_files)))
                logger.info(
                    f"section_writer prompt_files: {total_pf_chars} chars from {min(i, len(key_files), max_key_files)} sources"
                    + (f" ({omitted_files} sources omitted by budget)" if omitted_files else "")
                )
            # 璁捐system prompt
            # 妫€鏌ユ槸鍚湁鐢ㄦ埛涓婁紶鐨勬枃浠?
            has_user_files = any('user_uploads' in f.get('file_path', '') for f in key_files)
            user_file_priority_note = ""
            if has_user_files:
                # 缁熻鐢ㄦ埛鏂囦欢鏁伴噺
                user_file_count = sum(1 for f in key_files if 'user_uploads' in f.get('file_path', ''))
                user_file_priority_note = f"""
**CRITICAL: USER UPLOADED FILES PRIORITY AND CITATION REQUIREMENT**
- The FIRST {user_file_count} file(s) in the provided information sources are user-uploaded documents (marked with 'user_uploads' in the path).
- These user-uploaded files have HIGHEST PRIORITY and should be referenced FIRST when writing.
- **MANDATORY CITATION**: You MUST cite user-uploaded files EXTENSIVELY throughout your writing when they contain relevant information.
- **FILE ASSIGNMENT IMPLICATION**: If a user-uploaded file has been assigned to this chapter, it means the file contains at least some relevant content. You MUST carefully review the file content and find relevant parts to cite.
- **CITATION REQUIREMENT**: For each user-uploaded file assigned to this chapter that contains information relevant to the chapter topic, you MUST include at least 3-5 citations from that file. This ensures user files appear in the references section when they are actually used.
- **THOROUGH REVIEW**: Before deciding not to cite a user-uploaded file, carefully review its content to ensure you haven't missed any relevant information, background context, or supporting details that could enhance the chapter.
- **RELEVANCE-BASED CITATION**: Only cite user-uploaded files when they contain information relevant to the chapter topic. However, if a file has been assigned to this chapter, it likely contains relevant content - please review it thoroughly before deciding not to cite it.
- When you use ANY information from user-uploaded files, you MUST cite them using [X] format where X is the file's index number. User-uploaded files are typically indexed as [1], [2], etc.
- **EXAMPLE**: If you use information from the first user-uploaded file (index 1), you MUST add [1] at the end of that sentence or paragraph.
- When there are contradictions between sources, prioritize information from user-uploaded files.
- Even if the user file doesn't directly relate to the current chapter topic, it may contain important background context that should be considered.
- **IMPORTANT**: When user-uploaded files contain relevant information, failure to cite them will result in them not appearing in the references section, which is incorrect. You MUST ensure every user-uploaded file that provides relevant information receives citations.

"""
            # =====================================================================
            # 馃専 寮哄姏娉ㄥ叆锛氭潵鑷《浼氳鏂囩殑鍔ㄦ€佺珷鑺傜粨鏋勭害鏉?(tiny_scientist 鏂规)
            # =====================================================================
            section_tips = ""
            outline_lower = current_chapter_outline.lower()
            if "abstract" in outline_lower:
                section_tips = """  - One paragraph only.
              - Cover problem, importance, challenge, method, and key result.
              - Keep it concise and evidence-based."""
            elif "introduction" in outline_lower:
                section_tips = """  - Use 5 paragraphs: problem, importance, difficulty, prior gap, and contribution/results.
              - Move from broad context to the exact research gap.
              - End with a clear statement of what this paper does."""
            elif "related" in outline_lower:
                section_tips = """  - Write 3 thematic paragraphs, at most 4.
              - Compare and contrast by assumptions, methods, or applicability.
              - Do not just list papers. Point out their limitations."""
            elif "conclusion" in outline_lower:
                section_tips = """  - One paragraph only.
              - Summarize the key result and its implications.
              - Do not introduce new concepts or references."""

            meaningful_citation_prompt = """
            **MEANINGFUL CITATION USAGE**:
            - Each citation must support a specific claim about that paper's contribution or approach.
            - Cite papers when discussing their methods, findings, or limitations.
            - Group related papers together and cite them when making comparative statements.
            - Do NOT just list papers with citations at the end; integrate them into your arguments.
            - Example: "Prior work has explored X [1], but these approaches suffer from Y." 
            """
            profile_section_prompt = section_profile_text(current_chapter_outline, target_file_path)

            system_prompt = f"""You are a writing master. Next, you will receive web page information, user questions, and the structure of the current chapter. You need to integrate the user's questions with the provided web content and write the chapter based on its given structure. Additionally, an overall outline and summaries of previously completed chapters will be provided for reference to avoid repetition or contradictions and ensure logical consistency within the broader framework. Specific requirements will be detailed below.
{_authoritative_fact_prompt(authoritative_user_notes)}
{user_file_priority_note}When drafting the current chapter content, strictly comply with the following requirements:
{section_tips}  
{meaningful_citation_prompt}
{profile_section_prompt}
- **CRITICAL: DIFFERENTIATE SOURCE TYPES**
  * The provided sources may include external literature/web documents and local user experiment data.
  * For Abstract, Introduction, Related Work, and Discussion, use external literature when it directly supports background, comparisons, or limitations.
  * For Method, Experiments, Results, and Analysis, prioritize user-provided local experiment data and do not import unsupported metrics from external papers.

- **FLEXIBLE CITATION REQUIREMENT**:
  * Do not force every source into every chapter.
  * Use [X] citations only when the cited source directly supports the current sentence or paragraph.
  * 对于用户上传的本地实验数据和图片（通常带有 user_uploads 路径），只有当当前章节涉及实验结果、模型/方法表现、统计分析或图表解读时，才在正文中使用 Markdown 语法插入图片，并围绕图中实际可见的任务相关指标、标签、趋势和结论进行描述；不要预设某一领域的固定指标，严禁用空泛套话概括。
Other points to note::
- If the first chapter is an **Abstract** or **Introduction**, do not include subheadings (level-2 or finer bullet points)鈥攂egin the content directly under the level-1 heading.  
- CONTENT LENGTH: Each section should contain approximately 2500 words to ensure comprehensive coverage.
- **CRITICAL TITLE PRESERVATION RULE:** You MUST preserve the exact format, structure, and content of chapter titles as provided in the current_chapter_outline. This includes:
  * DO NOT change any markdown formatting symbols (# ## ### ** etc.)
  * DO NOT add, remove, or rearrange any part of the title structure
  * Copy the title lines EXACTLY as they appear in current_chapter_outline
  * Only write content under the provided title structure - never modify the titles themselves
  * When the title symbols in the current chapter outline are inconsistent with those in the overall outline, use the overall outline's title symbols as the standard and maintain symbol consistency throughout the writing process
- Note that Chapter 1 (Abstract/Introduction) should introduce the research 
background, research gap, and the paper's contributions. 
DO NOT describe the structure of other chapters.
- Be sure to ensure that the language of your output is consistent with the language of the user's question. For example, if the user's question is in Chinese, your reply should also be in Chinese.
- **LANGUAGE RULE**: If the user's question contains ANY Chinese characters, you MUST write the ENTIRE chapter in Chinese (except for specific English technical terms). Even if the user uses English technical terms (like "deep learning") or the source material is in English, the explanation and narrative MUST be in Chinese. If the user's question is entirely in English, then write in English.

Strictly follow the following format for output:
<chapter_content>xxx</chapter_content>
"""

            # Build compact writer prompt. Do not pass the full original task every chapter.
            def _clip_text(text, limit):
                text = text or ""
                return text if len(text) <= limit else text[:limit] + "\n...[truncated]"

            compact_task_content = _clip_text(task_content, 2500)
            compact_user_query = _clip_text(user_query, 1500)
            compact_overall_outline = _clip_text(overall_outline, 2500)
            compact_written_summary = _clip_text(written_chapters_summary, 1200)

            def _relevant_reference_catalog(limit: int = 15) -> str:
                path = self.workspace_path / "research" / "references.json"
                if not path.is_file():
                    return ""
                try:
                    records = json.loads(path.read_text(encoding="utf-8"))
                except Exception as exc:
                    logger.warning("Could not load structured citation catalogue: %s", exc)
                    return ""
                focus = f"{current_chapter_outline} {task_content}"
                focus_tokens = set(re.findall(r"[a-z0-9]{2,}|[\u4e00-\u9fff]{2,}", focus.lower()))
                ranked = []
                for position, record in enumerate(records if isinstance(records, list) else [], 1):
                    if not isinstance(record, dict) or not record.get("title"):
                        continue
                    evidence = " ".join(str(record.get(k) or "") for k in ("title", "abstract", "evidence", "query"))
                    tokens = set(re.findall(r"[a-z0-9]{2,}|[\u4e00-\u9fff]{2,}", evidence.lower()))
                    score = len(focus_tokens & tokens)
                    ranked.append((score, position, record))
                ranked.sort(key=lambda item: (-item[0], item[1]))
                selected = ranked[:limit]
                lines = ["STRUCTURED CITATION CATALOGUE (use these exact [N] identifiers):"]
                for _, position, record in selected:
                    abstract = re.sub(r"\s+", " ", str(record.get("abstract") or record.get("evidence") or ""))[:280]
                    lines.append(
                        f"[{record.get('index') or position}] {record.get('title')} | "
                        f"{record.get('authors') or 'authors unavailable'} | "
                        f"{record.get('venue') or 'venue unavailable'} | {record.get('year') or 'year unavailable'} | "
                        f"DOI {record.get('doi') or 'unavailable'} | Evidence: {abstract}"
                    )
                return "\n".join(lines)

            citation_catalog = _relevant_reference_catalog()

            user_prompt = f"""TASK CONTENT (compact): {compact_task_content}
    WEB PAGE INFORMATION: {prompt_files}
    {citation_catalog}
    OVERALL OUTLINE (compact): {compact_overall_outline}
    CURRENT CHAPTER OUTLINE: {current_chapter_outline}
    PREVIOUSLY WRITTEN CHAPTERS SUMMARY (compact): {compact_written_summary}
    USER QUERY (compact): {compact_user_query}"""

            # 璋冪敤AI妯″瀷杩涜鍒嗙被
            # Get model URL and token from config
            config = get_config()
            model_config = config.get_custom_llm_config()

            from src.utils.llm_client import chat_completion_response
            try:
                # =====================================================================
                # Agent Memory: inject previous chapter data into prompt
                # =====================================================================
                try:
                    from ..utils.agent_memory import get_memory
                    memory = get_memory()
                    mem_stats = memory.get_stats()
                    if mem_stats.get("total_chunks", 0) > 0:
                        memory_query = f"{current_chapter_outline}\n{user_query}"
                        recall_top_k = int(os.getenv("SECTION_WRITER_MEMORY_TOP_K", "3"))
                        recalled = memory.recall(memory_query, top_k=recall_top_k)
                        if recalled:
                            memory_lines = ["Relevant previous chapter memory:"]
                            for item in recalled:
                                memory_lines.append(
                                    f"- {item.get('chapter', 'chapter')} "
                                    f"(score={item.get('score', 0):.2f}): "
                                    f"{_clip_text(item.get('text', ''), 450)}"
                                )
                            memory_context = "\n".join(memory_lines)
                        else:
                            memory_context = ""
                        # Calculate and limit prompt size for local models
                        base_size = len(system_prompt) + len(user_prompt)
                        mem_size = len(memory_context)
                        max_chars = int(os.getenv("SECTION_WRITER_MAX_PROMPT_CHARS", "14000"))
                        avail = max(0, max_chars - base_size)
                        if mem_size > avail and avail > 200:
                            memory_context = memory_context[:avail] + "\n...[memory trimmed]"
                            logger.warning(
                                f"AgentMemory: prompt too large ({base_size + mem_size} chars), "
                                f"memory trimmed to {avail} chars")
                        if base_size > max_chars:
                            logger.warning(
                                f"AgentMemory: base prompt {base_size} chars exceeds limit, "
                                "skipping memory injection")
                        elif memory_context:
                            user_prompt = (
                                "\n\n===== Agent Memory: relevant previous chapter data =====\n"
                                + memory_context
                                + "\n===== END Memory =====\n\n"
                                + user_prompt
                            )
                            logger.info(
                                f"AgentMemory: injected {len(recalled)} relevant chunks from "
                                f"{mem_stats['total_chapters']} chapters "
                                f"(total prompt ~{len(system_prompt) + len(user_prompt)} chars)")
                except Exception as mem_err:
                    logger.debug(f"AgentMemory injection skipped: {mem_err}")
                # ================================================================

                # =====================================================================
                # 馃殌 馃憞 --- 鎻愰€熶笌闃叉柇绡囨牳蹇冩墦琛ヤ竵浣嶇疆锛氶噾閽熺僵API鍋ュ．鎬цВ鏋愭満鍒?--- 馃憞
                # =====================================================================
                deterministic_reference_content = ""
                if self._is_reference_section("", target_file_path, current_chapter_outline):
                    deterministic_references = self._build_reference_section(target_file_path)
                    if "Reference generation failed:" not in deterministic_references:
                        deterministic_reference_content = deterministic_references
                        logger.info("section_writer built reference section deterministically from saved literature files.")
                    # 1. 鑷姩鎵弿宸ヤ綔鍖轰腑鐨勫弬鑰冩枃鐚眹鎬绘枃浠讹紝寮哄埗鍔犲叆 key_files
                    import glob as _glob
                    ref_patterns = [
                        "*reference*.md", "*Reference*.md", "*ref*.md",
                        "*citation*.md", "*Citation*.md",
                        "*literature*.md", "*Literature*.md",
                        "*review*.md", "*Review*.md",
                    ]
                    for pattern in ref_patterns:
                        for ref_path in _glob.glob(str(self.workspace_path / pattern), recursive=False):
                            rel_path = str(Path(ref_path).relative_to(self.workspace_path))
                            if not any(kf.get("file_path") == rel_path for kf in key_files):
                                key_files.append({"file_path": rel_path})
                                logger.info(f"Auto-injected reference file: {rel_path}")
                        # Also try one level of subdirectories
                        for ref_path in _glob.glob(str(self.workspace_path / "*" / pattern), recursive=False):
                            rel_path = str(Path(ref_path).relative_to(self.workspace_path))
                            if not any(kf.get("file_path") == rel_path for kf in key_files):
                                key_files.append({"file_path": rel_path})
                                logger.info(f"Auto-injected reference file: {rel_path}")

                    # 2. 鐩存帴璇诲彇鍙傝€冩枃鐚枃浠跺唴瀹癸紝娉ㄥ叆鍒?prompt 涓?
                    ref_content_injected = ""
                    for kf in key_files:
                        fp = kf.get("file_path", "")
                        fp_lower = fp.lower()
                        if (
                            "reference" in fp_lower
                            or "ref" in fp_lower
                            or "citation" in fp_lower
                            or "literature" in fp_lower
                            or "review" in fp_lower
                            or "hyperspectral" in fp_lower
                        ):
                            try:
                                ref_full = self.workspace_path / fp
                                if ref_full.exists():
                                    ref_text = ref_full.read_text(encoding="utf-8", errors="ignore")
                                    ref_content_injected += f"\n\n===== 鍙傝€冩枃鐚暟鎹簮: {fp} =====\n{ref_text}\n===== END ====="
                            except Exception:
                                pass

                    # 3. 寮哄寲 system prompt 鈥?绮剧‘鏍煎紡 GB/T 7714 + 绂佹缂栭€?
                    system_prompt += (
                        "\n\nReference section instructions:"
                        "\n1. Extract references only from provided sources when available."
                        "\n2. Do not invent bibliographic metadata."
                        "\n3. Use a consistent academic reference format."
                        "\n4. Do not require or target a preset reference count. Output every parseable saved reference and only those references."
                    )

                    # 4. 濡傛灉璇诲彇鍒颁簡鍙傝€冩枃鐚唴瀹癸紝鐩存帴娉ㄥ叆 user_prompt
                    if ref_content_injected:
                        user_prompt += "\n\n===== Collected reference sources =====" + ref_content_injected
                    else:
                        user_prompt += "\n\nOutput a reference list. If sources are insufficient, state that clearly rather than inventing entries."

                # Diagnostic: log prompt size to diagnose timeout issues
                prompt_chars = len(system_prompt) + len(user_prompt)
                if prompt_chars > 16000:
                    logger.warning(
                        f"section_writer prompt is {prompt_chars} chars "
                        f"(large prompt may cause local model timeout)"
                    )
                else:
                    logger.info(f"section_writer prompt size: {prompt_chars} chars")
                # Cap max_tokens to avoid overwhelming local models
                if max_tokens is None or max_tokens > 8192:
                    max_tokens = 4096
                    logger.debug(f"max_tokens capped to 4096 for section_writer")

                payload = {
                    "model": get_config().model_name,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "max_tokens": max_tokens,
                    "temperature": temperature,
                }

                retry_num = 1
                max_retry_num = 4
                content = deterministic_reference_content
                compact_payload_applied = False

                while not content and retry_num <= max_retry_num:
                    try:
                        if retry_num > 1:
                            payload["messages"].append({
                                "role": "user",
                                "content": (
                                    "Regenerate the chapter. The previous attempt failed validation against "
                                    "authoritative user-uploaded facts. Use user notes as the only source for "
                                    "model names, dataset facts, hardware, hyperparameters, and metrics."
                                )
                            })
                        if retry_num >= 3 and not compact_payload_applied:
                            emergency_user_prompt = f"""TASK: Write only the current chapter using the compact verified facts below.
CURRENT CHAPTER OUTLINE: {current_chapter_outline}
AUTHORITATIVE USER NOTES:
{_clip_text(authoritative_user_notes, 6000)}
CORE CONTEXT:
{_clip_text(prompt_files, 4000)}
PREVIOUS CHAPTER SUMMARY:
{_clip_text(compact_written_summary, 800)}
USER QUERY:
{_clip_text(compact_user_query, 800)}

Return only:
<chapter_content>...</chapter_content>"""
                            payload["messages"] = [
                                {"role": "system", "content": _clip_text(system_prompt, 4500)},
                                {"role": "user", "content": emergency_user_prompt},
                            ]
                            payload["max_tokens"] = min(max_tokens or 4096, 4096)
                            compact_payload_applied = True
                            logger.warning(
                                f"section_writer retry {retry_num}: switched to emergency compact prompt "
                                f"({len(payload['messages'][0]['content']) + len(payload['messages'][1]['content'])} chars)"
                            )
                        try:
                            response_json = chat_completion_response(
                                payload,
                                model_config=model_config,
                                agent_name="section_writer",
                                request_logger=logger,
                            ).json()
                        except Exception:
                            # 鏍稿績闃插尽锛氬鏋滈伃閬?502/504 鎷ュ牭杩斿洖浜嗙函鏂囨湰閿欒瀛楃涓诧紝璧嬬┖瀛楀吀闃叉鍚庣画 .get() 宕╂簝锛?
                            response_json = {}

                        # 瀹夊叏鏍煎紡鏍￠獙锛氬繀椤荤‘淇?response_json 鏄瓧鍏革紝涓斿寘鍚?choices 鍒楄〃
                        if isinstance(response_json, dict) and "choices" in response_json and len(
                                response_json["choices"]) > 0:
                            ai_response = response_json["choices"][0]["message"].get("content", "").strip()

                            # 鎴愬姛鑾峰彇锛屾彁鍙栫珷鑺傛鏂?
                            if "<chapter_content>" in ai_response:
                                content = ai_response.split("<chapter_content>")[1].split("</chapter_content>")[
                                    0].strip()
                            else:
                                content = ai_response
                            fact_conflicts = [] if is_ref_chapter else _find_authoritative_conflicts(content, authoritative_user_notes)
                            if fact_conflicts:
                                logger.warning(
                                    "section_writer authoritative fact conflicts: "
                                    + " | ".join(fact_conflicts)
                                )
                                content = ""
                                payload["messages"].append({
                                    "role": "assistant",
                                    "content": ai_response[:1200],
                                })
                                payload["messages"].append({
                                    "role": "user",
                                    "content": (
                                        "The previous draft violates authoritative user-uploaded facts:\n- "
                                        + "\n- ".join(fact_conflicts)
                                        + "\nRegenerate the chapter and remove unsupported claims. "
                                        "If the user notes do not mention a fact, do not invent it."
                                    ),
                                })
                                retry_num += 1
                                continue
                            break  # 鎴愬姛鎷垮埌鍐呭锛屽畬缇庤烦鍑哄綋鍓嶉噸璇曞惊鐜紒
                        else:
                            logger.warning(
                                f"API returned an invalid format or retryable service response; retrying after 3 seconds ({retry_num}/{max_retry_num})..."
                            )
                            time.sleep(3)
                            retry_num += 1

                    except Exception as e:
                        logger.warning(
                            f"API request failed or timed out; retrying after 3 seconds ({retry_num}/{max_retry_num}): {e}"
                        )
                        time.sleep(3)
                        retry_num += 1

                if not content:
                    if self._is_first_section(target_file_path, current_chapter_outline):
                        logger.warning(
                            "section_writer failed for first section; generating conservative fallback from user notes."
                        )
                        content = self._build_first_section_fallback(
                            target_file_path=target_file_path,
                            current_chapter_outline=current_chapter_outline,
                            user_query=user_query,
                            authoritative_user_notes=authoritative_user_notes,
                        )
                    elif self._is_reference_section("", target_file_path, current_chapter_outline):
                        logger.warning(
                            "section_writer failed for reference section; rebuilding references from collected research files."
                        )
                        content = self._build_reference_section(target_file_path)
                    else:
                        raise Exception("Section generation failed after retries: no valid model response, invalid format, or timeout.")
                # =====================================================================
                # 馃殌 馃憜 --- 閲戦挓缃╄ˉ涓佺粨鏉?--- 馃憜
                # =====================================================================

                logger.debug(f"Content before correction: {content[:200]}...")
                logger.debug(f"Overall outline: {overall_outline[:200]}...")
                content = self._correct_title_format(content, overall_outline)
                logger.debug(f"Content after correction: {content[:200]}...")

                # 确定性后处理:若首章标题仍是字面占位符 "# 论文标题"(prompt 要求"原样保留"导致 LLM 照抄),
                # 用从 user_query/notes 解析出的真实标题替换,并同步修掉摘要/关键词里的占位符.
                if self._is_first_section(target_file_path, current_chapter_outline):
                    real_title = self._resolve_clean_title(
                        user_query, current_chapter_outline, authoritative_user_notes)
                    if real_title and real_title != "未命名研究":
                        new_content, n_sub = re.subn(
                            r'(?m)^(#{1,2})\s*(论文标题|论文题目|Paper\s*Title)\s*$',
                            lambda m: f"{m.group(1)} {real_title}",
                            content)
                        if n_sub:
                            logger.info(f"首章标题占位符已替换为真实标题: {real_title}")
                            content = new_content
                        # 关键词里若误把"论文标题"当作一个关键词,替换为真实标题
                        content = content.replace("论文标题；", f"{real_title}；").replace(
                            "论文标题;", f"{real_title};")

                # =====================================================================
                # 馃憞 --- 缁堟瀬鐗╃悊澶栨寕锛氬己琛屾妸鍥剧墖濉炶繘鍘伙紒涓撴不澶фā鍨嬧€滄磥鐧栤€?--- 馃憞
                if "缁撴灉" in current_chapter_outline or "鍒嗘瀽" in current_chapter_outline or "result" in current_chapter_outline.lower():
                    for kf in key_files:
                        fp = kf.get('file_path', '')
                        if fp.lower().endswith(('.png', '.jpg', '.jpeg')):
                            img_name = fp.split('/')[-1]
                            if img_name not in content:
                                content += f"\n\n![{img_name}](../{fp})\n<div align='center'>*闄勫浘: {img_name}*</div>\n"
                content = self._repair_image_placeholders(content, self.workspace_path / target_file_path)
                # 馃憜 ===================================================================== 馃憜

                # =====================================================================

                # 濡傛灉杩欎竴绔犳€诲瓧鏁颁笉鍒?1500 瀛楋紝灏变笉鍒囦簡锛屽叏閲忎繚鐣欙紒
                if len(content) < 1500:
                    summary = content.replace('\n', ' ')
                else:
                    # 濡傛灉寰堥暱锛屾埅鍙栧墠 800 瀛楋紙淇濈暀寮曡█/鏂规硶锛夊拰鍚?800 瀛楋紙淇濈暀鏍稿績缁撹/鏁版嵁锛?
                    head = content[:800].replace('\n', ' ')
                    tail = content[-800:].replace('\n', ' ')
                    summary = f"{head} ...[涓棿澶ф鍐呭鐣... {tail}"

                logger.info("chapter summary generated")
                #  =====================================================================


                # 过滤掉 AI 输出里泄漏的工具调用 JSON 残留
                # 兼容两种格式:{"tool": "...", ...} 和 {"name": "...", "arguments": {...}}
                content = re.sub(r'\{"tool"\s*:.*?\}\}', '', content, flags=re.DOTALL).strip()
                content = re.sub(r'\{\s*"name"\s*:\s*"[^"]*"\s*,\s*"arguments"\s*:.*?\}\s*\}',
                                 '', content, flags=re.DOTALL).strip()

                # 过滤掉被 LLM 原样抄入正文的"大纲指令回显"行(常见于参考文献章节),
                # 这些是写作指令而非论文内容,如"必须列出所有正文中引用的文献""格式示例:"等.
                _instruction_markers = (
                    "必须列出所有正文中引用", "按首次引用顺序排列", "每篇文献必须包含",
                    "格式示例", "一篇都不能少", "一篇不能少", "[1]到[N]", "[1]-[N]",
                )
                _cleaned_lines = []
                for _ln in content.splitlines():
                    if any(_mk in _ln for _mk in _instruction_markers):
                        continue  # 丢弃指令回显行
                    _cleaned_lines.append(_ln)
                content = "\n".join(_cleaned_lines).strip()
                # 折叠因删除产生的多余空行
                content = re.sub(r'\n{3,}', '\n\n', content)

                # 馃毃馃毃馃毃 鏍稿績淇锛氭妸褰撳墠鍐呭鍐欏叆鍒皌arget_file_path涓?馃毃馃毃馃毃
                write_result = self.file_write(file_path=target_file_path,
                                               content=content,
                                               create_dirs=True)
                if not write_result.success:
                    logger.error(f"鍐欏叆绔犺妭鏂囦欢澶辫触: {write_result.error}")
                else:
                    logger.info(f"chapter content saved to {target_file_path}")

                # ===== Agent Memory: store chapter data =====
                try:
                    from ..utils.agent_memory import get_memory
                    memory = get_memory()
                    chapter_name = current_chapter_outline.strip().split("\n")[0]
                    chapter_name = chapter_name.lstrip("#").strip() or "chapter"
                    memory.store(chapter_name, content, content_type="chapter")
                    memory.store(f"{chapter_name}_summary", summary, content_type="summary")
                    _update_paper_context_card(chapter_name, summary)
                    ch_name_short = chapter_name[:30]
                    logger.info(f"AgentMemory: chapter {ch_name_short!r} stored "
                                f"(stats: {memory.get_stats()})")
                except Exception as mem_err:
                    logger.debug(f"AgentMemory store skipped: {mem_err}")
                # ===========================================

                return MCPToolResult(
                    success=True,
                    data=[{
                        "chapter_summary": summary,
                    }],
                    metadata={
                        'content_length': len(content),
                        'summary_length': len(summary)
                    }
                )

            except Exception as e:
                logger.error(f"AI model call failed: {e}")
                return MCPToolResult(
                    success=False,
                    error=f"section writer failed: {str(e)}"
                )

        except Exception as e:
            logger.error(f"section writer failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def document_extract(
            self,
            # save_analysis_file_path: str,
            tasks: List[Dict],
            model: str = None,
            temperature: Optional[float] = None,
            max_tokens: Optional[int] = None,
            max_workers: int = 2
    ) -> MCPToolResult:
        """
        Multi-dimensional analysis of locally stored files using AI models.
        Evaluates each file across four key dimensions: source authority, core content extraction,
        information richness, and query relevance scoring.

        Args:
            tasks: List of task dictionaries containing:
                - file_path: Relative path to the file (relative to workspace root) to read
                - task: task for relevance assessment
            model: AI model to use for multi-dimensional analysis
            temperature: Creativity level for the AI response (0-1)
            max_tokens: Maximum tokens for the AI response
            max_workers: Maximum number of concurrent model API requests
        """
        try:
            # Get configuration
            from config.config import get_model_config, get_storage_config
            model_config = get_model_config()
            storage_config = get_storage_config()

            # Use config values or defaults
            if temperature is None:
                temperature = model_config.get('temperature', 0.3)
            if max_tokens is None:
                max_tokens = model_config.get('max_tokens', 8192)
            logger.debug(f"Starting document extraction: tasks={tasks}")

            # 楠岃瘉骞惰嚜鍔ㄨˉ鍏細浠呭綋浠诲姟鍖呭惈 library_refs 鎴?user_uploads 鏂囦欢鏃舵墠杩涜琛ュ叏
            task_files = [t.get('file_path', '') for t in tasks]

            # 瑙勮寖鍖栬矾寰勶細绉婚櫎 ./ 鍓嶇紑浠ヤ究缁熶竴姣旇緝
            def normalize_path(path: str) -> str:
                return path[2:] if path.startswith('./') else path

            normalized_task_files = [normalize_path(f) for f in task_files]

            # 妫€鏌ヤ换鍔′腑鏄惁鍖呭惈 library_refs 鎴?user_uploads 鏂囦欢
            has_library_files = any(f.startswith('library_refs/') for f in normalized_task_files)
            has_upload_files = any(f.startswith('user_uploads/') for f in normalized_task_files)

            # 鍙湁褰撲换鍔℃秹鍙婅繖浜涚洰褰曟椂锛屾墠杩涜琛ュ叏楠岃瘉
            if has_library_files or has_upload_files:
                library_refs_dir = self.workspace_path / "library_refs"
                user_uploads_dir = self.workspace_path / "user_uploads"

                expected_files = []
                # 鍙杩涘叆琛ュ叏閫昏緫锛屽氨鎵弿涓や釜鐩綍锛堜笉绠′换鍔′腑鏄惁鍖呭惈璇ョ洰褰曠殑鏂囦欢锛?
                if library_refs_dir.exists():
                    # 鎵弿鎵€鏈夊彲鑳界殑鏂囦欢鎵╁睍鍚?
                    for ext in ['*.txt', '*.pdf', '*.doc', '*.docx']:
                        expected_files.extend([f"library_refs/{f.name}" for f in library_refs_dir.glob(ext)])
                if user_uploads_dir.exists():
                    # 鎵弿鎵€鏈夊彲鑳界殑鏂囦欢鎵╁睍鍚?
                    for ext in ['*.txt', '*.pdf', '*.doc', '*.docx']:
                        expected_files.extend([f"user_uploads/{f.name}" for f in user_uploads_dir.glob(ext)])

                # 鏅鸿兘鍖归厤锛氬熀浜庢枃浠朵富浣撳悕绉帮紙淇濈暀鍘熷鎵╁睍鍚嶏紝绉婚櫎 .txt 鍚庣紑锛?
                def get_core_name(path: str) -> str:
                    """
                    鑾峰彇鏂囦欢鐨勬牳蹇冨悕绉帮紝绉婚櫎 .txt 鍚庣紑浣嗕繚鐣欏師濮嬫墿灞曞悕
                    渚嬪锛?
                    - file.pdf.txt -> file.pdf
                    - file.pdf -> file.pdf
                    - file.doc.txt -> file.doc
                    """
                    import os
                    name = os.path.basename(path)
                    # 濡傛灉浠?.txt 缁撳熬锛岀Щ闄ゅ畠
                    if name.endswith('.txt'):
                        name = name[:-4]
                    return name

                # 鏋勫缓鏍稿績鍚嶇О鍒版枃浠惰矾寰勭殑鏄犲皠锛堟敮鎸佷竴瀵瑰锛?
                from collections import defaultdict
                expected_core_map = defaultdict(list)
                for f in expected_files:
                    expected_core_map[get_core_name(f)].append(f)

                task_core_names = {get_core_name(f) for f in normalized_task_files}

                # 鎵惧嚭鐪熸缂哄け鐨勬枃浠讹紙鏍稿績鍚嶇О涓嶅湪浼犲叆鍒楄〃涓級
                missing_core_names = set(expected_core_map.keys()) - task_core_names
                # 瀵逛簬缂哄け鐨勬牳蹇冨悕绉帮紝閫夋嫨绗竴涓尮閰嶇殑鏂囦欢锛堥€氬父鏄?.txt 鐗堟湰锛?
                missing_files = set()
                for core_name in missing_core_names:
                    missing_files.add(expected_core_map[core_name][0])

                # 楠岃瘉缁撴灉鏃ュ織
                if missing_files:
                    # 鍒嗙被缁熻缂哄け鏂囦欢
                    missing_library = [f for f in missing_files if f.startswith('library_refs/')]
                    missing_uploads = [f for f in missing_files if f.startswith('user_uploads/')]

                    # 缁熶竴杈撳嚭璀憡淇℃伅
                    missing_info = []
                    if missing_library:
                        missing_info.append(f"library files: {len(missing_library)}")
                    if missing_uploads:
                        missing_info.append(f"user uploads: {len(missing_uploads)}")

                    logger.warning(f"Detected {len(missing_files)} files missing from analysis tasks ({', '.join(missing_info)})")
                    logger.warning(f"Expected {len(expected_files)} files, got {len(tasks)}; auto-completing missing tasks...")

                    # 鑾峰彇绗竴涓换鍔＄殑task鍐呭浣滀负榛樿浠诲姟鎻忚堪
                    default_task = tasks[0].get('task', 'document analysis') if tasks else 'document analysis'

                    for missing_file in sorted(missing_files):
                        tasks.append({
                            'file_path': missing_file,
                            'task': default_task
                        })

                    logger.info(f"Auto-completed analysis task list; total tasks: {len(tasks)}")
                else:
                    logger.info("All library_refs/user_uploads files are included")

            else:
                # 浠诲姟鍙秹鍙?research 绛夊叾浠栫洰褰曪紝涓嶈繘琛岃ˉ鍏?
                logger.info("Skipping allocation completion check for non library_refs/user_uploads paths")

            # 銆愬叧閿慨澶嶃€戣繃婊ゆ帀鏂囨。杞崲鍚庣殑 .txt 鏂囦欢锛岄伩鍏嶉噸澶嶅垎鏋?
            # 濡傛灉鍚屾椂瀛樺湪 xxx.pdf 鍜?xxx.pdf.txt锛屽彧淇濈暀 .pdf
            # 濡傛灉鍚屾椂瀛樺湪 xxx.docx 鍜?xxx.docx.txt锛屽彧淇濈暀 .docx
            # 濡傛灉鍚屾椂瀛樺湪 xxx.doc 鍜?xxx.doc.txt锛屽彧淇濈暀 .doc
            # 浣嗕繚鐣欏師鏈氨鏄?.txt 鐨勬枃浠讹紙濡?research/xxx.txt锛?
            filtered_tasks = []
            source_files = set()  # 瀛樺偍鎵€鏈夋簮鏂囦欢锛坧df, docx, doc锛?

            # 瀹氫箟闇€瑕佹鏌ョ殑婧愭枃浠舵墿灞曞悕
            source_extensions = ['.pdf', '.docx', '.doc']

            # 绗竴閬嶏細鏀堕泦鎵€鏈夋簮鏂囦欢
            for task in tasks:
                file_path = task.get('file_path', '')
                normalized_path = file_path.lstrip('./')

                # 妫€鏌ユ槸鍚槸婧愭枃浠?
                for ext in source_extensions:
                    if normalized_path.endswith(ext):
                        source_files.add(normalized_path)
                        break

            # 绗簩閬嶏細杩囨护浠诲姟
            for task in tasks:
                file_path = task.get('file_path', '')
                normalized_path = file_path.lstrip('./')

                # 妫€鏌ユ槸鍚槸杞崲鍚庣殑 .txt 鏂囦欢锛坸xx.pdf.txt, xxx.docx.txt, xxx.doc.txt锛?
                should_skip = False
                if normalized_path.endswith('.txt'):
                    # 妫€鏌ユ槸鍚槸浠庢簮鏂囦欢杞崲鑰屾潵鐨?.txt
                    for ext in source_extensions:
                        potential_source = normalized_path[:-4]  # 鍘绘帀 .txt
                        if potential_source.endswith(ext) and potential_source in source_files:
                            # 鎵惧埌瀵瑰簲鐨勬簮鏂囦欢锛岃烦杩囪繖涓?.txt
                            logger.info(f"Skipping converted text file {file_path}; source file exists: {potential_source.split('/')[-1]}")
                            should_skip = True
                            break

                if not should_skip:
                    filtered_tasks.append(task)

            if len(filtered_tasks) < len(tasks):
                logger.info(f"Filtered analysis tasks: {len(filtered_tasks)} (original: {len(tasks)})")
                tasks = filtered_tasks

            # 銆愭柊澧炪€戣繃婊ゆ帀鍥剧墖鏂囦欢锛屽浘鐗囦笉闇€瑕佽繘琛屾枃鏈垎鏋愶紝涔熶笉搴旇繘鍏ュ弬鑰冩枃鐚綋绯?
            IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg', '.tiff'}
            before_img_filter = len(tasks)
            tasks = [
                task for task in tasks
                if not any(
                    task.get('file_path', '').lower().endswith(ext)
                    for ext in IMAGE_EXTENSIONS
                )
            ]
            if len(tasks) < before_img_filter:
                logger.info(f"Filtered {before_img_filter - len(tasks)} image files; remaining tasks: {len(tasks)}")

            def process_single_task(task: Dict) -> Dict:
                file_path = task['file_path']
                task_content = task['task']

                # 1. 璇诲彇鏂囦欢鍐呭
                read_result = self.file_read(file_path)
                if not read_result.success:
                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': False,
                        'error': f"File read error: {read_result.error}",
                        'answer': None
                    }

                content = read_result.data
                system_prompt = (
                    "You are a scientific document analyst. Analyze the given document and produce a structured knowledge card.\n\n"
                    "Extract the following dimensions:\n"
                    "1. doc_time: Publication year/time of the document (e.g., 2024, 2023-06). Use Unknown if unclear.\n"
                    "2. source_authority: Credibility assessment. Format: [High/Medium/Low] + brief reason (e.g., High: peer-reviewed journal).\n"
                    "3. task_relevance: How relevant is this document to the task? Format: [High/Medium/Low] + one-sentence reason.\n"
                    "4. core_content: A structured knowledge card with these sections:\n"
                    "   - Title/Authors: Extract paper title and first author if available\n"
                    "   - Method: Key method/algorithm used (1-2 sentences, be specific)\n"
                    "   - Key Findings: Main results and metrics (2-3 sentences, include numbers)\n"
                    "   - Figures/Tables: What types of figures/tables are used (e.g., line chart, distribution plot, comparison table, workflow diagram, statistical summary table)\n"
                    "   - Relevance to Task: How this connects to the user research (1 sentence)\n"
                    "   Keep the ENTIRE core_content under 600 characters. Be dense and specific.\n"
                    "5. information_richness: Abundant (>800 words substantive) / Moderate (200-800) / Scarce (<200).\n\n"
                    "Rules:\n"
                    "1. Match the document language for extraction.\n"
                    "2. For core_content, prioritize specific methods, metrics, and findings over vague descriptions.\n"
                    "3. If the document is the user uploaded experiment data, describe what metrics and figures it contains.\n\n"
                    "Output ONLY a valid JSON object:\n"
                    "{\n"
                    "  \"doc_time\": \"xxx\",\n"
                    "  \"source_authority\": \"xxx\",\n"
                    "  \"task_relevance\": \"xxx\",\n"
                    "  \"core_content\": \"xxx\",\n"
                    "  \"information_richness\": \"xxx\"\n"
                    "}\n\n"
                    "Important: Return ONLY the JSON object, no additional text."
                )

                # 鏋勫缓鐢ㄦ埛鎻愮ず
                user_prompt = (
                    f"DOCUMENT CONTENT:\n{content}\n"
                    # f"DOCUMENT LEN: The length of the file content is{len(content)}\n"
                    f"TASK FOR RELEVANCE ASSESSMENT: {task_content}"
                )

                # Get model URL and token from config
                config = get_config()
                model_config = config.get_custom_llm_config()

                model_url = model_config.get('url') or os.getenv('MODEL_REQUEST_URL',
                                                                 'http://127.0.0.1:8088/v1/chat/completions')
                model_token = model_config.get('token') or os.getenv('MODEL_REQUEST_TOKEN', '')
                actual_model = get_config().model_name

                headers = {
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {model_token}',
                    'csb-token': model_token
                }

                payload = {
                    "model": actual_model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": temperature,
                    "max_tokens": max_tokens
                }

                try:
                    max_retries = 3
                    response_json = None

                    from src.utils.llm_client import chat_completion_response

                    response_json = chat_completion_response(
                        payload,
                        model_config=model_config,
                        agent_name="document_extract",
                        request_logger=logger,
                    ).json()

                    if not response_json or "choices" not in response_json:
                        raise Exception("Failed to get valid response from API")

                    answer = response_json["choices"][0]["message"]["content"]

                    session_context = self.get_session_context()
                    session_id = session_context.get("session_id")
                    # 鍒囨崲淇濆瓨鐨勬柟寮?
                    conversation_history = [
                        # {"role": "system", "content": "system 1"},
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                        {"role": "assistant", "content": answer}
                    ]

                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': True,
                        'answer': answer,
                        'metadata': {
                            'file_size': len(content),
                            'line_count': len(content.splitlines())
                        }
                    }

                except Exception as e:
                    logger.error(f"Model API call failed for file '{file_path}': {e}")
                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': False,
                        'error': f"Model API error: {str(e)}"
                    }

            # 4. 骞跺彂澶勭悊鎵€鏈変换鍔?
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(tasks))) as executor:
                future_to_task = {executor.submit(process_single_task, task): task for task in tasks}

                for future in as_completed(future_to_task):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        task = future_to_task[future]
                        logger.error(f"Task processing failed for file '{task['file_path']}': {e}")
                        results.append({
                            'file_path': task['file_path'],
                            'task': task['task'],
                            'success': False,
                            'error': f"Task processing exception: {str(e)}"
                        })

            # 5. 淇濇寔鍘熷浠诲姟椤哄簭
            task_order = {task['file_path']: i for i, task in enumerate(tasks)}
            results.sort(key=lambda x: task_order.get(x['file_path'], float('inf')))

            # 淇濆瓨缁撴灉鍒版枃浠?
            def parse_answer_to_structured_data(answer_text: str, file_path: str) -> Dict[str, str]:
                """Parse model JSON safely and avoid index errors."""
                # 榛樿缁撴瀯
                structured_data = {
                    "file_path": file_path,
                    "doc_time": "Unknown",
                    "source_authority": "Unknown",
                    "task_relevance": "Unknown",
                    "information_richness": "Unknown",
                    "core_content": "Unknown"
                }

                if not answer_text:
                    return structured_data

                try:
                    # 浣跨敤姝ｅ垯琛ㄨ揪寮忔彁鍙栨渶澶栧眰鐨勫ぇ鎷彿鍐呭锛屽交搴曡В鍐?split 瀵艰嚧鐨?list index out of range 闂

                    json_match = re.search(r'\{.*\}', answer_text, re.DOTALL)

                    if json_match:
                        json_str = json_match.group(0)
                        parsed_data = json.loads(json_str)

                        # 鏇存柊瑙ｆ瀽鍒扮殑鍊?
                        if isinstance(parsed_data, dict):
                            for key in ["doc_time", "source_authority", "task_relevance", "core_content",
                                        "information_richness"]:
                                if key in parsed_data:
                                    structured_data[key] = str(parsed_data[key])
                    else:
                        structured_data["core_content"] = f"Parse warning: model did not return JSON in the required format. Preview: {answer_text[:100]}..."

                    return structured_data

                except Exception as e:
                    # 鎹曡幏鎵€鏈夎В鏋愰敊璇紝闃叉鏁翠釜鏅鸿兘浣撳穿婧?
                    structured_data["core_content"] = f"Parse exception: {str(e)}. Raw response preview: {answer_text[:100]}..."
                    return structured_data

            # 杞崲缁撴灉骞惰繃婊ゅ潖鏁版嵁
            structured_results = []
            for result in results:
                f_path = result.get('file_path', 'Unknown')
                if result.get('success', False) and result.get('answer'):
                    # 闃插尽锛氬鏋滆繑鍥炲唴瀹瑰寘鍚槑鏄剧殑浜岃繘鍒跺ご锛堝PDF鐗瑰緛锛夛紝鏍囪涓轰笉鍙
                    ans_text = result['answer']
                    if "%PDF-" in ans_text[:100] or "\x00" in ans_text[:100]:
                        structured_results.append({
                            "file_path": f_path,
                            "doc_time": "Unsupported",
                            "source_authority": "Binary Data",
                            "task_relevance": "N/A",
                            "information_richness": "Scarcity",
                            "core_content": "The file content appears to be binary data and cannot be processed by the language model."
                        })
                        continue

                    structured_data = parse_answer_to_structured_data(ans_text, f_path)
                    structured_results.append(structured_data)
                else:
                    # 澶辫触浠诲姟鐨勫鐞?
                    structured_results.append({
                        "file_path": f_path,
                        "doc_time": "Failed",
                        "source_authority": "Processing failed",
                        "task_relevance": "N/A",
                        "information_richness": "Unknown",
                        "core_content": f"浠诲姟鎵ц澶辫触: {result.get('error', '鏈煡閿欒')}"
                    })

            # 淇濆瓨缁撴灉鍒?JSONL 鏂囦欢
            analysis_path = storage_config.get('document_analysis_path', './doc_analysis')
            full_save_path = self.workspace_path / analysis_path / "file_analysis.jsonl"
            full_save_path.parent.mkdir(parents=True, exist_ok=True)

            # 璇诲彇鏃ф暟鎹互鍚堝苟
            existing_data = {}
            if full_save_path.exists():
                try:
                    with open(full_save_path, "r", encoding='utf-8', errors='ignore') as f:
                        for line in f:
                            try:
                                data = json.loads(line.strip())
                                if data.get('file_path'):
                                    existing_data[data['file_path']] = data
                            except:
                                continue
                except Exception as e:
                    logger.error(f"Failed to read existing analysis file: {e}")

            # 鍚堝苟鏂版棫缁撴灉
            for result in structured_results:
                if result.get('file_path'):
                    existing_data[result['file_path']] = result

            # 鎺掑簭锛氱敤鎴蜂笂浼犳枃浠朵紭鍏?
            user_uploaded_files = []
            other_files = []
            for file_data in existing_data.values():
                file_path = file_data.get('file_path', '')
                if 'user_uploads' in file_path or file_path.startswith('./user_uploads/'):
                    user_uploaded_files.append(file_data)
                else:
                    other_files.append(file_data)

            sorted_data = user_uploaded_files + other_files

            # 鍐欏叆鏂囦欢
            with open(full_save_path, mode="w", encoding='utf-8') as f:
                for data in sorted_data:
                    f.write(json.dumps(data, ensure_ascii=False) + "\n")

            # 鏃ュ織杈撳嚭涓庣粺璁?
            library_count = sum(1 for d in other_files if 'library_refs' in d.get('file_path', ''))
            research_count = sum(1 for d in other_files if 'research' in d.get('file_path', ''))
            if user_uploaded_files:
                logger.info(
                    f"Analysis saved: user_uploads({len(user_uploaded_files)}) + library_refs({library_count}) + research({research_count})")

            successful_tasks = len([r for r in results if r.get('success', False)])
            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_tasks': len(tasks),
                    'successful_tasks': successful_tasks,
                    'failed_tasks': len(tasks) - successful_tasks,
                    'model': model
                }
            )

        except Exception as e:
            logger.error(f"Unexpected error during document extraction: {e}")
            return MCPToolResult(success=False, error=str(e))

    def document_qa(
            self,
            tasks: List[Dict],
            model: str = None,
            temperature: float = 0.3,
            max_tokens: int = 8192,
            max_workers: int = 5
    ) -> MCPToolResult:
        # Answer questions based on content stored in local files.
        try:
            if model is None:
                model = get_config().model_name
            logger.info(f"鎴戠幇鍦ㄥ紑濮嬭皟鐢╠ocument_qa浜嗭細{tasks}")

            # 澶勭悊鍗曚釜浠诲姟
            def process_single_task(task: Dict) -> Dict:
                file_path = task['file_path']
                question = task['question']

                # 1. 璇诲彇鏂囦欢鍐呭
                read_result = self.file_read(file_path)
                if not read_result.success:
                    return {
                        'file_path': file_path,
                        'question': question,
                        'success': False,
                        'error': f"File read error: {read_result.error}",
                        'answer': None
                    }

                content = read_result.data

                # 2. 鏋勫缓绯荤粺鎻愮ず
                system_prompt = (
                    "You are an expert document analyst. Answer the user's question "
                    "based ONLY on the provided context. If the answer cannot be found "
                    "in the context, say 'I don't know'.\n\n"
                    "CONTEXT:\n{context}"
                ).format(context=content)

                # 3. Call the configured LLM through the unified client.
                from config.config import get_config
                from src.utils.llm_client import chat_completion_response

                try:
                    config_instance = get_config()
                    custom_model_config = config_instance.get_custom_llm_config()
                    actual_model = custom_model_config.get('model') or config_instance.model_name

                    payload = {
                        "model": actual_model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": question}
                        ],
                        "temperature": temperature,
                        "max_tokens": max_tokens
                    }

                    response = chat_completion_response(
                        payload,
                        model_config=custom_model_config,
                        agent_name="document_qa",
                        request_logger=logger,
                    )
                    response_json = response.json()

                    if not response_json or "choices" not in response_json:
                        raise Exception(f"Failed to get valid response from API: {response_json}")

                    answer = response_json["choices"][0]["message"]["content"]

                    return {
                        'file_path': file_path,
                        'question': question,  # 娉ㄦ剰锛歞ocument_qa 閲岀敤鐨勬槸 question
                        'success': True,
                        'answer': answer,
                        'metadata': {
                            'file_size': len(content),
                            'line_count': len(content.splitlines())
                        }
                    }

                except Exception as e:
                    logger.error(f"Model API call failed for file '{file_path}': {e}")
                    return {
                        'file_path': file_path,
                        'question': question,
                        'success': False,
                        'error': f"Model API error: {str(e)}"
                    }

            # 4. 骞跺彂澶勭悊鎵€鏈変换鍔?
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(tasks))) as executor:
                future_to_task = {executor.submit(process_single_task, task): task for task in tasks}

                for future in as_completed(future_to_task):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        task = future_to_task[future]
                        logger.error(f"Task processing failed for file '{task['file_path']}': {e}")
                        results.append({
                            'file_path': task['file_path'],
                            'question': task['question'],
                            'success': False,
                            'error': f"Task processing exception: {str(e)}"
                        })

            # 5. 淇濇寔鍘熷浠诲姟椤哄簭
            task_order = {task['file_path']: i for i, task in enumerate(tasks)}
            results.sort(key=lambda x: task_order.get(x['file_path'], float('inf')))

            # 6. 缁熻缁撴灉
            successful_tasks = len([r for r in results if r.get('success', False)])

            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_tasks': len(tasks),
                    'successful_tasks': successful_tasks,
                    'failed_tasks': len(tasks) - successful_tasks,
                    'model': model,
                    'concurrent_workers': min(max_workers, len(tasks))
                }
            )

        except Exception as e:
            logger.error(f"Context-based QA batch processing failed: {e}")
            logger.error(f"document qa batch processing failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ FILE DOWNLOAD TOOLS ================

    def download_files(
            self,
            urls: List[str],
            target_directory: str = None,
            overwrite: bool = False,
            max_file_size_mb: int = 1024
    ) -> MCPToolResult:
        """
        Download human-readable research files such as PDFs, documents, and data files.
        
        Use this tool for downloading research papers, documentation, reports, data files (CSV, JSON, XML),
        academic publications, and other human-readable content that you can analyze.
        
        WARNING: Do NOT use this tool for downloading web pages (HTML/HTM files) or other non-readable formats.
        For web page content extraction, use the url_crawler tool instead.
        
        Args:
            urls: List of URLs to download (PDFs, DOCs, research papers, data files, etc.)
            target_directory: Directory to save files (relative to session workspace)
            overwrite: Whether to overwrite existing files
            max_file_size_mb: Maximum file size in MB
        """
        try:
            if target_directory:
                # Ensure target_directory is relative to session workspace for security
                download_dir = self._safe_join(target_directory)
            else:
                download_dir = self.workspace_path / "downloads"

            download_dir.mkdir(parents=True, exist_ok=True)

            def download_single_file(url: str) -> Dict[str, Any]:
                # Download a single file with timeout protection
                try:
                    # Parse URL to get filename
                    parsed_url = urlparse(url)
                    filename = os.path.basename(parsed_url.path) or 'downloaded_file'

                    # Ensure filename has extension
                    if '.' not in filename:
                        filename += '.html'  # Default extension

                    if os.path.isabs(filename):
                        raise Exception(f"Path '{filename}' is absolute. Only relative paths are allowed.")
                    # 妫€娴嬫槸鍚负PDF鏂囦欢锛堢綉椤典笅杞斤級
                    is_pdf_from_web = filename.lower().endswith('.pdf')

                    file_path = download_dir / filename
                    if not os.path.realpath(file_path).startswith(self.full_workspace_path):
                        raise Exception(f"Path '{filename}' is outside workspace directory.")

                    # Check if file exists
                    if file_path.exists() and not overwrite:
                        return {
                            'url': url,
                            'success': False,
                            'error': 'File already exists',
                            'file_path': str(file_path)
                        }

                    # 馃憞 --- 鏍稿績鏇挎崲锛氬姞涓婇槻灏佹潃浼澶?+ 涓ユ牸瓒呮椂 --- 馃憞
                    import urllib3
                    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                    }
                    
                    # 馃殌 浼樺寲: 闄嶄綆瓒呮椂鏃堕棿浠?5绉掑埌30绉?閬垮厤涓€涓枃浠跺崱浣忓お涔?
                    response = requests.get(url, stream=True, timeout=(10, 30), headers=headers, verify=False)
                    # timeout=(connect_timeout, read_timeout) = (10绉掕繛鎺? 30绉掕鍙?
                    response.raise_for_status()
                    # 馃憜 --- 鏇挎崲缁撴潫 --- 馃憜

                    # Check file size from content-length header
                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > max_file_size_mb * 1024 * 1024:
                        return {
                            'url': url,
                            'success': False,
                            'error': f'File too large (>{max_file_size_mb}MB)',
                            'file_path': None
                        }

                    # Save file
                    with open(file_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)

                    return {
                        'url': url,
                        'success': True,
                        'file_path': str(file_path),
                        'file_size': file_path.stat().st_size
                    }

                except requests.exceptions.Timeout as e:
                    logger.warning(f"Download timeout (30s): {url}; skipping this file")
                    return {
                        'url': url,
                        'success': False,
                        'error': f'Download timeout (30s): {str(e)}',
                        'file_path': None
                    }
                except requests.exceptions.ConnectionError as e:
                    logger.warning(f"Download connection error: {url}; skipping this file")
                    return {
                        'url': url,
                        'success': False,
                        'error': f'Connection failed: {str(e)}',
                        'file_path': None
                    }
                except Exception as e:
                    logger.warning(f"Download failed: {url}; error: {str(e)}")
                    return {
                        'url': url,
                        'success': False,
                        'error': str(e),
                        'file_path': None
                    }

            # Process downloads concurrently
            results = []
            max_concurrent_downloads = min(5, len(urls))  # Limit concurrent downloads to avoid overwhelming servers
            with ThreadPoolExecutor(max_workers=max_concurrent_downloads) as executor:
                # Submit all download tasks
                future_to_url = {executor.submit(download_single_file, url): url for url in urls}

                # Collect results as they complete
                for future in as_completed(future_to_url):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        url = future_to_url[future]
                        logger.error(f"Download task failed for '{url}': {e}")
                        results.append({
                            'url': url,
                            'success': False,
                            'error': f"Download task exception: {str(e)}",
                            'file_path': None
                        })

            # Sort results to maintain original URL order
            url_order = {urls[i]: i for i in range(len(urls))}
            results.sort(key=lambda x: url_order.get(x['url'], float('inf')))

            # Generate status message following context management philosophy
            successful_downloads = len([r for r in results if r.get('success', False)])
            failed_downloads = len(results) - successful_downloads

            status_msg = f"File download task completed. Processed {len(urls)} URLs with {successful_downloads} successful downloads and {failed_downloads} failures. Files saved to {download_dir.relative_to(self.workspace_path)}. Use file reading tools to examine the downloaded files."

            return MCPToolResult(
                success=True,
                data=status_msg,
                metadata={
                    'download_directory': str(download_dir),
                    'total_urls': len(urls),
                    'successful_downloads': successful_downloads,
                    'failed_downloads': failed_downloads
                }
            )

        except Exception as e:
            logger.error(f"Download files failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def process_user_uploaded_files(
            self,
            file_ids: List[str],
            backend_url: str = "http://localhost:5000",
            target_subdir: str = "user_uploads"
    ) -> MCPToolResult:
        # Download user-uploaded files into the workspace.
        try:
            if not file_ids:
                return MCPToolResult(success=False, error="No file IDs provided")

            user_files_dir = self.workspace_path / target_subdir
            user_files_dir.mkdir(parents=True, exist_ok=True)

            # 璋冪敤 Flask 鍚庣 API
            response = requests.post(
                f"{backend_url}/api/user_files/download_and_parse",
                json={"file_ids": file_ids},
                timeout=60
            )

            if not response.ok:
                return MCPToolResult(success=False, error=f"HTTP {response.status_code}")

            result_data = response.json()
            processed_files = []
            # 鐢ㄤ簬璺熻釜鏂囦欢鍚嶏紝閬垮厤閲嶅鏂囦欢鍚嶅啿绐?
            used_filenames = set()

            for file_info in result_data.get('files', []):
                if not file_info.get('success'):
                    continue

                file_id = file_info.get('file_id', '')
                filename = file_info.get('filename', 'unknown.txt')

                # 瀹夊叏鍖栨枃浠跺悕锛氫繚鐣欎腑鏂囧拰鐗规畩瀛楃锛屽彧绉婚櫎Windows涓嶅厑璁哥殑瀛楃
                # Windows淇濈暀瀛楃鍜岃矾寰勫垎闅旂
                forbidden_chars = r'<>:"/\|?*'
                forbidden_chars += ''.join(chr(i) for i in range(32))  # 鎺у埗瀛楃
                safe_filename = "".join(c for c in filename if c not in forbidden_chars)
                safe_filename = safe_filename.strip(' .')  # 绉婚櫎棣栧熬绌烘牸鍜岀偣鍙?
                if not safe_filename or safe_filename == '.':
                    safe_filename = f'file_{file_id[:8]}.pdf' if filename.endswith(
                        '.pdf') else f'file_{file_id[:8]}.txt'

                # 澶勭悊鏂囦欢鍚嶅啿绐侊細濡傛灉鏂囦欢鍚嶅凡瀛樺湪锛屾坊鍔?file_id 鍓嶇紑
                if safe_filename in used_filenames:
                    # 鎻愬彇鏂囦欢鎵╁睍鍚?
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{file_id[:8]}_{name_part}.{ext_part}"
                    else:
                        safe_filename = f"{file_id[:8]}_{safe_filename}"

                # 濡傛灉杩樻槸鍐茬獊锛堟瀬缃曡鎯呭喌锛夛紝娣诲姞鏃堕棿鎴?
                if safe_filename in used_filenames:
                    import time
                    timestamp = str(int(time.time() * 1000))[-6:]
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{name_part}_{timestamp}.{ext_part}"
                    else:
                        safe_filename = f"{safe_filename}_{timestamp}"

                used_filenames.add(safe_filename)
                target_path = user_files_dir / safe_filename

                # 鏍规嵁鏂囦欢绫诲瀷鍐冲畾濡備綍澶勭悊
                file_type = file_info.get('file_type', '').lower()
                source_path = file_info.get('source_path')  # 鍘熷鏂囦欢璺緞锛堢敤浜庡鍒讹級

                # 浜岃繘鍒舵枃浠讹紙PDF銆乄ord銆丒xcel绛夛級闇€瑕佺洿鎺ュ鍒跺師濮嬫枃浠?
                binary_extensions = {
                    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.rtf', '.odt', '.epub',
                    '.zip', '.rar', '.7z', '.tar', '.gz', '.jpg', '.jpeg', '.png', '.bmp', '.webp'
                }
                is_binary = file_type in binary_extensions or any(
                    filename.lower().endswith(ext) for ext in binary_extensions)

                logger.info(
                    f"Processing file: {filename}, file_type={file_type}, is_binary={is_binary}, source_path={source_path}")

                if is_binary:
                    if not source_path:
                        logger.error(f"No source_path provided for binary file: {filename}")
                        continue

                    # 澶嶅埗浜岃繘鍒舵枃浠?
                    try:
                        source_file = Path(source_path)
                        if not source_file.exists():
                            logger.error(f"Source file not found: {source_path} (absolute: {source_file.absolute()})")
                            continue

                        # 鎵ц澶嶅埗
                        shutil.copy2(source_file, target_path)
                        copied_size = target_path.stat().st_size
                        source_size = source_file.stat().st_size
                        logger.info(
                            f"Copied binary file: {source_file} ({source_size} bytes) -> {target_path} ({copied_size} bytes)")

                        # 楠岃瘉澶嶅埗鏄惁鎴愬姛
                        if copied_size != source_size:
                            logger.warning(f"File size mismatch: source={source_size}, target={copied_size}")

                        # 瀵逛簬PDF鍜孌OCX鏂囦欢锛屽皾璇曟彁鍙栧苟缂撳瓨鏂囨湰鐗堟湰锛堝彲閫変紭鍖栵級
                        if file_type == '.pdf':
                            try:
                                # 鎻愬彇PDF鏂囨湰
                                extracted_text = self._read_pdf_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 淇濆瓨鏂囨湰缂撳瓨锛堟枃浠跺悕.pdf.txt锛?
                                    text_cache_path = target_path.with_suffix('.pdf.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for PDF: {text_cache_path} ({len(extracted_text)} chars)")
                                    # 鏇存柊杩斿洖璺緞锛屼紭鍏堜娇鐢ㄦ枃鏈増鏈紙Agent鍙互鐩存帴璇诲彇鏂囨湰锛?
                                    # 娉ㄦ剰锛歅DF鍘熸枃浠嶇劧淇濈暀锛屾枃鏈紦瀛樻槸鍙€夌殑
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for PDF {filename}: {e}")
                                # 缂撳瓨澶辫触涓嶅奖鍝嶄富娴佺▼锛岀户缁娇鐢≒DF鍘熸枃

                        elif file_type == '.docx':
                            try:
                                # 鎻愬彇DOCX鏂囨湰
                                extracted_text = self._read_docx_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 淇濆瓨鏂囨湰缂撳瓨锛堟枃浠跺悕.docx.txt锛?
                                    text_cache_path = target_path.with_suffix('.docx.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for DOCX: {text_cache_path} ({len(extracted_text)} chars)")
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for DOCX {filename}: {e}")
                                # 缂撳瓨澶辫触涓嶅奖鍝嶄富娴佺▼锛岀户缁娇鐢―OCX鍘熸枃

                        elif file_type == '.doc':
                            try:
                                # 鎻愬彇DOC鏂囨湰
                                extracted_text = self._read_doc_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 淇濆瓨鏂囨湰缂撳瓨锛堟枃浠跺悕.doc.txt锛?
                                    text_cache_path = target_path.with_suffix('.doc.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for DOC: {text_cache_path} ({len(extracted_text)} chars)")
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for DOC {filename}: {e}")
                                # 缂撳瓨澶辫触涓嶅奖鍝嶄富娴佺▼锛岀户缁娇鐢―OC鍘熸枃

                    except Exception as e:
                        logger.error(f"Error copying binary file {filename}: {e}", exc_info=True)
                        continue
                else:
                    # 鏂囨湰鏂囦欢锛氬啓鍏ユ枃鏈唴瀹?
                    try:
                        content = file_info.get('content', '')
                        with open(target_path, 'w', encoding='utf-8') as f:
                            f.write(content)
                        logger.info(f"Written text file: {target_path} ({len(content)} chars)")
                    except Exception as e:
                        logger.error(f"Error writing text file {filename}: {e}", exc_info=True)
                        continue

                processed_files.append({
                    'file_id': file_id,
                    'filename': filename,
                    'local_path': f"./{target_subdir}/{safe_filename}",
                    'content_length': file_info.get('content_length', 0),
                    'is_user_uploaded': True,
                    'priority': 'high',
                    'success': True
                })

            return MCPToolResult(
                success=True,
                data={
                    'files': processed_files,
                    'total_files': len(processed_files),
                    'user_files_directory': str(user_files_dir)
                }
            )

        except Exception as e:
            logger.error(f"Error processing user files: {e}")
            return MCPToolResult(success=False, error=str(e))

    def process_library_files(
            self,
            file_ids: List[str],
            backend_url: str = "http://localhost:5000",
            target_subdir: str = "library_refs"
    ) -> MCPToolResult:
        # Download selected library files into the workspace.
        try:
            if not file_ids:
                return MCPToolResult(success=False, error="No file IDs provided")

            library_files_dir = self.workspace_path / target_subdir
            library_files_dir.mkdir(parents=True, exist_ok=True)

            # 璋冪敤 Flask 鍚庣 API锛堜笌 user_uploads 浣跨敤鐩稿悓鐨?API锛?
            response = requests.post(
                f"{backend_url}/api/user_files/download_and_parse",
                json={"file_ids": file_ids},
                timeout=60
            )

            if not response.ok:
                return MCPToolResult(success=False, error=f"HTTP {response.status_code}")

            result_data = response.json()
            processed_files = []
            used_filenames = set()

            for file_info in result_data.get('files', []):
                if not file_info.get('success'):
                    continue

                file_id = file_info.get('file_id', '')
                filename = file_info.get('filename', 'unknown.txt')

                # 瀹夊叏鍖栨枃浠跺悕
                forbidden_chars = r'<>:"/\|?*'
                forbidden_chars += ''.join(chr(i) for i in range(32))
                safe_filename = "".join(c for c in filename if c not in forbidden_chars)
                safe_filename = safe_filename.strip(' .')
                if not safe_filename or safe_filename == '.':
                    safe_filename = f'file_{file_id[:8]}.pdf' if filename.endswith(
                        '.pdf') else f'file_{file_id[:8]}.txt'

                # 澶勭悊鏂囦欢鍚嶅啿绐?
                if safe_filename in used_filenames:
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{file_id[:8]}_{name_part}.{ext_part}"
                    else:
                        safe_filename = f"{file_id[:8]}_{safe_filename}"

                if safe_filename in used_filenames:
                    import time
                    timestamp = str(int(time.time() * 1000))[-6:]
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{name_part}_{timestamp}.{ext_part}"
                    else:
                        safe_filename = f"{safe_filename}_{timestamp}"

                used_filenames.add(safe_filename)
                target_path = library_files_dir / safe_filename

                # 鏍规嵁鏂囦欢绫诲瀷鍐冲畾濡備綍澶勭悊
                file_type = file_info.get('file_type', '').lower()
                source_path = file_info.get('source_path')

                binary_extensions = {
                    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.rtf', '.odt', '.epub',
                    '.zip', '.rar', '.7z', '.tar', '.gz', '.jpg', '.jpeg', '.png', '.bmp', '.webp'
                }
                is_binary = file_type in binary_extensions or any(
                    filename.lower().endswith(ext) for ext in binary_extensions)

                if is_binary and source_path:
                    # 澶嶅埗浜岃繘鍒舵枃浠?
                    try:
                        source_file = Path(source_path)
                        if source_file.exists():
                            shutil.copy2(source_file, target_path)

                            # 鍒涘缓鏂囨湰缂撳瓨锛堜笌 user_uploads 鐩稿悓锛?
                            if file_type == '.pdf':
                                try:
                                    extracted_text = self._read_pdf_text(target_path)
                                    if extracted_text and len(extracted_text.strip()) > 100:
                                        text_cache_path = target_path.with_suffix('.pdf.txt')
                                        with open(text_cache_path, 'w', encoding='utf-8') as f:
                                            f.write(extracted_text)
                                        logger.info(f"Created text cache for library PDF: {text_cache_path}")
                                except Exception as e:
                                    logger.debug(f"Failed to create text cache for PDF {filename}: {e}")

                            elif file_type == '.docx':
                                try:
                                    extracted_text = self._read_docx_text(target_path)
                                    if extracted_text and len(extracted_text.strip()) > 100:
                                        text_cache_path = target_path.with_suffix('.docx.txt')
                                        with open(text_cache_path, 'w', encoding='utf-8') as f:
                                            f.write(extracted_text)
                                        logger.info(f"Created text cache for library DOCX: {text_cache_path}")
                                except Exception as e:
                                    logger.debug(f"Failed to create text cache for DOCX {filename}: {e}")
                    except Exception as e:
                        logger.error(f"Failed to copy library file {filename}: {e}")
                        continue
                else:
                    # 鏂囨湰鏂囦欢鐩存帴鍐欏叆
                    content = file_info.get('content', '')
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(content)

                processed_files.append({
                    'file_id': file_id,
                    'filename': safe_filename,
                    'file_path': str(target_path.relative_to(self.workspace_path)),
                    'file_type': file_type,
                    'is_library_file': True,
                    'priority': 'normal',  # 鍏抽敭锛氫笉鏍囪涓洪珮浼樺厛绾?
                    'success': True
                })

            logger.info(f"Library files processed: {len(processed_files)} files saved to {target_subdir}/")

            return MCPToolResult(
                success=True,
                data={
                    'files': processed_files,
                    'total_files': len(processed_files),
                    'library_files_directory': str(library_files_dir)
                }
            )

        except Exception as e:
            logger.error(f"Error processing library files: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ WORKSPACE TOOLS ================

    def list_workspace(
            self,
            path: str = None,
            recursive: bool = False,
            include_hidden: bool = False,
            max_depth: int = 3
    ) -> MCPToolResult:
        """
        List files and directories in workspace with tree structure visualization
        """
        try:
            if path:
                target_path = self._safe_join(path)
            else:
                target_path = self.workspace_path

            if not target_path.exists():
                return MCPToolResult(success=False, error=f"Path does not exist: {target_path}")

            if not target_path.is_dir():
                return MCPToolResult(success=False, error=f"Path is not a directory: {target_path}")

            items = []
            tree_structure = []

            def _list_items(current_path: Path, current_depth: int = 0):
                if current_depth > max_depth: return
                try:
                    all_items = list(current_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]
                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))
                    for item in all_items:
                        item_info = {
                            'name': item.name,
                            'path': str(item.relative_to(self.workspace_path)),
                            'type': 'directory' if item.is_dir() else 'file',
                            'size': item.stat().st_size if item.is_file() else None,
                            'modified': item.stat().st_mtime,
                            'depth': current_depth
                        }
                        items.append(item_info)
                        if recursive and item.is_dir():
                            _list_items(item, current_depth + 1)
                except PermissionError:
                    pass

            def _generate_tree_structure(current_path: Path, prefix: str = "", is_last: bool = True,
                                         current_depth: int = 0):
                if current_depth > max_depth: return
                try:
                    all_items = list(current_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]
                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))
                    for i, item in enumerate(all_items):
                        is_last_item = i == len(all_items) - 1
                        current_symbol = "鈹斺攢鈹€ " if is_last_item else "鈹溾攢鈹€ "
                        extension = "    " if is_last_item else "鈹?  "

                        if item.is_dir():
                            name_with_indicator = f"馃搧 {item.name}/"
                        else:
                            name_with_indicator = f"馃搫 {item.name}"

                        tree_line = prefix + current_symbol + name_with_indicator
                        tree_structure.append(tree_line)

                        # 馃殌 绗竴灞傜啍鏂細濡傛灉鏂囦欢鏁伴噺瓒呰繃200涓紝鐩存帴鍋滄閫掑綊锛岄槻姝㈠唴瀛樻寰幆
                        if len(tree_structure) > 200:
                            return

                        if recursive and item.is_dir():
                            _generate_tree_structure(item, prefix + extension, is_last_item, current_depth + 1)
                except PermissionError:
                    tree_structure.append(prefix + "鈹斺攢鈹€ [Permission Denied]")

            _list_items(target_path)
            root_name = target_path.name if target_path.name else "workspace"
            tree_structure.append(f"馃搧 {root_name}/")

            if recursive:
                _generate_tree_structure(target_path)
            else:
                try:
                    all_items = list(target_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]
                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))
                    for i, item in enumerate(all_items):
                        is_last_item = i == len(all_items) - 1
                        symbol = "鈹斺攢鈹€ " if is_last_item else "鈹溾攢鈹€ "
                        name_with_indicator = f"馃搧 {item.name}/" if item.is_dir() else f"馃搫 {item.name}"
                        tree_structure.append(symbol + name_with_indicator)
                except PermissionError:
                    pass

            tree_string = "\n".join(tree_structure)

            # 馃殌 绗簩灞傛牳蹇冮槻鐖嗙浘锛氫弗鏍奸檺鍒惰繑鍥炵粰澶фā鍨嬬殑瀛楃涓查暱搴?
            MAX_TREE_LENGTH = 4000
            if len(tree_string) > MAX_TREE_LENGTH or len(tree_structure) > 200:
                truncated_msg = "\n... [System truncation: directory output is too large, so it has been shortened. Do not keep listing this directory; proceed with the task using this summary.] ..."
                tree_string = tree_string[:MAX_TREE_LENGTH] + truncated_msg

            return MCPToolResult(
                success=True,
                data={
                    'items': items[:50],  # 鎴柇瀛楀吀
                    'tree_structure': tree_string,
                    'tree_lines': tree_structure[:50]  # 鎴柇鍒楄〃
                },
                metadata={'total_items': len(items), 'is_truncated': len(tree_structure) > 50}
            )

        except Exception as e:
            return MCPToolResult(success=False, error=str(e))

    # ================ FILE EDITING TOOLS ================
    def str_replace_based_edit_tool(
            self,
            action: str,
            file_path: str,
            content: str = None,
            old_str: str = None,
            new_str: str = None,
            line_number: int = None,
            max_char_len: int = 10000,
    ) -> MCPToolResult:
        # Comprehensive file editing tool.
        try:
            full_path = self._safe_join(file_path)

            if action == 'create':
                if full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File already exists: {file_path}"
                    )

                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(content or '', encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"File created: {file_path}",
                    metadata={'file_size': full_path.stat().st_size}
                )

            elif action == 'view':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                content = full_path.read_text(encoding='utf-8')
                if len(content) > max_char_len:
                    content = ("Due to the content being too long, only the first 10,000 characters are returned. "
                               "It is recommended to use other tools such as `document_qa` to extract the required content from the file. "
                               "Below is the returned portion of the file content: \n\n") + content[:max_char_len]

                return MCPToolResult(
                    success=True,
                    data=content,
                    metadata={
                        'file_size': len(content),
                        'line_count': len(content.splitlines())
                    }
                )

            elif action == 'str_replace':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                if not old_str or new_str is None:
                    return MCPToolResult(
                        success=False,
                        error="Both old_str and new_str are required for str_replace"
                    )

                original_content = full_path.read_text(encoding='utf-8')

                if old_str not in original_content:
                    return MCPToolResult(
                        success=False,
                        error=f"String not found: {old_str[:50]}..."
                    )

                new_content = original_content.replace(old_str, new_str)
                full_path.write_text(new_content, encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"Replaced {original_content.count(old_str)} occurrence(s)",
                    metadata={
                        'old_size': len(original_content),
                        'new_size': len(new_content)
                    }
                )

            elif action == 'insert':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                if line_number is None or content is None:
                    return MCPToolResult(
                        success=False,
                        error="Both line_number and content are required for insert"
                    )

                lines = full_path.read_text(encoding='utf-8').splitlines()

                if line_number < 0 or line_number > len(lines):
                    return MCPToolResult(
                        success=False,
                        error=f"Invalid line number: {line_number}"
                    )

                lines.insert(line_number, content)
                full_path.write_text('\n'.join(lines), encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"Inserted content at line {line_number}",
                    metadata={'new_line_count': len(lines)}
                )

            elif action == 'append':
                if not full_path.exists():
                    full_path.touch()

                with open(full_path, 'a', encoding='utf-8') as f:
                    f.write(content or '')

                return MCPToolResult(
                    success=True,
                    data=f"Appended content to {file_path}",
                    metadata={'file_size': full_path.stat().st_size}
                )

            elif action == 'delete':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                full_path.unlink()

                return MCPToolResult(
                    success=True,
                    data=f"Deleted file: {file_path}"
                )

            else:
                return MCPToolResult(
                    success=False,
                    error=f"Unknown action: {action}"
                )

        except Exception as e:
            logger.error(f"File edit failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def analyze_image(
            self,
            file_path: str,
            prompt: str = "Describe this image in detail. If it contains charts or tables, extract the visible values and labels."
    ) -> MCPToolResult:
        try:
            full_path = self._safe_join(file_path)
            if not full_path.exists():
                return MCPToolResult(success=False, error=f"Image file does not exist: {file_path}")

            import base64

            # 鏄惧紡MIME鏄犲皠锛屼笉渚濊禆mimetypes鐚滄祴
            ext = full_path.suffix.lower()
            mime_map = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.webp': 'image/webp',
                '.gif': 'image/gif',
                '.bmp': 'image/png'
            }
            mime_type = mime_map.get(ext, 'image/jpeg')

            with open(full_path, "rb") as f:
                encoded_string = base64.b64encode(f.read()).decode('utf-8')

            from config.config import get_config
            config = get_config()
            model_config = config.get_custom_llm_config()

            model_url = model_config.get('url')
            model_token = model_config.get('token')
            model_name = model_config.get('model')

            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {model_token}',
                'csb-token': model_token
            }

            # 鍏抽敭锛氳瑙夎姹備笉浼?chat_template 鍜?spaces_between_special_tokens
            # 杩欎袱涓瓧娈垫槸鏂囨湰妯″瀷涓撶敤鐨勶紝浼犲叆浼氱牬鍧忓妯℃€佸浘鐗囪В鏋?
            payload = {
                "model": model_name,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {
                            "url": f"data:{mime_type};base64,{encoded_string}"
                        }}
                    ]
                }],
                "max_tokens": 1024
            }

            response = requests.post(
                url=model_url,
                headers=headers,
                json=payload,
                timeout=60
            )
            response_json = response.json()

            if "choices" in response_json and response_json["choices"]:
                text_result = response_json["choices"][0]["message"]["content"]
                return MCPToolResult(
                    success=True,
                    data=f"Vision model analysis result:\n{text_result}"
                )
            else:
                return MCPToolResult(
                    success=False,
                    error=f"Model did not return a valid result: {response_json}"
                )

        except Exception as e:
            logger.error(f"鍥剧墖鍒嗘瀽澶辫触: {e}")
            return MCPToolResult(success=False, error=str(e))

    def markdown_to_pdf(self, markdown_path: str, output_path: str = None) -> MCPToolResult:
        # Convert Markdown to PDF.
        try:
            full_md_path = self._safe_join(markdown_path)
            if not full_md_path.exists():
                return MCPToolResult(success=False, error=f"Markdown file does not exist: {markdown_path}")

            if output_path is None:
                output_path = markdown_path.replace('.md', '.pdf', 1)

            full_pdf_path = self._safe_join(output_path)
            full_pdf_path.parent.mkdir(parents=True, exist_ok=True)

            with open(full_md_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = f.read()

            # Use Chrome headless + MathJax for high-quality PDF with proper math rendering
            success = _generate_pdf_with_chrome(md_content, full_pdf_path, str(full_md_path.parent))
            if not success:
                return MCPToolResult(success=False, error="Chrome headless PDF generation failed. Please check Chrome installation.")

            logger.info(f"PDF generated: {markdown_path} -> {output_path}")
            return MCPToolResult(
                success=True,
                data={
                    "pdf_path": output_path,
                    "pdf_absolute_path": str(full_pdf_path)
                }
            )

        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return MCPToolResult(success=False, error=f"PDF generation failed: {str(e)}")

    # ================ BASIC FILE TOOLS ================

    def file_read(self, file_path: str, encoding: str = 'utf-8', max_char_len: int = 10000) -> MCPToolResult:
        # Read file content
        try:
            full_path = self._safe_join(file_path)
            if not full_path.exists():
                return MCPToolResult(success=False, error=f"File does not exist: {file_path}")

            file_size = full_path.stat().st_size
            if file_size > 10 * 1024 * 1024:
                return MCPToolResult(success=False, error="File exceeds 10MB; direct reading is refused. Use a specialized extraction or analysis tool.")

            # 馃殌 鏍稿績淇锛氶潤榛樿瑙夎В鏋?+ 瀹岀編鍏滃簳鏈哄埗
            if full_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif']:
                return MCPToolResult(
                    success=True,
                    data=(
                        f'Image file: {file_path}\\n'
                        f'Size: {file_size} bytes\\n'
                        'This file was not automatically sent to a vision model to avoid API timeouts. '
                        'If visual data extraction is required, call analyze_image with a focused prompt. '
                        f'When writing the paper, insert it with Markdown such as ![Figure]({file_path}).'
                    ),
                    metadata={'file_size': file_size, 'is_image': True}
                )
            # PDF 瑙ｆ瀽涓庢枃鏈В鏋愪繚鐣欎笉鍙?
            if full_path.suffix.lower() == '.pdf':
                content = self._read_pdf_text(full_path)
                if not content:
                    return MCPToolResult(success=False, error="PDF text extraction failed.")
            else:
                with open(full_path, 'rb') as f:
                    if b'\x00' in f.read(1024):
                        return MCPToolResult(success=False, error="Unsupported binary file detected.")
                content = full_path.read_text(encoding=encoding, errors='ignore')

            if len(content) > max_char_len:
                content = f"Content is too long; returning the first {max_char_len} characters only.\n\n" + content[:max_char_len]

            return MCPToolResult(success=True, data=content, metadata={'file_size': file_size})

        except Exception as e:
            return MCPToolResult(success=False, error=str(e))

    def read_image_for_llm(self, file_path: str) -> MCPToolResult:
        # Read an image and return a data URL for LLM vision calls.
        try:
            full_path = self._safe_join(file_path)
            if not full_path.exists():
                return MCPToolResult(success=False, error=f"Image does not exist: {file_path}")

            import base64
            import mimetypes
            mime_type, _ = mimetypes.guess_type(full_path)
            if not mime_type or not mime_type.startswith('image/'):
                mime_type = 'image/jpeg'

            with open(full_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')

            return MCPToolResult(
                success=True,
                # 娉ㄦ剰杩欓噷鐨勬暟鎹粨鏋勶紝澶фā鍨嬭兘鐩存帴鐞嗚В data:image 鍗忚
                data={
                    "is_vision_content": True,
                    "image_url": f"data:{mime_type};base64,{encoded_string}"
                },
                metadata={'file_size': full_path.stat().st_size}
            )
        except Exception as e:
            logger.error(f"Image read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _read_pdf_text(self, path: Path) -> str:
        # Read PDF text with multiple fallbacks.
        # 楠岃瘉鏂囦欢鏄惁瀛樺湪涓斾笉涓虹┖
        if not path.exists():
            logger.error(f"PDF file does not exist: {path}")
            return ''

        file_size = path.stat().st_size
        if file_size == 0:
            logger.error(f"PDF file is empty: {path}")
            return ''

        if file_size < 100:  # PDF 鏂囦欢鑷冲皯搴旇鏈夊嚑鐧惧瓧鑺?
            logger.warning(f"PDF file suspiciously small ({file_size} bytes): {path}")

        # 妫€鏌ユ枃浠跺ご鏄惁涓?PDF 鏍煎紡锛?PDF-锛?
        try:
            with open(path, 'rb') as f:
                header = f.read(5)
                if header != b'%PDF-':
                    logger.error(f"File is not a valid PDF (header: {header}): {path}")
                    return ''
        except Exception as e:
            logger.error(f"Failed to read file header for {path}: {e}")
            return ''

        # 1. 浼樺厛浣跨敤 PyMuPDF
        try:
            import fitz  # PyMuPDF
            text = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    text.append(page.get_text())
            result = "\n".join(text)
            if result.strip():
                logger.info(f"Successfully extracted PDF text using PyMuPDF: {path}")
                return result
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {path}: {e}")

        # 2. 鍏舵浣跨敤 pdfminer.six
        try:
            from pdfminer.high_level import extract_text
            result = extract_text(str(path)) or ''
            if result.strip():
                logger.info(f"Successfully extracted PDF text using pdfminer.six: {path}")
                return result
        except Exception as e:
            logger.debug(f"pdfminer.six failed for {path}: {e}")

        # 3. 鍐嶆浣跨敤 PyPDF2锛堟枃鏈川閲忎竴鑸級
        try:
            import PyPDF2
            text = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or '')
            result = "\n".join(text)
            if result.strip():
                logger.info(f"Successfully extracted PDF text using PyPDF2: {path}")
                return result
        except Exception as e:
            logger.debug(f"PyPDF2 failed for {path}: {e}")

        logger.warning(f"All PDF extraction methods failed for {path}")
        return ''

    def _read_docx_text(self, path: Path) -> str:
        # Read text from a DOCX file.
        try:
            from docx import Document
            doc = Document(str(path))

            # 鎻愬彇鎵€鏈夋钀芥枃鏈?
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # 鎻愬彇琛ㄦ牸涓殑鏂囨湰
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(' | '.join(row_text))

            # 鍚堝苟娈佃惤鍜岃〃鏍兼枃鏈?
            all_text = paragraphs + tables_text
            result = '\n'.join(all_text)

            if result.strip():
                logger.info(f"Successfully extracted text from DOCX: {path} ({len(result)} chars)")
                return result
            else:
                logger.warning(f"DOCX file appears to be empty: {path}")
                return ''

        except ImportError:
            logger.error(f"python-docx library not installed. Please install: pip install python-docx")
            return ''
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {path}: {e}")
            return ''

    def _read_doc_text(self, path: Path) -> str:
        # Read text from legacy DOC files with platform-specific fallbacks.
        import sys
        import subprocess

        # Windows浼樺厛灏濊瘯win32com锛堝鏋滃畨瑁呬簡MS Word锛屾晥鏋滄渶濂斤級
        if sys.platform == 'win32':
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(path.absolute()))
                text = doc.Content.Text
                doc.Close()
                word.Quit()

                if text.strip():
                    logger.info(f"Successfully extracted text from DOC using win32com: {path} ({len(text)} chars)")
                    return text.strip()
            except ImportError:
                logger.debug(f"win32com not installed (pip install pywin32)")
            except Exception as e:
                logger.debug(f"win32com failed for {path}: {e}")

        # 鏂规1: 浣跨敤antiword锛堣法骞冲彴锛孡inux鏈嶅姟鍣ㄦ帹鑽愶級
        try:
            result = subprocess.run(
                ['antiword', str(path)],
                capture_output=True,
                text=True,
                timeout=30,
                check=False  # 涓嶆姏鍑哄紓甯革紝閫氳繃returncode鍒ゆ柇
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                logger.info(f"Successfully extracted text from DOC using antiword: {path} ({len(text)} chars)")
                return text
            else:
                logger.debug(f"antiword returned code {result.returncode}, stderr: {result.stderr}")
        except FileNotFoundError:
            logger.debug(
                f"antiword not found. Install: apt-get install antiword (Linux) or choco install antiword (Windows)")
        except Exception as e:
            logger.debug(f"antiword failed for {path}: {e}")

        # 鏂规2: 浣跨敤textract锛圥ython搴擄紝璺ㄥ钩鍙颁絾渚濊禆杈冨锛?
        try:
            import textract
            text = textract.process(str(path)).decode('utf-8')
            if text.strip():
                logger.info(f"Successfully extracted text from DOC using textract: {path} ({len(text)} chars)")
                return text.strip()
        except ImportError:
            logger.debug(f"textract not installed (pip install textract)")
        except Exception as e:
            logger.debug(f"textract failed for {path}: {e}")

        # 鎵€鏈夋柟娉曢兘澶辫触
        logger.warning(f"All DOC extraction methods failed for {path}. Recommendations:")
        logger.warning(f"  1. Install antiword: apt-get install antiword (Linux) or choco install antiword (Windows)")
        logger.warning(f"  2. Install pywin32: pip install pywin32 (Windows with MS Word)")
        logger.warning(f"  3. Install textract: pip install textract (requires system dependencies)")
        logger.warning(f"  4. Or convert the file to .docx/.pdf format manually")
        return ''  # 杩斿洖绌哄瓧绗覆锛岃璋冪敤鏂瑰鐞?

    def _normalize_report_part_path(self, file_path: str) -> str:
        # Normalize misformatted report chapter filenames.
        try:
            path_obj = Path(file_path)
            filename = path_obj.name
            parts = path_obj.parts
            # Only normalize files under a report directory to avoid
            # touching unrelated paths.
            if "report" not in parts and path_obj.parent.name != "report":
                return file_path

            match = re.match(r'^part_(\d+)[\._]\d+.*\.md$', filename)
            if not match:
                return file_path

            chapter_index = match.group(1)
            normalized = path_obj.with_name(f"part_{chapter_index}.md")
            return normalized.as_posix()
        except Exception:
            # On any error, fall back to the original path
            return file_path

    def _clean_report_artifacts(self, content: str) -> str:
        # Clean internal marker tokens from final report content.
        if not content:
            return content

        try:
            # Remove any [unusedXX] style control tokens
            content = re.sub(r"\[unused\d+\]", "", content)

            # Normalize [webpaeg22] or [webpage22] -> [22]
            content = re.sub(r"\[webp(?:aeg|age)(\d+)\]", r"[\1]", content)
        except Exception:
            # On regex errors, return original content to be safe
            return content

        return content

    # ================ ENHANCED FILE ANALYSIS TOOLS ================

    def file_stats(self, file_path: str) -> MCPToolResult:
        # Get comprehensive file statistics without reading full content.
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            if not full_path.is_file():
                return MCPToolResult(
                    success=False,
                    error=f"Path is not a file: {file_path}"
                )

            # Get basic file stats
            stat_info = full_path.stat()
            file_size = stat_info.st_size

            # Quick content analysis without loading full file
            encoding = 'utf-8'
            line_count = 0
            word_count = 0
            char_count = 0
            first_lines = []
            last_lines = []

            try:
                with open(full_path, 'r', encoding=encoding, errors='ignore') as f:
                    # Read first few lines for preview
                    for i, line in enumerate(f):
                        line_count += 1
                        if i < 5:  # First 5 lines
                            first_lines.append(line.rstrip())

                        char_count += len(line)
                        word_count += len(line.split())

                        # For efficiency, stop detailed counting after reasonable limit
                        if line_count > 10000:
                            # Estimate remaining based on average
                            remaining_size = file_size - f.tell()
                            if remaining_size > 0:
                                avg_line_size = f.tell() / line_count
                                estimated_remaining_lines = int(remaining_size / avg_line_size)
                                line_count += estimated_remaining_lines

                                # Estimate words and chars
                                avg_chars_per_line = char_count / min(line_count, 10000)
                                avg_words_per_line = word_count / min(line_count, 10000)
                                char_count += int(remaining_size)
                                word_count += int(estimated_remaining_lines * avg_words_per_line)
                            break

                # Get last few lines if file is reasonable size
                if file_size < 1024 * 1024:  # Less than 1MB
                    with open(full_path, 'r', encoding=encoding, errors='ignore') as f:
                        lines = f.readlines()
                        last_lines = [line.rstrip() for line in lines[-5:]]
                        if line_count <= 10000:  # Recalculate if we estimated
                            line_count = len(lines)
                            char_count = sum(len(line) for line in lines)
                            word_count = sum(len(line.split()) for line in lines)

            except Exception as e:
                # Try binary mode to at least get size info
                encoding = 'binary'
                char_count = file_size

            # Determine file type
            file_extension = full_path.suffix.lower()
            file_type = self._detect_file_type(full_path, file_extension)

            # Reading recommendation
            reading_recommendation = self._get_reading_recommendation(
                file_size, line_count, word_count, file_type
            )

            stats = {
                'file_path': file_path,
                'file_size_bytes': file_size,
                'file_size_human': self._format_file_size(file_size),
                'line_count': line_count,
                'word_count': word_count,
                'character_count': char_count,
                'encoding': encoding,
                'file_type': file_type,
                'file_extension': file_extension,
                'modified_time': stat_info.st_mtime,
                'is_large_file': file_size > 1024 * 1024,  # > 1MB
                'is_very_large_file': file_size > 10 * 1024 * 1024,  # > 10MB
                'first_lines_preview': first_lines,
                'last_lines_preview': last_lines,
                'reading_recommendation': reading_recommendation
            }

            return MCPToolResult(
                success=True,
                data=stats,
                metadata={
                    'analysis_method': 'efficient_sampling' if line_count > 10000 else 'full_analysis'
                }
            )

        except Exception as e:
            logger.error(f"File stats failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    @staticmethod
    def _detect_file_type(file_path: Path, extension: str) -> str:
        # Detect file type based on extension and content

        # Extension-based detection
        type_map = {
            '.py': 'python_code',
            '.js': 'javascript_code',
            '.ts': 'typescript_code',
            '.java': 'java_code',
            '.cpp': 'cpp_code',
            '.c': 'c_code',
            '.html': 'html_markup',
            '.css': 'css_stylesheet',
            '.json': 'json_data',
            '.xml': 'xml_data',
            '.yaml': 'yaml_config',
            '.yml': 'yaml_config',
            '.md': 'markdown_document',
            '.txt': 'plain_text',
            '.csv': 'csv_data',
            '.sql': 'sql_code',
            '.sh': 'shell_script',
            '.dockerfile': 'docker_config',
            '.env': 'environment_config'
        }

        if extension in type_map:
            return type_map[extension]

        # Content-based detection for unknown extensions
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                first_line = f.readline().strip()

                if first_line.startswith('#!'):
                    return 'executable_script'
                elif first_line.startswith('<?xml'):
                    return 'xml_data'
                elif first_line.startswith('{') or first_line.startswith('['):
                    return 'json_data'
                elif 'DOCTYPE html' in first_line or '<html' in first_line:
                    return 'html_markup'
        except:
            pass

        return 'unknown_text'

    def _format_file_size(self, size_bytes: int) -> str:
        # Format file size in human readable format
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"

    def _get_reading_recommendation(self, file_size: int, line_count: int,
                                    word_count: int, file_type: str) -> Dict[str, Any]:
        # Provide intelligent recommendations for how to read the file

        recommendations = {
            'strategy': 'full_read',
            'reason': 'File is small enough for full reading',
            'alternatives': []
        }

        # Large file strategies
        if file_size > 1024 * 1024:  # > 1MB
            recommendations['strategy'] = 'selective_read'
            recommendations['reason'] = 'File is large, consider targeted approaches'
            recommendations['alternatives'] = [
                'Use file_grep_with_context to search for specific content',
                'Use content_preview to get overview before full read',
                'Use file_read_lines to read specific sections',
                'Content indexing has been disabled'
            ]

        elif line_count > 1000:
            recommendations['strategy'] = 'preview_first'
            recommendations['reason'] = 'Many lines, preview recommended before full read'
            recommendations['alternatives'] = [
                'Use content_preview for quick overview',
                'Use file_grep_with_context for specific searches'
            ]

        # File type specific recommendations
        if file_type in ['json_data', 'xml_data']:
            recommendations['alternatives'].append('Consider parsing structure instead of full text read')
        elif file_type.endswith('_code'):
            recommendations['alternatives'].append('Use grep to find specific functions/classes')
        elif file_type == 'csv_data':
            recommendations['alternatives'].append('Consider reading headers first with file_read_lines')

        return recommendations

    # ================ BASIC FILE TOOLS ================
    def file_write(
            self,
            file_path: str,
            content: str,
            encoding: str = 'utf-8',
            create_dirs: bool = True
    ) -> MCPToolResult:
        # Write text content to a workspace file.
        try:
            # 澧炲姞瀹夊叏妫€鏌ワ細涓ョ鍚?.zip 鍚庣紑鐨勬枃浠跺啓鍏ュ瓧绗覆鍐呭
            if file_path.lower().endswith(('.zip', '.rar', '.7z', '.tar', '.gz')):
                return MCPToolResult(success=False, error="Refusing to write text into an archive file.")
            # 1. 璺緞棰勫鐞嗭細鑷姩绾犳閿欒鐨?PDF 鍚庣紑
            try:
                path_obj = Path(file_path)
                parts = [p for p in path_obj.parts if p not in ('.',)]
                if parts and parts[0] == 'research' and path_obj.suffix.lower() == '.pdf':
                    new_path_obj = path_obj.with_suffix('.txt')
                    logger.info(f"Correcting research PDF text path to TXT: {file_path} -> {new_path_obj.as_posix()}")
                    file_path = new_path_obj.as_posix()
            except Exception:
                pass

            # 2. 瑙勮寖鍖栨姤鍛婄珷鑺傛枃浠跺悕 (濡?part_2.1..md -> part_2.md)
            try:
                if hasattr(self, '_normalize_report_part_path'):
                    file_path = self._normalize_report_part_path(file_path)
            except Exception:
                pass

            # 3. 娓呯悊鍐呭涓殑 Artifacts (閽堝鎶ュ憡鏂囦欢)
            try:
                path_obj = Path(file_path)
                if (".md" in file_path.lower()) and (path_obj.parent.name == "report" or "report" in file_path):
                    if hasattr(self, '_clean_report_artifacts'):
                        content = self._clean_report_artifacts(content)
            except Exception:
                pass

            # 4. 瀹夊叏鏍￠獙锛氱‘淇濅笉鍦ㄥ伐浣滅┖闂村鍐欏叆
            full_path = self._safe_join(file_path)

            if create_dirs:
                full_path.parent.mkdir(parents=True, exist_ok=True)

            # 5. 銆愭牳蹇冧慨澶嶃€戜娇鐢?"w" 妯″紡瑕嗙洊鍐欏叆锛岀‘淇濇暟鎹敮涓€鎬?
            # 鍚屾椂纭繚 content 鏄瓧绗覆锛岄槻姝㈤潪瀛楃涓茬被鍨嬪鑷村啓鍏ュけ璐?
            if not isinstance(content, str):
                content = str(content)

            # 6. 銆愪慨澶嶅弬鑰冩枃鐚贡鐮併€戠‘淇濆唴瀹逛娇鐢║TF-8缂栫爜
            # 鍏堝皾璇曟娴嬪苟淇甯歌鐨勭紪鐮侀棶棰?
            try:
                # 妫€鏌ユ槸鍚槸涔辩爜锛堝寘鍚繃澶氱殑鏇挎崲瀛楃锛?
                if content.count('\ufffd') + content.count('锟') > len(content) * 0.1:
                    logger.warning("Possible encoding issue detected: content contains many replacement characters")
            except Exception:
                pass

            # 寮哄埗浣跨敤UTF-8缂栫爜鍐欏叆
            # 淇瀛楅潰 \n 鏈瑙ｆ瀽涓虹湡瀹炴崲琛岀殑闂
            content = content.replace('\\n', '\n')
            with open(full_path, "w", encoding="utf-8", errors='replace') as f:
                f.write(content)

            file_size = full_path.stat().st_size
            return MCPToolResult(
                success=True,
                data=f"Successfully wrote {len(content)} characters to {file_path}",
                metadata={
                    'file_size': file_size,
                    'encoding': encoding,
                    'is_overwrite': True
                }
            )

        except Exception as e:
            logger.error(f"鏂囦欢鍐欏叆澶辫触 ({file_path}): {e}")
            return MCPToolResult(success=False, error=str(e))

    #==================鎵цpython鑳藉姏鐨勮剼鏈?================
    def run_python_script(self, script_path: str, timeout_seconds: int = 1800) -> Dict[str, Any]:
        # 鎵ц鎸囧畾鐨?Python 鑴氭湰鏂囦欢
        import subprocess
        import os
        import sys

        # 鑷姩灏嗗叾瀹氫綅鍒板綋鍓?session 鐨勫伐浣滅┖闂?(Workspace)
        cwd = self.workspace_path if self.workspace_path else None

        # 濡傛灉鏄浉瀵硅矾寰勶紝杞崲涓哄伐浣滅┖闂翠笅鐨勭粷瀵硅矾寰?
        if cwd and not os.path.isabs(script_path):
            script_path = os.path.join(cwd, script_path)

        if not os.path.exists(script_path):
            return {"success": False, "error": f"File not found: {script_path}"}

        try:
            result = subprocess.run(
                [sys.executable, script_path],
                capture_output=True,
                timeout=timeout_seconds,
                cwd=cwd,
            )
            def decode_output(value: bytes) -> str:
                for encoding in ("utf-8", "gb18030"):
                    try:
                        return (value or b"").decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return (value or b"").decode("utf-8", errors="replace")

            stdout = decode_output(result.stdout)
            stderr = decode_output(result.stderr)
            if result.returncode != 0:
                logger.error(
                    "Python script failed: script=%s returncode=%s stderr=%s stdout=%s",
                    script_path, result.returncode, stderr[-3000:], stdout[-1000:],
                )
            return {
                "success": result.returncode == 0,
                "stdout": stdout[-4000:],
                "stderr": stderr[-4000:],
                "returncode": result.returncode,
                "command": [sys.executable, script_path],
                "working_directory": str(cwd) if cwd else None,
            }
        except subprocess.TimeoutExpired:
            return {"success": False, "error": f"Command timed out after {timeout_seconds} seconds"}
        except Exception as e:
            return {"success": False, "error": f"Command failed: {str(e)}"}

    # ================ SEARCH TOOLS ================

    def file_grep_search(
            self,
            pattern: str,
            file_pattern: str = "*",
            recursive: bool = True,
            ignore_case: bool = False,
            max_matches: int = 100
    ) -> MCPToolResult:
        # Search for pattern in files using grep-like functionality
        try:
            import fnmatch

            flags = re.IGNORECASE if ignore_case else 0
            regex = re.compile(pattern, flags)

            matches = []
            search_path = self.workspace_path

            def _search_file(file_path: Path):
                try:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    lines = content.splitlines()

                    for line_num, line in enumerate(lines, 1):
                        if regex.search(line):
                            matches.append({
                                'file': str(file_path.relative_to(self.workspace_path)),
                                'line_number': line_num,
                                'line_content': line.strip(),
                                'match_start': regex.search(line).start() if regex.search(line) else 0
                            })

                            if len(matches) >= max_matches:
                                return False  # Stop searching

                    return True

                except Exception:
                    return True  # Continue searching other files

            # Search files
            if recursive:
                for file_path in search_path.rglob(file_pattern):
                    if file_path.is_file():
                        if not _search_file(file_path):
                            break
            else:
                for file_path in search_path.glob(file_pattern):
                    if file_path.is_file():
                        if not _search_file(file_path):
                            break

            return MCPToolResult(
                success=True,
                data=matches,
                metadata={
                    'pattern': pattern,
                    'total_matches': len(matches),
                    'truncated': len(matches) >= max_matches
                }
            )

        except Exception as e:
            logger.error(f"Grep search failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def file_find_by_name(
            self,
            name_pattern: str,
            recursive: bool = True,
            case_sensitive: bool = False,
            max_results: int = 100
    ) -> MCPToolResult:
        # Find files by name pattern
        try:
            import fnmatch

            if not case_sensitive:
                name_pattern = name_pattern.lower()

            matches = []
            search_path = self.workspace_path

            def _match_name(file_path: Path) -> bool:
                name = file_path.name
                if not case_sensitive:
                    name = name.lower()

                return fnmatch.fnmatch(name, name_pattern)

            # Search files
            if recursive:
                for file_path in search_path.rglob("*"):
                    if _match_name(file_path):
                        matches.append({
                            'name': file_path.name,
                            'path': str(file_path.relative_to(self.workspace_path)),
                            'type': 'directory' if file_path.is_dir() else 'file',
                            'size': file_path.stat().st_size if file_path.is_file() else None
                        })

                        if len(matches) >= max_results:
                            break
            else:
                for file_path in search_path.iterdir():
                    if _match_name(file_path):
                        matches.append({
                            'name': file_path.name,
                            'path': str(file_path.relative_to(self.workspace_path)),
                            'type': 'directory' if file_path.is_dir() else 'file',
                            'size': file_path.stat().st_size if file_path.is_file() else None
                        })

                        if len(matches) >= max_results:
                            break

            return MCPToolResult(
                success=True,
                data=matches,
                metadata={
                    'pattern': name_pattern,
                    'total_matches': len(matches),
                    'truncated': len(matches) >= max_results
                }
            )

        except Exception as e:
            logger.error(f"File find failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def file_read_lines(
            self,
            file_path: str,
            start_line: int = 1,
            end_line: int = None,
            max_lines: int = 1000
    ) -> MCPToolResult:
        # Read specific line ranges from a file without loading the entire file.
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            if start_line < 1:
                return MCPToolResult(
                    success=False,
                    error="start_line must be >= 1"
                )

            lines_read = []
            current_line = 0

            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    current_line += 1

                    # Skip lines before start_line
                    if current_line < start_line:
                        continue

                    # Stop if we've reached end_line
                    if end_line and current_line > end_line:
                        break

                    # Safety check for max_lines
                    if len(lines_read) >= max_lines:
                        break

                    lines_read.append({
                        'line_number': current_line,
                        'content': line.rstrip('\n\r')
                    })

            # Calculate actual end line
            actual_end_line = lines_read[-1]['line_number'] if lines_read else start_line - 1

            return MCPToolResult(
                success=True,
                data={
                    'file_path': file_path,
                    'start_line': start_line,
                    'end_line': actual_end_line,
                    'lines': lines_read,
                    'line_count': len(lines_read)
                },
                metadata={
                    'total_lines_read': len(lines_read),
                    'truncated_due_to_max_lines': len(lines_read) >= max_lines
                }
            )

        except Exception as e:
            logger.error(f"File read lines failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ BASH TOOLS ================

    def bash(
            self,
            command: str,
            timeout: int = 30,
            capture_output: bool = True,
            working_directory: str = None
    ) -> MCPToolResult:
        # Execute bash command
        try:
            # Security check - prevent dangerous commands
            dangerous_patterns = [
                r'rm\s+-rf\s+/',
                r'sudo\s+rm',
                r'mkfs',
                r'dd\s+if=.*of=/dev/',
                r'>\s*/dev/sd[a-z]',
                r'cat\s+.*>\s*/dev/sd[a-z]'
            ]

            for pattern in dangerous_patterns:
                if re.search(pattern, command, re.IGNORECASE):
                    return MCPToolResult(
                        success=False,
                        error=f"Potentially dangerous command blocked: {command}"
                    )

            # Set working directory
            cwd = self.workspace_path
            if working_directory:
                cwd = Path(working_directory)
                if not cwd.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"Working directory does not exist: {working_directory}"
                    )

            # =========================================================
            # 馃殌 缁堟瀬鏅鸿兘缂栫爜瑙ｇ爜鍣細瀹岀編鍏煎 Windows GBK 涓?UTF-8
            # =========================================================
            # 鏍稿績鏀瑰姩锛氬幓鎺?text=True 鍜?encoding锛岀洿鎺ヨ幏鍙栧簳灞傜殑鍘熺敓 bytes锛堝瓧鑺傛祦锛?
            result = subprocess.run(
                command,
                shell=True,
                cwd=str(cwd),
                capture_output=capture_output,
                timeout=timeout
            )

            # 鏅鸿兘鍙岄噸瑙ｇ爜鍑芥暟锛氬厛璇?utf-8锛屼笉琛屽氨璇?gbk锛屽潥鍐充笉涓腑鏂囷紒
            def smart_decode(b_data: bytes) -> str:
                if not b_data:
                    return ""
                try:
                    return b_data.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        return b_data.decode('gbk')
                    except UnicodeDecodeError:
                        return b_data.decode('utf-8', errors='ignore')

            # 瀹夊叏瑙ｇ爜鏍囧噯杈撳嚭鍜岄敊璇緭鍑?
            stdout_text = smart_decode(result.stdout) if capture_output and result.stdout else None
            stderr_text = smart_decode(result.stderr) if capture_output and result.stderr else None
            if result.returncode != 0:
                logger.error(
                    "Shell command failed: returncode=%s cwd=%s command=%r stderr=%s stdout=%s",
                    result.returncode, cwd, command,
                    (stderr_text or "")[-3000:], (stdout_text or "")[-1000:],
                )

            return MCPToolResult(
                success=result.returncode == 0,
                error=(stderr_text or f"Command exited with code {result.returncode}") if result.returncode else None,
                data={
                    'stdout': stdout_text,
                    'stderr': stderr_text,
                    'returncode': result.returncode,
                    'command': command
                },
                metadata={
                    'execution_time': timeout,
                    'working_directory': str(cwd)
                }
            )

        except subprocess.TimeoutExpired:
            return MCPToolResult(
                success=False,
                error=f"Command timed out after {timeout} seconds"
            )
        except Exception as e:
            logger.error(f"Bash command failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def assign_multi_tasks_to_info_seeker(
            self,
            tasks: List[Dict[str, str]],
            max_workers: int = 4
    ) -> MCPToolResult:
        # Assign multiple tasks to InformationSeekerAgents for parallel execution.
        try:
            # Validate task count (1-4 tasks)
            if not (1 <= len(tasks) <= 4):
                return MCPToolResult(
                    success=False,
                    error=f"Invalid task count ({len(tasks)}). Must assign 1-4 tasks."
                )

            # Import here to avoid circular imports
            try:
                from agents import TaskInput, create_information_seeker
            except ImportError:
                from ..agents import TaskInput, create_information_seeker

            results = []
            lock = threading.Lock()

            def process_task(task: Dict[str, str]):
                # Process a single task with thread-safe result collection
                try:
                    # Get workspace_id from task or set default
                    task_workspace_id = task.get("workspace_id")
                    if not task_workspace_id:
                        task_workspace_id = f"info_seeker_task_{int(time.time())}"

                    # Get current_task_status from task or set default
                    task_status = task.get("current_task_status")
                    if not task_status:
                        task_status = "Task assigned to InformationSeekerAgent for execution"

                    # Create TaskInput object
                    task_input = TaskInput(
                        task_content=task["task_content"],
                        task_steps_for_reference=task.get("task_steps_for_reference"),
                        deliverable_contents=task.get("deliverable_contents"),
                        current_task_status=task_status,
                        workspace_id=task_workspace_id,
                        acceptance_checking_criteria=task.get("acceptance_checking_criteria")
                    )

                    # Create and execute with info seeker agent
                    info_seeker = create_information_seeker(
                        workspace_path=str(self.workspace_path),
                    )

                    logger.info(f"Assigning task to InformationSeekerAgent: {task['task_content'][:800]}...")

                    # Execute the task
                    result = info_seeker.execute_task(task_input)

                    # Prepare response data
                    response_data = {
                        "task_assignment": {
                            "task_content": task["task_content"],
                            "task_executor": "info_seeker",
                            "workspace_id": task_workspace_id,
                            "acceptance_criteria": task.get("acceptance_checking_criteria")
                        },
                        "execution_result": {
                            "success": result.success,
                            "iterations": result.iterations,
                            "execution_time": result.execution_time,
                            "agent_name": result.agent_name
                        }
                    }

                    # Include result data if successful
                    if result.success and result.result:
                        response_data["task_result"] = result.result

                    # Include error if failed
                    if not result.success and result.error:
                        response_data["execution_result"]["error"] = result.error

                    # Include reasoning trace summary
                    if result.reasoning_trace:
                        response_data["execution_result"]["reasoning_steps"] = len([
                            step for step in result.reasoning_trace if step.get("type") == "reasoning"
                        ])
                        response_data["execution_result"]["action_steps"] = len([
                            step for step in result.reasoning_trace if step.get("type") == "action"
                        ])

                    # Thread-safe result collection
                    with lock:
                        results.append(response_data)

                    return response_data

                except Exception as e:
                    error_msg = f"Task processing failed: {str(e)}"
                    logger.error(error_msg)
                    with lock:
                        results.append({
                            "task_content": task.get("task_content", "Unknown task"),
                            "success": False,
                            "error": error_msg
                        })
                    return None

            # Execute tasks in parallel with thread pool
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = [executor.submit(process_task, task) for task in tasks]
                # Wait for all tasks to complete
                for future in futures:
                    future.result()  # Raise exceptions if any

            # Check overall success
            all_success = all(task_result.get("success", False) for task_result in results)

            return MCPToolResult(
                success=all_success,
                data={"tasks": results},
                error=None if all_success else "Some tasks failed",
                metadata={
                    "tool_name": "assign_multi_tasks_to_info_seeker",
                    "task_count": len(tasks),
                    "success_count": sum(1 for r in results if r.get("success")),
                    "failure_count": sum(1 for r in results if not r.get("success"))
                }
            )

        except Exception as e:
            logger.error(f"Multi-task assignment failed: {e}")
            return MCPToolResult(
                success=False,
                error=f"Multi-task assignment failed: {str(e)}"
            )


    def assign_task_to_info_seeker(
            self,
            task_content: str,
            task_steps_for_reference: str = None,
            deliverable_contents: str = None,
            acceptance_checking_criteria: str = None,
            workspace_id: str = None,
            current_task_status: str = None
    ) -> MCPToolResult:
        # Assign a task to the InformationSeekerAgent.
        try:
            # Import here to avoid circular imports
            try:
                from agents import TaskInput, create_information_seeker
            except ImportError:
                from ..agents import TaskInput, create_information_seeker

            # Set default workspace if not provided
            if not workspace_id:
                workspace_id = f"info_seeker_task_{int(time.time())}"

            # Set default status if not provided
            if not current_task_status:
                current_task_status = "Task assigned to InformationSeekerAgent for execution"

            # Create TaskInput object
            task_input = TaskInput(
                task_content=task_content,
                task_steps_for_reference=task_steps_for_reference,
                deliverable_contents=deliverable_contents,
                current_task_status=current_task_status,
                task_executor="info_seeker",
                workspace_id=workspace_id,
                acceptance_checking_criteria=acceptance_checking_criteria
            )

            # Create and execute with info seeker agent
            info_seeker = create_information_seeker(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to InformationSeekerAgent: {task_content[:100]}...")

            # Execute the task
            result = info_seeker.execute_task(task_input)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "task_content": task_content,
                    "task_executor": "info_seeker",
                    "workspace_id": workspace_id,
                    "acceptance_criteria": acceptance_checking_criteria
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_info_seeker",
                    "task_executor": "info_seeker",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to info seeker: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to InformationSeekerAgent: {str(e)}"
            )

    def assign_task_to_writer(
            self,
            # save_analysis_file_path: str,
            task_content: str,
            user_query: str,
            key_files: List[Dict[str, str]]
    ) -> MCPToolResult:
        # Assign a task to the WriterAgent.
        try:
            # Import here to avoid circular imports
            try:
                from agents.base_agent import WriterAgentTaskInput
                from agents import create_writer_agent
            except ImportError:
                from ..agents.base_agent import WriterAgentTaskInput
                from ..agents import create_writer_agent

            # Generate workspace ID using timestamp
            workspace_id = f"writer_task_{int(time.time())}"

            # Create WriterAgentTaskInput object
            task_input = WriterAgentTaskInput(
                # save_analysis_file_path=save_analysis_file_path,
                user_query=user_query,
                task_content=task_content,
                key_files=key_files,
                workspace_id=workspace_id,
            )

            # Create and execute with writer agent
            writer = create_writer_agent(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to WriterAgent: {task_content[:800]}...")

            # Execute the task
            result = writer.execute_task(task_input)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "task_content": task_content,
                    "task_executor": "writer",
                    "workspace_id": workspace_id,
                    "user_query": user_query,
                    "key_files_count": len(key_files)
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_writer",
                    "task_executor": "writer",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to writer: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to WriterAgent: {str(e)}"
            )

    def assign_task_to_section_writer(
            self,
            task_content: str,
            write_file_path: str,
            user_query: str,
            overall_outline: str,
            current_chapter_outline: str,
            key_files: List[Dict[str, str]]
    ) -> MCPToolResult:
        # Assign a task to the SectionWriterAgent.
        try:
            # Import here to avoid circular imports
            try:
                from agents.base_agent import SectionWriterTaskInput
                from agents.section_writer import create_section_writer
            except ImportError:
                from ..agents.base_agent import SectionWriterTaskInput
                from ..agents.section_writer import create_section_writer

            # Generate workspace ID using timestamp
            workspace_id = f"section_writer_task_{int(time.time())}"

            # Create SectionWriterTaskInput object
            task_input = SectionWriterTaskInput(
                task_content=task_content,
                write_file_path=write_file_path,
                current_chapter_outline=current_chapter_outline,
                overall_outline=overall_outline,
                user_query=user_query,
                key_files=key_files,
                workspace_id=workspace_id,
            )

            # Create and execute with section writer agent
            section_writer = create_section_writer(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to SectionWriterAgent: {write_file_path}")

            # Execute the task
            result = section_writer.execute_task(task_input, write_file_path)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "write_file_path": write_file_path,
                    "current_chapter_outline": current_chapter_outline,
                    "task_executor": "section_writer",
                    "workspace_id": workspace_id,
                    "key_files_count": len(key_files)
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_section_writer",
                    "task_executor": "section_writer",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to section writer: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to SectionWriterAgent: {str(e)}"
            )

    def assign_task_to_agent(
            self,
            task_content: str,
            task_executor: str,
            task_steps_for_reference: str = None,
            deliverable_contents: str = None,
            acceptance_checking_criteria: str = None,
            workspace_id: str = None,
            current_task_status: str = None
    ) -> MCPToolResult:
        # Generic task assignment to another agent.
        try:
            # Validate task_executor
            if task_executor not in ["info_seeker", "writer"]:
                return MCPToolResult(
                    success=False,
                    error=f"Invalid task_executor '{task_executor}'. Must be 'info_seeker' or 'writer'"
                )

            # Route to specific assignment method
            if task_executor == "info_seeker":
                return self.assign_task_to_info_seeker(
                    task_content=task_content,
                    task_steps_for_reference=task_steps_for_reference,
                    deliverable_contents=deliverable_contents,
                    acceptance_checking_criteria=acceptance_checking_criteria,
                    workspace_id=workspace_id,
                    current_task_status=current_task_status
                )
            else:  # writer
                return MCPToolResult(
                    success=False,
                    error="Writer agent assignment via assign_task_to_agent is no longer supported. Use assign_task_to_writer directly with summary and key_files parameters."
                )

        except Exception as e:
            logger.error(f"Error in generic task assignment: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task: {str(e)}"
            )

    def semantic_search(self, **kwargs) -> MCPToolResult:
        # Search semantically through system-maintained knowledge index.
        try:
            # Extract parameters
            query = kwargs.get('query', '')
            max_tokens = kwargs.get('max_tokens', 2000)
            max_results = kwargs.get('max_results', 5)
            similarity_threshold = kwargs.get('similarity_threshold', 0.7)
            filters = kwargs.get('filters', {})

            if not query:
                return MCPToolResult(
                    success=False,
                    error="query is required for semantic search"
                )

            # Check OpenAI availability
            if not hasattr(self.config, 'get_openai_client') or not self.config.get_openai_client():
                return MCPToolResult(
                    success=False,
                    error="OpenAI API key required for embeddings. Please set OPENAI_API_KEY."
                )

            # Use Faiss-based system if available for high performance
            if FAISS_AVAILABLE and get_optimized_knowledge_manager:
                try:
                    manager = get_optimized_knowledge_manager(self.config)
                    search_result = manager.search(
                        query=query,
                        max_tokens=max_tokens,
                        max_results=max_results,
                        similarity_threshold=similarity_threshold,
                        filters=filters
                    )

                    if search_result['success']:
                        return MCPToolResult(
                            success=True,
                            data=search_result
                        )
                    else:
                        # Fallback to JSON-based search on error
                        logger.warning(
                            f"Faiss search failed: {search_result.get('error')}, falling back to JSON search")
                except Exception as e:
                    logger.warning(f"Faiss search error: {e}, falling back to JSON search")

            # Fallback to JSON-based search
            logger.info("Using JSON-based search (install faiss-cpu for better performance)")
            client = self.config.get_openai_client()

            # Use system-managed index (session_knowledge.json)
            index_file = "session_knowledge.json"
            try:
                with open(index_file, 'r', encoding='utf-8') as f:
                    index_data = json.load(f)
            except FileNotFoundError:
                return MCPToolResult(
                    success=True,
                    data={
                        'query': query,
                        'results': [],
                        'total_matches': 0,
                        'message': 'No knowledge index found yet. System will build index as agents complete tasks.',
                        'search_metadata': {
                            'similarity_threshold': similarity_threshold,
                            'max_tokens_requested': max_tokens,
                            'embedding_model': 'text-embedding-3-small',
                            'vector_store': 'JSON-based (fallback)'
                        }
                    }
                )

            if not index_data:
                return MCPToolResult(
                    success=True,
                    data={
                        'query': query,
                        'results': [],
                        'total_matches': 0,
                        'message': 'Knowledge index is empty.',
                        'search_metadata': {
                            'similarity_threshold': similarity_threshold,
                            'max_tokens_requested': max_tokens,
                            'embedding_model': 'text-embedding-3-small'
                        }
                    }
                )

            # Generate query embedding
            response = client.embeddings.create(
                input=query,
                model="text-embedding-3-small"
            )
            query_embedding = response.data[0].embedding

            # Calculate similarities and apply filters
            candidate_results = []
            for item in index_data:
                # Apply filters
                if filters.get('task_name') and item['task_name'] != filters['task_name']:
                    continue
                if filters.get('file_path') and filters['file_path'] not in item['file_path']:
                    continue
                if filters.get('is_final_output') is not None and item['is_final_output'] != filters['is_final_output']:
                    continue

                # Calculate cosine similarity
                import numpy as np
                item_embedding = np.array(item['embedding'])
                query_emb = np.array(query_embedding)

                similarity = np.dot(query_emb, item_embedding) / (
                        np.linalg.norm(query_emb) * np.linalg.norm(item_embedding)
                )

                if similarity >= similarity_threshold:
                    candidate_results.append({
                        'task_name': item['task_name'],
                        'file_path': item['file_path'],
                        'file_desc': item['file_desc'],
                        'is_final_output': item['is_final_output'],
                        'chunk_index': item['chunk_index'],
                        'content': item['chunk_content'],
                        'similarity_score': float(similarity),
                        'token_count': item.get('token_count', len(item['chunk_content'].split()))
                    })

            # Sort by similarity
            candidate_results.sort(key=lambda x: x['similarity_score'], reverse=True)

            # Apply token limit - intelligent selection
            selected_results = []
            total_tokens = 0

            for result in candidate_results:
                result_tokens = result['token_count']
                if total_tokens + result_tokens <= max_tokens and len(selected_results) < max_results:
                    selected_results.append(result)
                    total_tokens += result_tokens
                elif len(selected_results) < max_results:
                    # Try to fit a shorter excerpt if we have space
                    remaining_tokens = max_tokens - total_tokens
                    if remaining_tokens > 100:  # Minimum meaningful excerpt
                        words = result['content'].split()
                        excerpt_words = words[:remaining_tokens]
                        excerpt_content = ' '.join(excerpt_words) + '...' if len(words) > remaining_tokens else result[
                            'content']

                        result_copy = result.copy()
                        result_copy['content'] = excerpt_content
                        result_copy['token_count'] = len(excerpt_words)
                        result_copy['is_excerpt'] = True

                        selected_results.append(result_copy)
                        total_tokens += len(excerpt_words)
                        break

            return MCPToolResult(
                success=True,
                data={
                    'query': query,
                    'results': selected_results,
                    'total_matches': len(candidate_results),
                    'tokens_used': total_tokens,
                    'search_metadata': {
                        'similarity_threshold': similarity_threshold,
                        'max_tokens_requested': max_tokens,
                        'max_results_requested': max_results,
                        'filters_applied': filters,
                        'embedding_model': 'text-embedding-3-small',
                        'total_candidates_found': len(candidate_results)
                    }
                }
            )

        except Exception as e:
            return MCPToolResult(
                success=False,
                error=f"Semantic search failed: {str(e)}"
            )

    def _create_text_chunks(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        # Split text into overlapping chunks for better search coverage.
        words = text.split()
        chunks = []

        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append(chunk_text)

            # Stop if we've reached the end
            if i + chunk_size >= len(words):
                break

        return chunks

    def knowledge_status(self, **kwargs) -> MCPToolResult:
        # Get status and statistics about the system-managed knowledge index.
        try:
            # Try Faiss-based system first for optimal performance
            if FAISS_AVAILABLE and get_optimized_knowledge_manager:
                try:
                    manager = get_optimized_knowledge_manager(self.config)
                    stats = manager.get_stats()

                    # Add system information
                    stats['vector_store_system'] = 'Faiss (High Performance)'
                    stats['performance'] = 'O(log n) search complexity'
                    stats['recommendation'] = 'Using optimal vector storage system'

                    return MCPToolResult(
                        success=True,
                        data=stats
                    )
                except Exception as e:
                    logger.warning(f"Faiss status error: {e}, checking JSON fallback")

            # Fallback to JSON-based system
            try:
                from knowledge.knowledge_manager import get_knowledge_manager
            except ImportError:
                from ..knowledge.knowledge_manager import get_knowledge_manager

            manager = get_knowledge_manager(self.config)
            stats = manager.get_index_stats()

            # Add performance information
            stats['vector_store_system'] = 'JSON-based (Fallback)'
            stats['performance'] = 'O(n) linear search'
            stats['recommendation'] = 'Install faiss-cpu for better performance: pip install faiss-cpu'

            return MCPToolResult(
                success=True,
                data=stats
            )

        except Exception as e:
            return MCPToolResult(
                success=False,
                error=f"Failed to get knowledge status: {str(e)}"
            )

    def search_pubmed_key_words(self, keywords, max_results=10) -> MCPToolResult:
        try:
            client = self._get_pubmed_client()
            pmids = client.search(str(keywords), max_results=max_results)
            articles = client.fetch_metadata(pmids)
            logger.info(
                "PubMed keyword search completed with one batched metadata fetch: query=%r results=%d",
                str(keywords)[:120], len(articles),
            )
            return MCPToolResult(
                success=True,
                data=articles,
                metadata={
                    "result_count": len(articles),
                    "retrieval_mode": "batch_metadata_and_abstract",
                    "ncbi_api_key_used": bool(client.api_key),
                },
            )
        except Exception as e:
            logger.warning("PubMed keyword search failed: %s", e)
            return MCPToolResult(success=False, error=f"PubMed search failed: {e}")

    def search_pubmed_advanced(
        self, term="", title="", author="", journal="",
        start_date="", end_date="", num_results=10,
    ) -> MCPToolResult:
        try:
            query = build_pubmed_query(
                term=term, title=title, author=author, journal=journal,
                start_date=start_date, end_date=end_date,
            )
            client = self._get_pubmed_client()
            pmids = client.search(query, max_results=num_results)
            articles = client.fetch_metadata(pmids)
            logger.info("PubMed advanced search completed: results=%d", len(articles))
            return MCPToolResult(
                success=True,
                data=articles,
                metadata={
                    "result_count": len(articles),
                    "retrieval_mode": "batch_metadata_and_abstract",
                    "ncbi_api_key_used": bool(client.api_key),
                },
            )
        except Exception as e:
            logger.warning("PubMed advanced search failed: %s", e)
            return MCPToolResult(success=False, error=f"PubMed advanced search failed: {e}")

    def get_pubmed_article(self, pmid) -> MCPToolResult:
        pmid = str(pmid or "").strip()
        logger.info("Retrieving selected PubMed article: PMID=%s", pmid)
        if not pmid.isdigit():
            return MCPToolResult(success=False, error="PMID must contain digits only")
        try:
            client = self._get_pubmed_client()
            records = client.fetch_metadata([pmid])
            if not records:
                return MCPToolResult(success=False, error=f"No PubMed record found for PMID {pmid}")
            record = records[0]

            pmc_id = record.get("pmc_id")
            if pmc_id:
                full_text = client.fetch_open_full_text(pmc_id)
                if full_text:
                    return MCPToolResult(success=True, data={
                        **record,
                        **full_text,
                        "full_text_available": True,
                        "access_status": "open_full_text",
                    })

            # Some PubMed-indexed papers are available through the user's
            # Elsevier entitlement even when they are not open in PMC.
            doi = record.get("doi")
            if doi and os.getenv("ELSEVIER_API_KEY", "").strip():
                try:
                    from src.tools.academic_search import fetch_sciencedirect_article
                    elsevier = fetch_sciencedirect_article(doi, id_type="doi")
                    if elsevier.get("content"):
                        return MCPToolResult(success=True, data={**record, **elsevier, "pmid": pmid})
                except Exception as exc:
                    logger.info("Elsevier fallback unavailable for PMID %s: %s", pmid, exc)

            return MCPToolResult(
                success=True,
                data={
                    **record,
                    "content": record.get("abstract", ""),
                    "full_text_available": False,
                    "access_status": "metadata_abstract_only",
                    "source_url": record.get("url", ""),
                },
                metadata={
                    "limitation": "Full text was unavailable; verified PubMed metadata and abstract were returned.",
                    "do_not_retry_full_text": True,
                },
            )
        except Exception as e:
            logger.warning("PubMed article retrieval failed for PMID %s: %s", pmid, e)
            return MCPToolResult(success=False, error=f"PubMed article retrieval failed: {e}")

    def get_sciencedirect_article(self, identifier, id_type="doi") -> MCPToolResult:
        """Retrieve an article through the official Elsevier API."""
        try:
            from src.tools.academic_search import fetch_sciencedirect_article

            data = fetch_sciencedirect_article(identifier, id_type=id_type)
            return MCPToolResult(success=True, data=data)
        except Exception as exc:
            logger.warning("Elsevier article retrieval failed: %s", exc)
            return MCPToolResult(success=False, error=f"Elsevier article retrieval failed: {exc}")

    def arxiv_search(self, query: str, max_results: int = 10) -> MCPToolResult:
        BASE_URL = "http://export.arxiv.org/api/query"
        params = {
            'search_query': query,
            'max_results': max_results,
            'sortBy': 'submittedDate',
            'sortOrder': 'descending'
        }
        try:
            response = requests.get(
                BASE_URL,
                params=params,
                timeout=(8, 20),
                verify=False,
                proxies=proxy,
            )
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.warning(f"arxiv_search timeout/network failure: {e}")
            return MCPToolResult(success=False, error=f"arxiv_search network failure: {e}")
        feed = feedparser.parse(response.content)
        papers = []
        for entry in feed.entries:
            try:
                authors = [author.name for author in entry.authors]
                published = datetime.strptime(entry.published, '%Y-%m-%dT%H:%M:%SZ')
                updated = datetime.strptime(entry.updated, '%Y-%m-%dT%H:%M:%SZ')
                pdf_url = next((link.href for link in entry.links if link.type == 'application/pdf'), '')
                papers.append(Paper(
                    paper_id=entry.id.split('/')[-1],
                    title=entry.title,
                    authors=authors,
                    abstract=entry.summary,
                    url=entry.id,
                    pdf_url=pdf_url,
                    published_date=published,
                    updated_date=updated,
                    source='arxiv',
                    categories=[tag.term for tag in entry.tags],
                    keywords=[],
                    doi=entry.get('doi', '')
                ).to_dict())
            except Exception as e:
                return MCPToolResult(success=False, error=f"鑾峰彇arxiv璁烘枃淇℃伅澶辫触!{e}")
        return MCPToolResult(success=True, data={"papers": papers})

    def _resolve_paper_save_dir(self, save_path: str, default_subdir: str) -> tuple[Path, str]:
        requested_path = save_path or default_subdir
        relative_path = requested_path.replace("\\", "/")
        if relative_path.startswith("./"):
            relative_path = relative_path[2:]
        if relative_path in ("", "."):
            relative_path = default_subdir
        save_dir = self._safe_join(relative_path)
        save_dir.mkdir(parents=True, exist_ok=True)
        workspace_relative = os.path.relpath(save_dir, self.workspace_path).replace(os.sep, "/")
        return save_dir, workspace_relative

    @staticmethod
    def _paper_cache_filename(paper_id: str) -> str:
        return paper_id.replace("/", "_").replace("\\", "_")

    def download_pdf(self, paper_id: str, save_path: str) -> str:
        pdf_url = f"https://arxiv.org/pdf/{paper_id}.pdf"
        response = requests.get(pdf_url, timeout=(8, 30), verify=False, proxies=proxy)
        response.raise_for_status()
        save_dir = Path(save_path)
        save_dir.mkdir(parents=True, exist_ok=True)
        output_file = save_dir / f"{self._paper_cache_filename(paper_id)}.pdf"
        with open(output_file, 'wb') as f:
            f.write(response.content)
        return str(output_file)

    def arxiv_read_paper(self, paper_id: str, save_path: str = "./arxiv") -> MCPToolResult:
        try:
            save_dir, relative_save_dir = self._resolve_paper_save_dir(save_path, "arxiv")
            file_stem = self._paper_cache_filename(paper_id)
            txt_path = save_dir / f"{file_stem}.txt"
            pdf_path = save_dir / f"{file_stem}.pdf"
            txt_relative_path = f"{relative_save_dir}/{file_stem}.txt"

            if txt_path.exists():
                with open(txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={"paper": content.decode('utf-8', errors='ignore')})

            if not pdf_path.exists():
                pdf_path = self.download_pdf(paper_id, str(save_dir))

            text_content = self._read_pdf_text(Path(pdf_path))

            if not text_content:
                return MCPToolResult(success=False, error="Failed to extract text from Arxiv PDF; content is empty.")

            # 馃憞 --- 鏍稿績淇锛氬己鍒朵繚瀛?TXT 骞跺悗鍙板垎鏋?--- 馃憞
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            try:
                self.document_extract(tasks=[{"file_path": txt_relative_path, "task": "Extract core info"}], max_workers=1)
            except Exception as e:
                logger.warning(f"Arxiv auto-extract failed: {e}")
            # 馃憜 --- 淇缁撴潫 --- 馃憜

            return MCPToolResult(success=True, data={"paper": text_content})

        except Exception as e:
            return MCPToolResult(success=False, error=f"鑾峰彇arxiv璁烘枃鍐呭澶辫触!{e}")
            # 涓嬭浇鏂囦欢锛坉ownload_pdf鐜板湪浼氱洿鎺ヤ繚瀛樹负.txt锛?
            txt_path = self.download_pdf(paper_id, save_path)
            with open(txt_path, 'rb') as f:
                content = f.read()
            return MCPToolResult(success=True, data={"paper": content.decode('utf-8', errors='ignore')})

        except Exception as e:
            return MCPToolResult(success=False, error=f"鑾峰彇arxiv璁烘枃鍐呭澶辫触!{e}")

    def medrxiv_search(self, query: str, max_results: int = 10, days: int = 30) -> List[Paper]:
        # Search for papers on medRxiv by category within the last N days.
        # Calculate date range: last N days
        try:
            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')

            # Format category: lowercase and replace spaces with underscores
            category = query.lower().replace(' ', '_')

            papers = []
            cursor = 0
            while len(papers) < max_results:
                url = f"{self.BASE_URL}/{start_date}/{end_date}/{cursor}"
                if category:
                    url += f"?category={category}"

                tries = 0
                while tries < self.max_retries:
                    try:
                        response = self.session.get(url, timeout=self.timeout, verify=False, proxies=proxy)
                        response.raise_for_status()
                        data = response.json()
                        collection = data.get('collection', [])
                        for item in collection:
                            date = datetime.strptime(item['date'], '%Y-%m-%d')
                            papers.append(Paper(
                                paper_id=item['doi'],
                                title=item['title'],
                                authors=item['authors'].split('; '),
                                abstract=item['abstract'],
                                url=f"https://www.medrxiv.org/content/{item['doi']}v{item.get('version', '1')}",
                                pdf_url=f"https://www.medrxiv.org/content/{item['doi']}v{item.get('version', '1')}.full.pdf",
                                published_date=date,
                                updated_date=date,
                                source="medrxiv",
                                categories=[item['category']],
                                keywords=[],
                                doi=item['doi']
                            ).to_dict())
                        if len(collection) < 100:
                            break  # No more results
                        cursor += 100
                        break  # Exit retry loop on success
                    except requests.exceptions.RequestException as e:
                        tries += 1
                        if tries == self.max_retries:
                            logger.error(f"Failed to connect to medRxiv API after {self.max_retries} attempts: {e}")
                            break
                        logger.error(f"Attempt {tries} failed, retrying...")
                else:
                    continue
                break
            return MCPToolResult(success=True, data={"paper": papers})
        except Exception as e:
            return MCPToolResult(success=False, error=f"鑾峰彇medrxiv璁烘枃鍐呭澶辫触!{e}")

    def medrxiv_download_pdf(self, paper_id: str, save_path: str) -> str:
        # Download a PDF for a given paper ID from medRxiv.
        if not paper_id:
            raise ValueError("Invalid paper_id: paper_id is empty")

        pdf_url = f"https://www.medrxiv.org/content/{paper_id}v1.full.pdf"
        tries = 0
        while tries < self.max_retries:
            try:
                # Add User-Agent to avoid potential 403 errors
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = self.session.get(pdf_url, timeout=self.timeout, headers=headers, verify=False, proxies=proxy)
                response.raise_for_status()
                save_dir = Path(save_path)
                save_dir.mkdir(parents=True, exist_ok=True)
                output_file = save_dir / f"{self._paper_cache_filename(paper_id)}.pdf"
                with open(output_file, 'wb') as f:
                    f.write(response.content)
                logger.info(f"Saved medrxiv paper {paper_id} directly as text")
                return str(output_file)
            except requests.exceptions.RequestException as e:
                tries += 1
                if tries == self.max_retries:
                    raise Exception(f"Failed to download PDF after {self.max_retries} attempts: {e}")
                logger.error(f"Attempt {tries} failed, retrying...")

    def medrxiv_read_paper(self, paper_id: str, save_path: str = "./medrxiv") -> MCPToolResult:
        # Read a paper and convert it to text format.
        try:
            save_dir, relative_save_dir = self._resolve_paper_save_dir(save_path, "medrxiv")
            file_stem = self._paper_cache_filename(paper_id)

            txt_path = save_dir / f"{file_stem}.txt"
            pdf_path = save_dir / f"{file_stem}.pdf"
            txt_relative_path = f"{relative_save_dir}/{file_stem}.txt"

            # Use cached text if it already exists.
            if txt_path.exists():
                with open(txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={"paper": content.decode('utf-8', errors='ignore')})

            # 濡傛灉涓嶅瓨鍦≒DF锛屽厛涓嬭浇
            if not pdf_path.exists():
                pdf_path = self.medrxiv_download_pdf(paper_id, str(save_dir))

            # 馃殌 鎻愬彇PDF鏂囨湰
            text_content = self._read_pdf_text(Path(pdf_path))

            if not text_content:
                return MCPToolResult(success=False, error="Failed to extract text from medrxiv PDF; content is empty.")

            # Save extracted text and run background analysis.
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text_content)

            try:
                self.document_extract(
                    tasks=[
                        {"file_path": txt_relative_path, "task": "Extract core information and references for paper writing"}],
                    max_workers=1
                )
                logger.info(f"Medrxiv background extraction wrote JSONL successfully: {txt_path}")
            except Exception as ext_err:
                logger.warning(f"Medrxiv background extraction warning: {ext_err}")

            return MCPToolResult(success=True, data={"paper": text_content})

        except Exception as e:
            logger.error(f"鑾峰彇medrxiv璁烘枃鍐呭澶辫触: {e}")
            return MCPToolResult(success=False, error=f"鑾峰彇medrxiv璁烘枃鍐呭澶辫触!{e}")


def normalize_company(param_name: str):
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            response = await func(*args, **kwargs)
            if response.get('statusCode') == 2:
                sig = inspect.signature(func)
                params = list(sig.parameters.keys())
                if param_name in params:
                    param_index = params.index(param_name)
                    if kwargs.get(param_name):
                        company_name = kwargs.get(param_name)
                    else:
                        args = list(args)
                        company_name = args[param_index]
                    n_company_name = await normalize_company_name(company_name)
                    if n_company_name and company_name != n_company_name:
                        if kwargs.get(param_name):
                            kwargs[param_name] = n_company_name
                        else:
                            args[param_index] = n_company_name
                            args = tuple(args)
                        return await func(*args, **kwargs)
            return response

        return wrapper

    return decorator


def build_pubmed_query(term=None, title=None, author=None, journal=None,
                       start_date=None, end_date=None):
    query_parts = []
    if term:
        query_parts.append(str(term).strip())
    if title:
        query_parts.append(f"{str(title).strip()}[Title]")
    if author:
        query_parts.append(f"{str(author).strip()}[Author]")
    if journal:
        query_parts.append(f"{str(journal).strip()}[Journal]")
    if start_date and end_date:
        query_parts.append(f"{start_date}:{end_date}[Date - Publication]")
    return " AND ".join(part for part in query_parts if part)


def generate_pubmed_search_url(term=None, title=None, author=None, journal=None,
                               start_date=None, end_date=None, num_results=10):
    # Generate a PubMed search URL from query fields.
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    query = quote(build_pubmed_query(
        term=term, title=title, author=author, journal=journal,
        start_date=start_date, end_date=end_date,
    ))
    params = {
        "db": "pubmed",
        "term": query,
        "retmax": num_results,
        "retmode": "xml"
    }

    return f"{base_url}?{'&'.join([f'{k}={v}' for k, v in params.items()])}"


def search_pubmed(search_url):
    # Parse PubMed IDs from search results.

    response = requests.get(search_url, verify=False, proxies=proxy)

    if response.status_code == 200:
        root = ET.fromstring(response.content)
        id_list = root.find("IdList")
        if id_list is not None:
            return [id.text for id in id_list.findall("Id")]
        else:
            logger.info("No results found.")
            return []
    else:
        logger.error(f"Error: Unable to fetch data (status code: {response.status_code})")
        return []


def get_pubmed_metadata(pmid):
    # Fetch PubMed metadata by PMID.
    url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmid}&retmode=xml"
    response = requests.get(url, verify=False, proxies=proxy)

    if response.status_code == 200:
        root = ET.fromstring(response.content)
        article = root.find(".//Article")
        if article is not None:
            title = article.find(".//ArticleTitle")
            title = title.text if title is not None else "No title available"

            abstract = article.find(".//Abstract/AbstractText")
            abstract = abstract.text if abstract is not None else "No abstract available"

            authors = []
            for author in article.findall(".//Author"):
                last_name = author.find(".//LastName")
                if last_name is not None and last_name.text:
                    authors.append(last_name.text)
            authors = ", ".join(authors) if authors else "No authors available"

            journal = article.find(".//Journal/Title")
            journal = journal.text if journal is not None else "No journal available"

            pub_date = article.find(".//PubDate/Year")
            pub_date = pub_date.text if pub_date is not None else "No publication date available"

            return {
                "PMID": pmid,
                "Title": title,
                "Authors": authors,
                "Journal": journal,
                "Publication Date": pub_date,
                "Abstract": abstract
            }
        else:
            logger.info(f"No article data found for PMID: {pmid}")
            return None
    else:
        logger.error(f"Error: Unable to fetch metadata (status code: {response.status_code})")
        return None


# ================ MCP TOOL SCHEMAS ================

MCP_TOOL_SCHEMAS = {
    "think": {
        "name": "think",
        "description": "Use the tool to think about something. It will not obtain new information or make any changes to the repository, but just log the thought. Use it when complex reasoning or brainstorming is needed.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "thought": {
                    "type": "string",
                    "description": "Your thoughts."
                }
            },
            "required": ["thought"]
        }
    },

    "reflect": {
        "name": "reflect",
        "description": "When multiple attempts yield no progress, use this tool to reflect on previous reasoning and planning, considering possible overlooked clues and exploring more possibilities. It will not obtain new information or make any changes to the repository.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "reflect": {
                    "type": "string",
                    "description": "The specific content of your reflection"
                }
            },
            "required": ["reflect"]
        }
    },

    "academic_search": {
        "name": "academic_search",
        "description": "Search structured scholarly metadata across Crossref, OpenAlex, and optionally ScienceDirect, then return source-labelled results.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "queries": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Academic search queries."
                },
                "sources": {
                    "type": "array",
                    "items": {"type": "string", "enum": ["crossref", "openalex", "sciencedirect"]},
                    "description": "Optional source list. Defaults to ACADEMIC_SEARCH_SOURCES."
                },
                "max_results_per_query": {"type": "integer", "default": 5, "minimum": 1, "maximum": 100},
                "max_workers": {"type": "integer", "default": 6, "minimum": 1, "maximum": 12}
            },
            "required": ["queries"]
        }
    },
    "batch_web_search": {
        "name": "batch_web_search",
        "description": "Search multiple queries using configurable search API with concurrent processing (no more than 8 search queries)",
        "inputSchema": {
            "type": "object",
            "properties": {
                "queries": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of search queries"
                },
                "max_results_per_query": {
                    "type": "integer",
                    "default": 4,
                    "description": "Maximum search results per query (limited to 10)"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent search requests"
                }
            },
            "required": ["queries"]
        }
    },

    "url_crawler": {
        "name": "url_crawler",
        "description": "Extract content from web pages using configurable URL crawler API. Input is a list of documents with metadata including URL and local file path for saving extracted content.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "documents": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "url": {
                                "type": "string",
                                "description": "Web page URL to extract content from"
                            },
                            "file_path": {
                                "type": "string",
                                "description": "Local path to save extracted full text content"
                            },
                            "title": {
                                "type": "string",
                                "description": "Title of the web page"
                            },
                            "time": {
                                "type": "string",
                                "description": "Publication time of the web page"
                            }
                        },
                        "required": ["url", "file_path"]
                    },
                    "description": "List of documents with metadata including URL and save path"
                },
                "max_tokens_per_url": {
                    "type": "integer",
                    "default": 4000,
                    "description": "Maximum tokens per URL result"
                },
                "include_metadata": {
                    "type": "boolean",
                    "default": True,
                    "description": "Whether to include extraction metadata"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 3,
                    "description": "Maximum number of concurrent extraction requests"
                }
            },
            "required": ["documents"]
        }
    },

    "concat_section_files": {
        "name": "concat_section_files",
        "description": "Concatenate the content of the saved section files into a single file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "final_file_path": {
                    "type": "string",
                    "description": "The final file path to save the concatenated content, save the file in the workspace **under the relative path `./report/`**, and specify the final_file_path as `./report/final_report.md`"
                },
                "section_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the saved section file"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "List of section files to concatenate"
                }
            },
            "required": ["section_files", "final_file_path"]
        }
    },

    # TODO 闇€瑕佷慨鏀箂chame鐨勬牸寮忥紝杩樻槸瀛樺湪閿欒
    "search_result_classifier": {
        "name": "search_result_classifier",
        "description": "Intelligently classify and organize search result files according to a structured outline for comprehensive long-form content generation. Analyzes files across fouer key dimensions (document time, source authority, core content, and task relevance) and assigns relevant files to appropriate outline sections. Files may be assigned to multiple sections when their content spans different topics.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "outline": {
                    "type": "string",
                    "description": "The outline here must be consistent with the content and structure of the outline generated above"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file containing research content"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "List of research files to be classified according to the outline"
                },
                "model": {
                    "type": "string",
                    "default": None,
                    "description": "AI model to use for classification and organization"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI classification (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                }
            },
            "required": ["key_files", "outline"]
        }
    },

    "document_qa": {
        "name": "document_qa",
        "description": "Answer questions based on content stored in local files. Each file has a corresponding question. Reads files and uses an AI model to answer each question using the respective file content as context.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "tasks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file (relative to workspace root)"
                            },
                            "question": {
                                "type": "string",
                                "description": "Question to ask about this file"
                            }
                        },
                        "required": ["file_path", "question"]
                    },
                    "description": "List of tasks, each containing a file path and a question"
                },
                "model": {
                    "type": "string",
                    "default": None,
                    "description": "AI model to use for generating answers"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI response (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent model API requests"
                }
            },
            "required": ["tasks"]
        }
    },

    "document_extract": {
        "name": "document_extract",
        "description": "Multi-dimensional analysis of locally stored files using AI models. Evaluates each file across four key dimensions: web page time extraction, source authority assessment, task relevance evaluation, and core content summarization (~300 words). Provides structured document analysis for research and content evaluation purposes.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "tasks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file (relative to workspace root)"
                            },
                            "task": {
                                "type": "string",
                                "description": "The content of the currently executed subtask"
                            }
                        },
                        "required": ["file_path", "task"]
                    },
                    "description": "List of tasks, each containing a file path and the current task"
                },
                "model": {
                    "type": "string",
                    "default": None,
                    "description": "AI model to use for generating answers"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI response (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent model API requests"
                }
            },
            "required": ["tasks"]
        }
    },

    "section_writer": {
        "name": "section_writer",
        "description": "Write the current chapter content based on given web information and chapter structure; also consider user questions, completed chapters, and overall outline to ensure content relevance while avoiding duplication or contradictions.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "written_chapters_summary": {
                    "type": "string",
                    "description": "The summary of the written chapters, including the content of the chapters and the reflections on the chapters. Note that this field should be concatenated with the summaries of all previously written chapters with '\\n', and do not modify the original summary. For example, if the current chapter is the third chapter, the value of this field is 'chapter 1 summary \\n chapter 2 summary'. If not, the value is set to 'No previous chapters written yet.'"
                },
                "task_content": {
                    "type": "string",
                    "description": "Detailed description of some requirements for writing the current chapter and avoidance prompts. If there are reflections from the `think` tool on previously written chapters, they can be added to this field."
                },
                "user_query": {
                    "type": "string",
                    "description": "The user query, ensure the drafted content is highly relevant to the user's inquiry."
                },
                "current_chapter_outline": {
                    "type": "string",
                    "description": "This field represents the current chapter structure to be drafted. When composing the chapter content, do not modify content and bold formatting symbols of the existing structure's titles!!!"
                },
                "overall_outline": {
                    "type": "string",
                    "description": "This field represents the overall outline of the article. When drafting the chapter content, you should consider the overall outline to ensure the chapter content is consistent with the overall outline."
                },
                "target_file_path": {
                    "type": "string",
                    "description": "The path to save the chapter content"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file containing research content"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "These files are the source materials required for drafting the current chapter."
                },
                "model": {
                    "type": "string",
                    "default": None,
                    "description": "AI model to use for classification and organization"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI classification (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 5000,
                    "description": "Maximum tokens for the AI response"
                },
            },
            "required": ["user_query", "current_chapter_outline", "overall_outline", "target_file_path", "key_files"]
        }
    },

    "download_files": {
        "name": "download_files",
        "description": "Download files from URLs to the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "urls": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of URLs to download"
                },
                "target_directory": {
                    "type": "string",
                    "description": "Directory to save files"
                },
                "overwrite": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to overwrite existing files"
                },
                "max_file_size_mb": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum file size in MB"
                }
            },
            "required": ["urls"]
        }
    },

    "process_user_uploaded_files": {
        "name": "process_user_uploaded_files",
        "description": "Process and download user-uploaded files from the Flask backend. This tool fetches files uploaded by users (e.g., PDFs, documents) and saves them to the workspace with high priority markers. Use this tool FIRST when user files are available to ensure they are analyzed before web search results.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of file IDs from user uploads"
                },
                "backend_url": {
                    "type": "string",
                    "default": "http://localhost:5000",
                    "description": "Flask backend URL"
                }
            },
            "required": ["file_ids"]
        }
    },

    "process_library_files": {
        "name": "process_library_files",
        "description": "Process and download user-selected files from the document library. This tool fetches files that users have selected from their document library and saves them to the workspace. These files are treated equally with web search results - the LLM will judge their relevance and decide whether to cite them based on task_relevance, source_authority, and information_richness dimensions. Use this tool when users have selected specific files from their document library.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of file IDs from document library"
                },
                "backend_url": {
                    "type": "string",
                    "default": "http://localhost:5000",
                    "description": "Flask backend URL"
                }
            },
            "required": ["file_ids"]
        }
    },

    "list_workspace": {
        "name": "list_workspace",
        "description": "List files and directories in the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Specify the directory path to list, using a relative path"
                },
                "recursive": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to list recursively"
                },
                "include_hidden": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to include hidden files"
                },
                "max_depth": {
                    "type": "integer",
                    "default": 3,
                    "description": "Maximum recursion depth"
                }
            },
            "required": []
        }
    },

    "str_replace_based_edit_tool": {
        "name": "str_replace_based_edit_tool",
        "description": "Create, view, and edit files with various operations",
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["create", "view", "str_replace", "insert", "append", "delete"],
                    "description": "Action to perform"
                },
                "file_path": {
                    "type": "string",
                    "description": "Path to the file"
                },
                "content": {
                    "type": "string",
                    "description": "Content for create/insert/append actions"
                },
                "old_str": {
                    "type": "string",
                    "description": "String to replace (for str_replace)"
                },
                "new_str": {
                    "type": "string",
                    "description": "Replacement string (for str_replace)"
                },
                "line_number": {
                    "type": "integer",
                    "description": "Line number for insert action"
                }
            },
            "required": ["action", "file_path"]
        }
    },

    "file_read": {
        "name": "file_read",
        "description": "Read file content",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                }
            },
            "required": ["file_path"]
        }
    },

    "load_json": {
        "name": "load_json",
        "description": "Read json format file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                }
            },
            "required": ["file_path"]
        }
    },

    "file_write": {
        "name": "file_write",
        "description": "Write content to file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "content": {
                    "type": "string",
                    "description": "Content to write"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                },
                "create_dirs": {
                    "type": "boolean",
                    "default": True,
                    "description": "Create parent directories"
                }
            },
            "required": ["file_path", "content"]
        }
    },

    "file_grep_search": {
        "name": "file_grep_search",
        "description": "Search for pattern in files",
        "inputSchema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "Regex pattern to search for"
                },
                "file_pattern": {
                    "type": "string",
                    "default": "*",
                    "description": "File pattern to search in"
                },
                "recursive": {
                    "type": "boolean",
                    "default": True,
                    "description": "Search recursively"
                },
                "ignore_case": {
                    "type": "boolean",
                    "default": False,
                    "description": "Ignore case in search"
                },
                "max_matches": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum number of matches"
                }
            },
            "required": ["pattern"]
        }
    },

    "file_find_by_name": {
        "name": "file_find_by_name",
        "description": "Find files by name pattern",
        "inputSchema": {
            "type": "object",
            "properties": {
                "name_pattern": {
                    "type": "string",
                    "description": "Name pattern to search for"
                },
                "recursive": {
                    "type": "boolean",
                    "default": True,
                    "description": "Search recursively"
                },
                "case_sensitive": {
                    "type": "boolean",
                    "default": False,
                    "description": "Case sensitive search"
                },
                "max_results": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum number of results"
                }
            },
            "required": ["name_pattern"]
        }
    },

    "bash": {
        "name": "bash",
        "description": "Execute bash command in the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "command": {
                    "type": "string",
                    "description": "Bash command to execute"
                },
                "timeout": {
                    "type": "integer",
                    "default": 30,
                    "description": "Command timeout in seconds"
                },
                "capture_output": {
                    "type": "boolean",
                    "default": True,
                    "description": "Whether to capture stdout/stderr"
                },
                "working_directory": {
                    "type": "string",
                    "description": "Working directory for command"
                }
            },
            "required": ["command"]
        }
    },

    "info_seeker_task_done": {
        "name": "info_seeker_task_done",
        "description": "Information Seeker Agent task completion reporting with information collection summary and related files.",
        "inputSchema": {
            "type": "object",
            "properties": {
                # "save_analysis_file_path": {
                #     "type": "string",
                #     "description": "The path to save the analysis file, save the analysis file in the workspace **under the relative path `./doc_analysis/`**, and specify the file path as `/doc_analysis/file_analysis.jsonl`"
                # },
                "task_summary": {
                    "type": "string",
                    "description": "Simple summary of what information has been collected for the current task and what new discoveries have been made.",
                    "format": "markdown"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file with collected content"
                            },
                        },
                        "required": ["file_path"]
                    },
                    "description": "Collect files highly relevant to this task. "
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the information gathering task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Brief analysis of task completion quality, information thoroughness, and any limitations or gaps."
                }
            },
            "required": ["task_summary", "key_files", "completion_status", "completion_analysis"]
        }
    },

    "section_writer_task_done": {
        "name": "section_writer_task_done",
        "description": "Section Writer Agent task completion reporting for chapter/section writing. Called when a chapter, section, or paragraph is completed to provide a brief overview of the written content and completion status.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "chapter_summary": {
                    "type": "string",
                    "description": "Brief summary of the content written in the current chapter/section, including main topics covered and key points addressed.",
                    "format": "markdown"
                },
                "key_topics_covered": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of main topics or themes addressed in the written chapter/section"
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the chapter/section writing task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Brief analysis of the writing task completion including: assessment of content quality, evaluation of outline adherence, identification of any challenges encountered, and overall evaluation of the writing process success."
                }
            },
            "required": ["chapter_summary", "key_topics_covered", "completion_status", "completion_analysis"]
        }
    },

    "writer_task_done": {
        "name": "writer_task_done",
        "description": "Writer Agent task completion reporting for complete long-form content. Called after all chapters/sections are written to provide a summary of the complete long article, final completion status and analysis, and the storage path of the final consolidated article.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "final_article_path": {
                    "type": "string",
                    "description": "The file path where the final article is saved."
                },
                "article_summary": {
                    "type": "string",
                    "description": "Comprehensive summary of the complete long-form article, including main themes, key points covered, and overall narrative structure.",
                    "format": "markdown"
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the complete long-form writing task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Analysis of the overall writing project completion including: assessment of article coherence and quality, evaluation of content organization and flow, identification of any challenges in the writing process, and overall evaluation of the long-form content creation success."
                }
            },
            "required": ["final_article_path", "article_summary", "completion_status", "completion_analysis"]
        }
    },

    "semantic_search": {
        "name": "semantic_search",
        "description": "Search semantically through system-maintained knowledge index using OpenAI embeddings",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query - can be natural language question or keywords"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens to return in results (controls result size)"
                },
                "max_results": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of results to return"
                },
                "similarity_threshold": {
                    "type": "number",
                    "default": 0.7,
                    "description": "Minimum similarity score (0-1) for results"
                },
                "filters": {
                    "type": "object",
                    "properties": {
                        "task_name": {
                            "type": "string",
                            "description": "Filter by specific task name"
                        },
                        "file_path": {
                            "type": "string",
                            "description": "Filter by files containing this path"
                        },
                        "is_final_output": {
                            "type": "boolean",
                            "description": "Filter by final output files only"
                        }
                    },
                    "description": "Optional filters to narrow search results"
                }
            },
            "required": ["query"]
        }
    },

    "knowledge_status": {
        "name": "knowledge_status",
        "description": "Get status and statistics about the system-managed knowledge index",
        "inputSchema": {
            "type": "object",
            "properties": {},
            "required": []
        }
    },

    "search_pubmed_key_words": {
        "name": "search_pubmed_key_words",
        "description": "Search PubMed and return batched, verified citation metadata and abstracts. No NCBI API key is required. Select only the most relevant few results for any later full-text call.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "keywords": {
                    "type": "string",
                    "description": "Search query string, only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Number of results to return (default: 10)"
                }

            },
            "required": ["keywords"]
        }
    },

    "search_pubmed_advanced": {
        "name": "search_pubmed_advanced",
        "description": "Perform an advanced search for biological articles on PubMed.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "term": {
                    "type": "string",
                    "description": "General search term, only supports english"
                },
                "title": {
                    "type": "string",
                    "description": "Search in title, only supports english"
                },
                "author": {
                    "type": "string",
                    "description": "Author name, only supports english"
                },
                "journal": {
                    "type": "string",
                    "description": "Journal name, only supports english"
                },
                "start_date": {
                    "type": "string",
                    "description": "Start date for search range (format: YYYY/MM/DD)"
                },
                "end_date": {
                    "type": "string",
                    "description": "End date for search range (format: YYYY/MM/DD)"
                },
                "num_results": {
                    "type": "integer",
                    "description": "Number of results to return (default: 10)"
                }
            },
            "required": []
        }
    },
    "get_pubmed_article": {
        "name": "get_pubmed_article",
        "description": "Retrieve one selected PubMed article by PMID. Use this only for the top 3-5 relevant search results, not every PMID. It tries open PMC/Europe PMC full text, then configured Elsevier access when a DOI exists, and otherwise returns verified metadata/abstract as a successful limited result that must not be retried.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "pmid": {
                    "type": "string",
                    "description": "PMID"
                }
            },
            "required": ["pmid"]
        }
    },
    "get_sciencedirect_article": {
        "name": "get_sciencedirect_article",
        "description": "Retrieve one selected Elsevier/ScienceDirect article through the official Elsevier API. Requires ELSEVIER_API_KEY in config/.env; full text depends on the key's institutional entitlement.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "identifier": {
                    "type": "string",
                    "description": "DOI, PII, or Scopus ID."
                },
                "id_type": {
                    "type": "string",
                    "enum": ["doi", "pii", "scopus_id"],
                    "default": "doi"
                }
            },
            "required": ["identifier"]
        }
    },
    "arxiv_search": {
        "name": "arxiv_search",
        "description": "Searcher for arXiv papers, return the metadata of papers. You can get paper_id with this function and then use it for reading paper.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query string, only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Max number of searched papers"
                },
            },
            "required": ["query"]
        }
    },
    "arxiv_read_paper": {
        "name": "arxiv_read_paper",
        "description": "Obtain Arxiv article content via paper_id. Before calling this function, first use arxiv_search to obtain the article's paper_id.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "paper_id": {
                    "type": "string",
                    "description": "arXiv paper ID"
                },
                "save_path": {
                    "type": "string",
                    "description": "Directory where the PDF is/will be saved"
                }

            },
            "required": ["paper_id"]
        }
    },
    "medrxiv_search": {
        "name": "medrxiv_search",
        "description": "Searcher for biologically relevant papers, return the metadata of papers. You can get paper_id with this function and then use it for medrxiv_read_paper.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Category name to search for (e.g., \"cardiovascular medicine\"), only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Max number of searched papers"
                },
                "days": {
                    "type": "integer",
                    "description": "Number of days to look back for papers."
                }
            },
            "required": ["query"]
        }
    },
    "medrxiv_read_paper": {
        "name": "medrxiv_read_paper",
        "description": "Obtain medrxiv article content via paper_id. Before calling this function, first use medrxiv_search to obtain the article's paper_id.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "paper_id": {
                    "type": "string",
                    "description": "medrxiv paper ID"
                },
                "save_path": {
                    "type": "string",
                    "description": "Directory where the PDF is/will be saved"
                }

            },
            "required": ["paper_id"]
        }
    },
    "file_stats": {
        "name": "file_stats",
        "description": "Get comprehensive file statistics without reading full content - perfect for deciding reading strategy",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the file (relative to workspace)"
                }
            },
            "required": ["file_path"]
        }
    },

    "file_read_lines": {
        "name": "file_read_lines",
        "description": "Read specific line ranges from a file without loading entire file - perfect for large files",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the file"
                },
                "start_line": {
                    "type": "integer",
                    "default": 1,
                    "description": "Starting line number (1-based)"
                },
                "end_line": {
                    "type": "integer",
                    "description": "Ending line number (1-based, None for end of file)"
                },
                "max_lines": {
                    "type": "integer",
                    "default": 1000,
                    "description": "Maximum number of lines to read (safety limit)"
                }
            },
            "required": ["file_path"]
        }
    },

    "run_python_script": {
            "name": "run_python_script",
            "description": "Execute a Python script file and return stdout/stderr.",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "script_path": {
                        "type": "string",
                        "description": "Path to the python script"
                    },
                    "timeout_seconds": {
                        "type": "integer",
                        "default": 1800,
                        "description": "Timeout in seconds"
                    }
                },
                "required": ["script_path"]
            }
        },

    "analyze_image": {
            "name": "analyze_image",
            "description": "Analyze an image file (png, jpg, etc.) using a Vision LLM backend and return pure text descriptions or extracted data. Use this tool explicitly when you need to understand image content, read charts, or extract metrics from pictures.",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Relative path to the image"
                    },
                    "prompt": {
                        "type": "string",
                        "description": "Specific instructions on what to extract or analyze from the image, e.g., 'Extract the visible labels, numeric values, trends, legends, and task-relevant conclusions from this chart.'",
                        "default": "Describe this image in detail. If it contains charts or tables, extract the visible values and labels."
                    }
                },
                "required": ["file_path", "prompt"]
            }
        },

    "markdown_to_pdf": {
        "name": "markdown_to_pdf",
        "description": "Convert a Markdown file to PDF format. Use this tool when you need to generate a PDF report from the markdown content.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "markdown_path": {
                    "type": "string",
                    "description": "Path to the Markdown file to convert (relative to workspace)"
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path for the PDF file (relative to workspace). If not specified, will use same name with .pdf extension"
                }
            },
            "required": ["markdown_path"]
        }
    },

    # NOTE: Task assignment tool schemas removed - these are now built-in methods of PlannerAgent
    # to avoid circular dependency issues with sub-agents trying to create MCP client connections
}


# ================ MAIN INTERFACE ================

def create_mcp_tools(workspace_path: str = None) -> MCPTools:
    # Create and return MCP tools instance.
    return MCPTools(workspace_path)


def get_tool_schemas() -> Dict[str, Any]:
    # Get all tool schemas for MCP registration.
    return MCP_TOOL_SCHEMAS


